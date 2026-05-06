"""
Generate the complete Pre-Capstone Major Assessment document.
AI-Powered Mobile Application for Cybercrime Reporting and Scam Detection in West Africa
By Wilsons Navid Wado Tiwa
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
import os

doc = Document()

DIAGRAM_DIR = r"C:\Users\LENOVO\Desktop\Capstone-Project\Major-assesment\diagrams"

# ─── Page Setup ────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# ─── Style Defaults ────────────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 2.0
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)

# ─── Page Numbers ──────────────────────────────────────────────────────────
def add_page_number(section):
    """Add centered page number to footer."""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Add page number field
    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._element.append(fldChar1)

    run2 = p.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    run2._element.append(instrText)

    run3 = p.add_run()
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run3._element.append(fldChar2)

    for r in p.runs:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(10)

def add_header_text(section, text):
    """Add running header text."""
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.italic = True

# Apply page numbers to all sections
for sec in doc.sections:
    add_page_number(sec)
    add_header_text(sec, "AI-Powered Mobile Application for Cybercrime Reporting \u2014 Pre-Capstone Proposal")

# ─── Helper Functions ──────────────────────────────────────────────────────

def add_page_break():
    doc.add_page_break()

def add_centered_text(text, size=14, bold=True, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    return p

def add_heading_text(text, size=14, bold=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    run.bold = bold
    return p

def add_subheading(text, size=12, bold=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    run.bold = bold
    return p

def add_body(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.first_line_indent = Inches(0.5)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def add_body_no_indent(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 2.0
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def add_numbered_item(number, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.left_indent = Inches(0.5)
    run = p.add_run(f"{number}. {text}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def add_hyperlink(paragraph, url, text=None):
    """Add a clickable hyperlink to a paragraph."""
    import re as _re
    display = text if text else url
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    # Font
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '24')  # 12pt = 24 half-points
    rPr.append(sz)
    # Blue color + underline
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = display
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._element.append(hyperlink)

def add_reference(text):
    """Add a reference entry, making any URL clickable."""
    import re as _re
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    # Split text on URLs and make them clickable
    url_pattern = r'(https?://[^\s,)]+)'
    parts = _re.split(url_pattern, text)
    for part in parts:
        if _re.match(url_pattern, part):
            add_hyperlink(p, part)
        else:
            run = p.add_run(part)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
    return p

def add_figure(filename, caption, width=Inches(5.5)):
    """Insert a diagram image with centered caption."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    img_path = os.path.join(DIAGRAM_DIR, filename)
    run.add_picture(img_path, width=width)
    # Caption
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after = Pt(12)
    r = cap.add_run(caption)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(10)
    r.italic = True
    return p

def add_bullet(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.left_indent = Inches(0.75)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    run = p.add_run(f"\u2022 {text}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p


# ═══════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════════════════

for _ in range(6):
    doc.add_paragraph()

add_centered_text("AI-Powered Mobile Application for Cybercrime Reporting and Scam Detection in West Africa", size=16, bold=True, space_after=24)
add_centered_text("Pre-Capstone Research Proposal", size=14, bold=False, space_after=48)
add_centered_text("Wilsons Navid Wado Tiwa", size=14, bold=True, space_after=6)
add_centered_text("African Leadership University", size=12, bold=False, space_after=6)
add_centered_text("Pre-Capstone Project", size=12, bold=False, space_after=6)
add_centered_text("Supervisor: TBD", size=12, bold=False, space_after=6)
add_centered_text("2026", size=12, bold=False)


# ═══════════════════════════════════════════════════════════════════════════
# DECLARATION
# ═══════════════════════════════════════════════════════════════════════════
add_page_break()
add_heading_text("Declaration", size=14)

add_body("I, Wilsons Navid Wado Tiwa, hereby declare that this research proposal titled \"AI-Powered Mobile Application for Cybercrime Reporting and Scam Detection in West Africa\" is my original work and has not been submitted for any other academic award at any institution. All sources of information used in this proposal have been duly acknowledged through proper referencing.")

for _ in range(4):
    doc.add_paragraph()

add_body_no_indent("Signature: ___________________________")
doc.add_paragraph()
add_body_no_indent("Date: ___________________________")


# ═══════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (placeholder)
# ═══════════════════════════════════════════════════════════════════════════
add_page_break()
add_heading_text("Table of Contents", size=14)

toc_items = [
    ("Declaration", "ii"),
    ("Table of Contents", "iii"),
    ("List of Tables", "iv"),
    ("List of Figures", "v"),
    ("List of Acronyms/Abbreviations", "vi"),
    ("", ""),
    ("CHAPTER ONE: INTRODUCTION", "1"),
    ("1.1 Introduction and Background", "1"),
    ("1.2 Problem Statement", "3"),
    ("1.3 Project\u2019s Main Objective", "5"),
    ("1.3.1 Specific Objectives", "5"),
    ("1.4 Research Questions", "5"),
    ("1.5 Project Scope", "6"),
    ("1.6 Significance and Justification", "7"),
    ("", ""),
    ("CHAPTER TWO: LITERATURE REVIEW", "9"),
    ("2.1 Introduction", "9"),
    ("2.2 Related Literature", "9"),
    ("2.3 Research Gap", "16"),
    ("2.4 Conceptual Framework", "17"),
    ("2.5 Conclusion", "18"),
    ("", ""),
    ("CHAPTER THREE: RESEARCH METHODOLOGY", "20"),
    ("3.1 Introduction", "20"),
    ("3.2 Research Design", "20"),
    ("3.3 Population and Sampling Strategy", "22"),
    ("3.4 Data Collection and Analysis", "23"),
    ("3.5 System Development and Design", "27"),
    ("3.5.1 Development Model", "27"),
    ("3.5.2 System Analysis", "28"),
    ("3.5.3 Description of the Proposed System", "29"),
    ("3.5.4 Functional and Non-Functional Requirements", "30"),
    ("3.5.5 System Architecture", "32"),
    ("3.5.6 System Design (UML Diagrams)", "33"),
    ("3.5.7 Development Tools", "36"),
    ("3.6 Ethical Considerations", "37"),
    ("", ""),
    ("References", "39"),
]

for item, page in toc_items:
    if not item:
        doc.add_paragraph()
        continue
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(item)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    if item.startswith("CHAPTER") or item in ["Declaration", "Table of Contents", "List of Tables", "List of Figures", "List of Acronyms/Abbreviations", "References"]:
        run.bold = True


# ═══════════════════════════════════════════════════════════════════════════
# LIST OF TABLES
# ═══════════════════════════════════════════════════════════════════════════
add_page_break()
add_heading_text("List of Tables", size=14)

tables_list = [
    "Table 1: Comparison of Existing Cybercrime Reporting Solutions",
    "Table 2: Functional Requirements of the Proposed System",
    "Table 3: Non-Functional Requirements of the Proposed System",
    "Table 4: Development Tools and Technologies",
    "Table 5: Scam Classification Categories and Descriptions",
]
for t in tables_list:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(t)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)


# ═══════════════════════════════════════════════════════════════════════════
# LIST OF FIGURES
# ═══════════════════════════════════════════════════════════════════════════
add_page_break()
add_heading_text("List of Figures", size=14)

figures_list = [
    "Figure 1: Agile Development Model Adapted for the Project",
    "Figure 2: System Architecture Diagram",
    "Figure 3: Entity Relationship Diagram (ERD)",
    "Figure 4: Class Diagram",
    "Figure 5: Use Case Diagram",
    "Figure 6: Convergent Parallel Mixed-Methods Design",
]
for f in figures_list:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(f)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)


# ═══════════════════════════════════════════════════════════════════════════
# LIST OF ACRONYMS
# ═══════════════════════════════════════════════════════════════════════════
add_page_break()
add_heading_text("List of Acronyms/Abbreviations", size=14)

acronyms = [
    ("AI", "Artificial Intelligence"),
    ("ANTIC", "Agence Nationale des Technologies de l\u2019Information et de la Communication"),
    ("API", "Application Programming Interface"),
    ("APA", "American Psychological Association"),
    ("APWG", "Anti-Phishing Working Group"),
    ("AU", "African Union"),
    ("BEC", "Business Email Compromise"),
    ("CNN", "Convolutional Neural Network"),
    ("EFCC", "Economic and Financial Crimes Commission"),
    ("ERD", "Entity Relationship Diagram"),
    ("GPS", "Global Positioning System"),
    ("GSM", "Global System for Mobile Communications"),
    ("IoT", "Internet of Things"),
    ("ITU", "International Telecommunication Union"),
    ("LSTM", "Long Short-Term Memory"),
    ("ML", "Machine Learning"),
    ("NCC", "Nigerian Communications Commission"),
    ("ngCERT", "Nigeria Computer Emergency Response Team"),
    ("NLP", "Natural Language Processing"),
    ("SDK", "Software Development Kit"),
    ("SIM", "Subscriber Identity Module"),
    ("SMS", "Short Message Service"),
    ("SUS", "System Usability Scale"),
    ("SVM", "Support Vector Machine"),
    ("TAM", "Technology Acceptance Model"),
    ("UI", "User Interface"),
    ("UML", "Unified Modeling Language"),
    ("URL", "Uniform Resource Locator"),
]

# Create a table for acronyms
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = "Acronym"
hdr[1].text = "Full Form"
for cell in hdr:
    for p in cell.paragraphs:
        p.runs[0].bold = True
        p.runs[0].font.name = 'Times New Roman'
        p.runs[0].font.size = Pt(12)

for abbr, full in acronyms:
    row = table.add_row().cells
    row[0].text = abbr
    row[1].text = full
    for cell in row:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(12)


# ═══════════════════════════════════════════════════════════════════════════
# CHAPTER ONE: INTRODUCTION
# ═══════════════════════════════════════════════════════════════════════════
add_page_break()
add_heading_text("CHAPTER ONE: INTRODUCTION", size=14)

add_subheading("1.1 Introduction and Background")

add_body("Cybercrime has become one of the defining security threats of our time. The expansion of digital connectivity around the world has opened the door to criminal activities that operate through electronic networks, and these activities have grown rapidly in both scale and complexity. According to the International Telecommunication Union (ITU), global cybercrime damages hit roughly $8.4 trillion in 2022, and that number is expected to climb past $10.5 trillion per year by 2025 (Morgan, 2022). Phishing, online fraud, identity theft, romance scams, financial exploitation \u2014 the list goes on. No continent, no sector, no demographic has been spared.")

add_body("Africa, in particular, has seen a dramatic spike in cybercrime over the last decade. The continent is undergoing a rapid digital shift, driven largely by rising mobile phone use and widening internet access. This shift has created real opportunities for economic growth, but it has also left millions of new users exposed to cyber threats they may not fully understand. The African Union\u2019s Cybersecurity Report put the annual cost of cybercrime to African economies at around $4 billion, a figure that represents a serious drag on socioeconomic progress (African Union, 2022). Within the continent, West Africa has emerged as a global hotspot for cybercriminal activity. Nigeria and Cameroon, in particular, rank among the hardest-hit nations (INTERPOL, 2021).")

add_body("Nigeria sits at the center of Africa\u2019s cybercrime crisis. It is the continent\u2019s largest economy and most populous country, home to over 220 million people, and it has more than 100 million internet users, the biggest online population in Africa (Internet World Stats, 2023). The country has long been associated with advance-fee fraud, often called \u201c419 scams\u201d after the relevant section of the Nigerian Criminal Code. But the playbook has evolved. Today, cybercriminals based in Nigeria are involved in business email compromise (BEC), romance scams, phishing campaigns, and cryptocurrency fraud. The Economic and Financial Crimes Commission (EFCC) reported that cybercrime-related losses in Nigeria topped $500 million in 2022, with thousands of ordinary citizens victimized each year (EFCC, 2023). Yet even those numbers likely understate the problem. The Nigeria Computer Emergency Response Team (ngCERT) has pointed out that the vast majority of incidents are never reported, due to stigma, a lack of awareness about where to turn, and the absence of easy-to-use reporting platforms (ngCERT, 2022).")

add_body("Cameroon tells a similar story, though with its own distinct characteristics. The country has roughly 15 million internet users, and its mobile money ecosystem is expanding fast. That growth has brought with it a surge in cybercrime, particularly mobile money fraud, SIM swap attacks, and social media scams (Ndifon & Achou, 2023). Cameroon\u2019s bilingual makeup, split between French-speaking and English-speaking populations, adds another layer of difficulty when it comes to public awareness campaigns. The National Agency for Information and Communication Technologies (ANTIC) has noted a steady increase in online fraud complaints, but Cameroon still lacks a centralized digital platform where citizens can report cybercrime or get immediate help (ANTIC, 2022). Making matters worse, the country\u2019s cybercrime legislation remains incomplete, leaving victims with few legal options.")

add_body("To understand why cybercrime thrives in West Africa, you have to look at the broader context. Rapid urbanization, high rates of youth unemployment, and limited digital literacy have created conditions where cybercrime flourishes, both as a form of exploitation and, disturbingly, as something some young people see as a viable economic path (Warner, 2011). Many victims, especially those in lower-income communities, simply do not have the knowledge to spot scam tactics, nor the financial cushion to recover from losses. Women, elderly individuals, and people going online for the first time are disproportionately targeted. Ironically, these same groups are the least likely to come forward and report what happened to them (Akanle et al., 2020).")

add_body("Technology has a long track record of being applied to societal problems, and crime prevention is no exception. Mobile applications have been used across sectors to facilitate incident reporting, share information, and engage communities. In public safety specifically, mobile-based reporting systems have shown real promise in enabling faster responses and improving data collection for law enforcement (Bott & Young, 2012). But when it comes to cybercrime victim reporting in West Africa, that area remains remarkably underdeveloped.")

add_body("Artificial Intelligence and Machine Learning offer genuine potential for tackling cybercrime detection and prevention. AI systems can sift through large volumes of data, pick out patterns tied to fraudulent activity, classify different scam types, and deliver real-time risk assessments (Buczak & Guven, 2016). In developed countries, machine learning has already proven its value in email filtering, financial fraud detection, and identifying phishing websites. The problem is that very little of this work has been adapted for mobile platforms built around the West African context, a context where mobile money fraud and locally tailored scam schemes are the dominant threats (Tade, 2022).")

add_body("Cross-platform frameworks like Flutter now make it practical to build applications that work smoothly on both Android and iOS, which matters a great deal in markets where Android dominates. Pair that with real-time database solutions like Firebase, and developers have the tools to create responsive apps that handle dynamic data, secure authentication, and instant notifications (Google Developers, 2023). These capabilities provide the technical foundation needed to build an accessible cybercrime reporting tool for West African users.")

add_body("There is also an ethical dimension worth considering. Combating cybercrime connects directly to one of the Millennium Project\u2019s 14 Global Challenges, specifically Challenge 12: How can transnational organized crime networks be stopped from becoming more powerful and sophisticated? Cybercrime is, by its nature, transnational. Perpetrators in West Africa routinely target victims in Europe, North America, and other parts of Africa, while also preying on local populations. Tackling this challenge calls for innovative, technology-driven approaches that give citizens the power to report incidents, access protective resources, and break the cycle of victimization (Glenn & Gordon, 2009).")

add_body("This project sets out to develop an AI-powered mobile application that allows cybercrime victims in Nigeria and Cameroon to report scams, get instant risk assessments, and access educational resources. The app will use machine learning for scam pattern detection, Flutter for cross-platform development, and Firebase for real-time data management. The goal is to fill a critical gap in cybercrime victim support across West Africa and, in doing so, contribute to the broader effort of combating cybercrime on the continent.")


# ─── 1.2 Problem Statement ────────────────────────────────────────────────
add_subheading("1.2 Problem Statement")

add_body("Even as cybercrime escalates across West Africa, victims face serious obstacles when they try to report scams or seek help. In Nigeria, the main channels for reporting are the EFCC and the Nigeria Police Force Cybercrime Unit. Both require in-person visits or formal written complaints, processes that are slow, intimidating, and effectively out of reach for many citizens, especially those living in rural areas (Adeniran, 2008). In Cameroon, ANTIC is the designated national cybersecurity body, but it has no public-facing digital reporting platform, and most citizens have never heard of it (ANTIC, 2022). The result is stark: INTERPOL\u2019s African Cyberthreat Assessment (2021) found that fewer than 20% of cybercrime incidents in Africa get formally reported. In countries with limited digital reporting infrastructure, the actual figure is almost certainly lower.")

add_body("Two existing solutions have tried to address parts of this problem. The first is ScamAdviser (www.scamadviser.com), an online tool that lets users check the trustworthiness of websites and online vendors before engaging with them. It serves a useful purpose for website verification, but it is primarily web-based, does not offer a dedicated mobile reporting function, and was never designed around the specific scam patterns common in West African markets, things like mobile money fraud, advance-fee schemes, SIM swap attacks, and phishing campaigns run through WhatsApp and SMS (ScamAdviser, 2023). The second is Trend Micro\u2019s Fraud Buster application, which uses AI to scan messages and URLs for potential threats. While effective at threat detection, it does not support victim reporting or resource provision, does not draw on region-specific scam data from West Africa, and lacks support for local languages such as Pidgin English, Yoruba, Igbo, or French, languages that are essential for reaching the people most at risk (Trend Micro, 2023).")

add_body("Both solutions fall short in three important ways. First, neither provides an integrated platform that brings together scam reporting, AI-driven risk assessment, and victim resources in a single mobile application. Second, neither was built for the West African user, meaning they do not account for low-bandwidth connectivity, multilingual populations, mobile money fraud patterns, or the need to connect with local law enforcement. Third, neither uses machine learning models trained on West African cybercrime data, which is essential for detecting the region-specific scam patterns that make up the bulk of incidents in Nigeria and Cameroon.")

add_body("The gap this project aims to fill is clear: there is no AI-powered, mobile-first platform that allows cybercrime victims in Nigeria and Cameroon to report incidents, receive intelligent scam analysis, and access targeted support, all in one place. This underreporting crisis feeds directly into the growing power of transnational organized crime networks, as outlined in Global Challenge 12. When cybercriminals can operate with near-total impunity, and when law enforcement is starved of the data it needs to identify patterns, disrupt networks, and build cases, the problem only gets worse.")


# ─── 1.3 Objectives ───────────────────────────────────────────────────────
add_subheading("1.3 Project\u2019s Main Objective")

add_body("This project aims to develop an AI-powered mobile application that allows cybercrime victims in Nigeria and Cameroon to report scams, receive instant risk assessments, and access educational and support resources. The ultimate goal is to improve cybercrime reporting rates and strengthen victim support infrastructure across West Africa.")

add_subheading("1.3.1 Specific Objectives")

add_numbered_item(1, "To design and build a cross-platform mobile application using Flutter that gives cybercrime victims in Nigeria and Cameroon an intuitive, multilingual interface for reporting scam incidents, with a target development cycle of three months.")
add_numbered_item(2, "To train a machine learning model on West African cybercrime datasets that can classify reported scams into five categories \u2014 advance-fee fraud, mobile money fraud, phishing, romance scams, and identity theft \u2014 with at least 85% accuracy.")
add_numbered_item(3, "To integrate a real-time risk assessment feature, powered by Firebase, that provides users with instant feedback on the threat level of their reported incident within 30 seconds of submission.")
add_numbered_item(4, "To build an educational resource module within the app containing at least 20 articles, guides, and videos on scam identification and prevention, available in English, French, and Pidgin English.")
add_numbered_item(5, "To evaluate the app\u2019s usability and effectiveness through pilot testing with at least 50 users in Lagos, Nigeria, and Douala, Cameroon, collecting feedback to refine the system over a two-month testing period.")


# ─── 1.4 Research Questions ───────────────────────────────────────────────
add_subheading("1.4 Research Questions")

add_numbered_item(1, "How can a mobile application be designed to offer an accessible, user-friendly platform for cybercrime victims in Nigeria and Cameroon to report scam incidents?")
add_numbered_item(2, "How accurately can machine learning algorithms detect and classify scam patterns specific to the West African context, including advance-fee fraud and mobile money scams?")
add_numbered_item(3, "How effective is a real-time, AI-powered risk assessment feature in helping users grasp the severity and nature of reported cybercrime incidents?")
add_numbered_item(4, "Which educational resources and support mechanisms are most effective in helping cybercrime victims in West Africa recover from scams and avoid future victimization?")
add_numbered_item(5, "How do pilot users in Lagos, Nigeria, and Douala, Cameroon, perceive the usability and usefulness of the application?")


# ─── 1.5 Project Scope ────────────────────────────────────────────────────
add_subheading("1.5 Project Scope")

add_body("Geographically, this project focuses on two pilot locations: Lagos, Nigeria, and Douala, Cameroon. Lagos was chosen because it is Nigeria\u2019s commercial capital and largest city, home to over 20 million residents, the highest concentration of internet users in the country, and the epicenter of both cybercrime perpetration and victimization in West Africa (NCC, 2023). Douala, Cameroon\u2019s economic hub with about 4 million residents, was selected for its rapidly growing mobile money ecosystem, its rising rates of digital fraud, and its value in representing the Francophone West African context (ANTIC, 2022). Testing in both cities allows the study to capture insights from Anglophone and Francophone populations alike.")

add_body("The pilot testing population will consist of 50 to 100 adults (aged 18 and older) across both cities. These participants will be active mobile phone users who have either experienced cybercrime or face elevated risk of victimization. In Lagos, recruitment will take place in the Ikeja and Victoria Island areas, drawing from community technology hubs, university campuses, specifically the University of Lagos and Yaba College of Technology, and local NGOs focused on digital literacy. In Douala, participants will be recruited from the Akwa and Bonaberi neighborhoods through the University of Douala, cybercaf\u00e9 networks, and community organizations.")

add_body("On the technical side, the application will be built with Flutter for cross-platform compatibility across Android and iOS, Firebase for real-time database management and user authentication, and Python-based machine learning libraries, TensorFlow and Scikit-learn, for the scam detection model. The model is scoped to five primary scam categories: advance-fee fraud (419 scams), mobile money fraud, phishing, romance scams, and identity theft. Training data will come from publicly available cybercrime datasets, supplemented with records from Nigerian and Cameroonian cybercrime reports, EFCC case files, and ANTIC incident logs.")

add_body("The project is planned to run over six months: three months for application development and machine learning model training, followed by three months of pilot testing, evaluation, and iteration. The app will support English, French, and Pidgin English to ensure accessibility for both target populations. At this stage, the project does not include direct integration with law enforcement databases, but the architecture is modular, designed so that future iterations could connect to EFCC, ANTIC, or INTERPOL systems.")


# ─── 1.6 Significance ─────────────────────────────────────────────────────
add_subheading("1.6 Significance and Justification")

add_body("If this application works as intended, it will give citizens in Nigeria and Cameroon a practical, technology-driven tool for reporting scam incidents and receiving immediate, intelligent feedback. That matters because the underreporting problem is severe. INTERPOL (2021) estimated that fewer than 20% of cybercrime cases in Africa are formally reported. A mobile-first, multilingual approach that removes the friction from reporting could meaningfully increase both the volume and quality of cybercrime data available to agencies like the EFCC and ANTIC, data they need to investigate and prosecute offenders more effectively.")

add_body("The project also speaks to the Millennium Project\u2019s Global Challenge 12 on transnational organized crime. West African cybercriminal networks are among the most sophisticated on the continent, running operations that span multiple countries and target victims worldwide (Tade, 2022). A platform that detects scam patterns and aggregates structured incident data could serve as a foundational intelligence tool, useful for mapping cybercrime trends, identifying criminal networks, and supporting coordinated law enforcement action at national, regional, and continental levels (African Union, 2022).")

add_body("Beyond enforcement, the app\u2019s educational module will equip vulnerable populations, including women, elderly users, and first-time internet adopters, with the knowledge to recognize and avoid scams, building a more digitally resilient society over time. The project stands to benefit multiple stakeholders: victims gain an accessible reporting channel and recovery resources; law enforcement agencies gain structured, analyzable incident data; researchers gain access to cybercrime trend datasets; and policymakers gain the evidence they need to push for stronger cybersecurity legislation in both Nigeria and Cameroon.")


# ═══════════════════════════════════════════════════════════════════════════
# CHAPTER TWO: LITERATURE REVIEW
# ═══════════════════════════════════════════════════════════════════════════
add_page_break()
add_heading_text("CHAPTER TWO: LITERATURE REVIEW", size=14)

# ─── 2.1 Introduction ─────────────────────────────────────────────────────
add_subheading("2.1 Introduction")

add_body("This chapter presents a review of existing literature relevant to the development of an AI-powered mobile application for cybercrime reporting and scam detection in West Africa, with a specific focus on Nigeria and Cameroon. The project is situated within the Millennium Project\u2019s Global Challenge 12, which asks: \u201cHow can transnational organized crime networks be stopped from becoming more powerful and sophisticated?\u201d (Glenn & Gordon, 2009). Cybercrime in West Africa has evolved from isolated fraud schemes into a transnational enterprise generating billions of dollars annually, with the African Union (2022) estimating the cost to African economies at approximately $4 billion per year. Fewer than 20% of cybercrime incidents in Africa are formally reported (INTERPOL, 2021), creating a data vacuum that allows criminal networks to operate with near impunity.")

add_body("The literature is organized thematically across seven areas: (1) the cybercrime landscape in West Africa, (2) artificial intelligence and machine learning for cybercrime detection, (3) mobile applications for crime reporting, (4) digital literacy and cybersecurity awareness, (5) transnational organized crime and Global Challenge 12, (6) mobile technology and cross-platform development, and (7) cybersecurity policy and legislation. For each theme, the review synthesizes key findings, identifies strengths and limitations in existing research, and connects the literature to the project\u2019s objectives.")


# ─── 2.2 Related Literature ───────────────────────────────────────────────
add_subheading("2.2 Related Literature")

# Theme 1
add_subheading("2.2.1 Cybercrime Landscape in West Africa", size=12)

add_body("Understanding the nature and scale of cybercrime in West Africa is a necessary starting point for any intervention aimed at supporting victims in the region. The literature reveals that cybercrime in West Africa is not simply a technical problem but a deeply embedded social phenomenon driven by poverty, limited digital literacy, and inadequate reporting infrastructure.")

add_body("Adeniran (2008) traces the emergence of the \u201cYahoo-boys\u201d subculture in Nigeria, arguing that economic deprivation and the glorification of fraudulent wealth have normalized cybercriminal behavior among young Nigerians. Building on this foundational work, Akanle et al. (2020) analyze the \u201cYahoo Yahoo\u201d phenomenon through a sociological lens, demonstrating that cybercrime is a symptom of systemic failures in governance and economic opportunity. They highlight that vulnerable populations, including women, the elderly, and first-time internet users, are disproportionately targeted, directly underscoring the need for a victim-centered reporting platform accessible to those most at risk.")

add_body("In Cameroon, Ndifon and Achou (2023) document the rise of mobile money fraud, SIM swap attacks, and social media scams, noting the absence of a centralized digital reporting platform and the challenges posed by the country\u2019s bilingual population. Their work is one of the few peer-reviewed analyses of Cameroonian cybercrime, and it supports the rationale for including Douala as a pilot location with multilingual support built into the application. Tade (2022) extends the analysis by examining how Nigerian cybercrime has evolved from advance-fee schemes to business email compromise and cryptocurrency fraud, emphasizing severe underreporting due to victim stigma and the absence of user-friendly reporting mechanisms. Consistent with findings from Akanle et al. (2020) and ngCERT (2022), Tade\u2019s work documents the exact problem the proposed application seeks to address.")


# Theme 2
add_subheading("2.2.2 Artificial Intelligence and Machine Learning for Cybercrime Detection", size=12)

add_body("The second body of literature examines the technical feasibility of using AI and machine learning for cybercrime detection. Buczak and Guven (2016), in a widely cited survey published in IEEE Communications Surveys and Tutorials, compare decision trees, support vector machines, neural networks, and ensemble methods for cybersecurity intrusion detection. Their algorithmic comparisons remain relevant and inform the selection of classification techniques for the proposed scam detection model.")

add_body("More recently, Salloum et al. (2022) systematically review natural language processing (NLP) techniques for phishing email detection, finding that NLP-based approaches consistently outperform rule-based methods. Similarly, Hajomer and Yunus (2023) propose a phishing detection model using NLP and deep learning that achieves high accuracy on over 50,000 web pages through word embeddings and recurrent neural networks. While both studies demonstrate the power of NLP for scam detection, they share a critical limitation: their training data is overwhelmingly English-language and Western-centric, raising questions about direct applicability to the multilingual West African context where SMS and WhatsApp-based campaigns are the dominant attack vectors (Ndifon & Achou, 2023).")

add_body("Chieloka and Ugwu (2021) provide the most directly relevant precedent by evaluating Random Forest, Gradient Boosting, and Neural Networks on Nigerian cybercrime data from law enforcement records. They find that ensemble methods achieve the highest accuracy for classifying advance-fee fraud and romance scams, corroborating Buczak and Guven (2016) in a regional setting. However, their small dataset underscores the difficulty of obtaining labeled cybercrime data in West Africa, a challenge this project must address through supplementary data collection during pilot testing.")


# Theme 3
add_subheading("2.2.3 Mobile Applications for Crime Reporting and Public Safety", size=12)

add_body("A third strand of literature examines the use of mobile platforms for crime reporting in developing countries. Bott and Young (2012) provide a conceptual foundation by examining how crowdsourcing technologies, including platforms like Kenya\u2019s Ushahidi, can improve governance and public safety by empowering communities and providing authorities with real-time data. Agushaka et al. (2021) build on this foundation by designing an intelligent crime management system integrating mobile reporting and data analytics for Nigerian law enforcement, demonstrating significant reductions in reporting time compared to paper-based methods. However, their system was designed around law enforcement workflows rather than victim needs, highlighting an opportunity for the proposed application to add AI-driven classification and victim-centered resources.")

add_body("Mwiya and Phiri (2020) propose a GSM and GIS-based crime reporting system for Zambia that works over basic mobile networks, addressing the challenge of limited internet connectivity. Their positive user acceptance testing results reinforce the importance of designing for low-bandwidth environments. Together, these studies validate the viability of mobile crime reporting in Sub-Saharan Africa while revealing that no existing platform combines AI-powered scam detection, victim reporting, and educational resources in a single cybercrime-focused application.")


# Theme 4
add_subheading("2.2.4 Digital Literacy and Cybersecurity Awareness", size=12)

add_body("The literature on digital literacy highlights severe awareness gaps that any cybercrime reporting platform must address. Kshetri (2019) analyzes Africa\u2019s cybersecurity vulnerabilities, including rapid digitization without proportional security investment and a shortage of approximately 100,000 cybersecurity professionals across the continent. This analysis contextualizes the challenges documented by Ndifon and Achou (2023) and Tade (2022) in Nigeria and Cameroon specifically.")

add_body("Okuku et al. (2015) find significant cybersecurity knowledge gaps even among Kenyan university students regarding phishing, password security, and social engineering. If regular internet users at universities cannot identify common scam techniques, awareness among the elderly and rural populations targeted by scams is likely far lower (Akanle et al., 2020). These findings underscore the need for multilingual educational content within the proposed application. Ogbanufe and Gerhart (2020) contribute a design insight by demonstrating that strong technology identity relationships lead to sustained engagement, suggesting that the application must be designed not merely to be functional but to become a trusted daily resource for users\u2019 digital safety practices.")


# Theme 5
add_subheading("2.2.5 Transnational Organized Crime and Global Challenge 12", size=12)

add_body("The Millennium Project\u2019s foundational document (Glenn & Gordon, 2009) identifies 15 Global Challenges, including Challenge 12 on stopping transnational organized crime. The document estimates that transnational organized crime generates 3% to 7% of global GDP annually, and it highlights the convergence between cybercrime and traditional organized crime. INTERPOL\u2019s African Cyberthreat Assessment (2021) provides the most authoritative evidence of the scale of the problem on the continent, finding that fewer than 20% of cybercrime incidents in Africa are formally reported and identifying Nigeria and Cameroon among the hardest-hit nations.")

add_body("The African Union (2022) complements this by documenting that only 11 of 55 member states had enacted comprehensive cybersecurity legislation, estimating cybercrime costs at $4 billion annually, and recommending investments in digital literacy and strengthened national CERTs. Together, these sources situate the proposed application within the global effort to combat organized criminal networks and demonstrate that underreporting directly enables these networks to grow in power and sophistication.")


# Theme 6
add_subheading("2.2.6 Mobile Technology and Cross-Platform Development", size=12)

add_body("The technical feasibility of the proposed application is supported by a growing body of literature on cross-platform mobile development. Google Developers (2023) documents Flutter as an open-source SDK for building natively compiled cross-platform applications from a single codebase, a critical capability for markets where Android dominates but iOS users cannot be excluded. Flutter\u2019s widget-based architecture and Dart language offer performance advantages over hybrid alternatives like React Native, though as official documentation it carries an inherent promotional bias and does not address low-bandwidth deployment challenges that Mwiya and Phiri (2020) identified as central to Sub-Saharan African contexts.")

add_body("Iliescu and Olariu (2021) validate the integration of Firebase\u2019s real-time database, authentication, and cloud messaging within a Flutter tourism application, demonstrating the technical stack\u2019s capability for handling dynamic data and user authentication. However, their tourism use case has less demanding security requirements than cybercrime reporting, and the authors do not discuss data sovereignty concerns relevant to storing sensitive victim data on Google\u2019s cloud infrastructure, an issue particularly pertinent given the African Union\u2019s (2022) emphasis on strengthening domestic digital capacity.")

add_body("Mamoun et al. (2021) address this gap more directly by developing a Flutter-based healthcare application that handles sensitive personal data, demonstrating cross-platform compatibility, data security, and real-time synchronization with high user satisfaction scores. Their work provides the closest precedent to the proposed cybercrime reporting application, as both domains require secure handling of personal information and intuitive interfaces for non-technical users. A notable tension exists in this literature between the convenience of cloud-based solutions like Firebase and the data localization preferences expressed by African cybersecurity bodies (ANTIC, 2022; ngCERT, 2022). The proposed application navigates this tension by using Firebase\u2019s encryption and security rules while designing the architecture to be modular enough for future migration to locally hosted infrastructure if regulatory requirements demand it.")


# Theme 7
add_subheading("2.2.7 Cybersecurity Policy and Legislation in Africa", size=12)

add_body("The final theme examines the institutional and policy context in which the proposed application would operate. The policy landscape in both target countries reveals a significant gap between the recognition of cybercrime as a national security threat and the institutional capacity to address it. ANTIC (2022) documents a steady increase in online fraud complaints in Cameroon, particularly mobile money fraud and social media scams, while candidly acknowledging the absence of a public-facing digital reporting platform. Cameroon\u2019s cybercrime legislation remains incomplete, leaving victims with limited legal recourse even when they do come forward. As an official government publication, ANTIC\u2019s report carries high credibility, though the inherent bias of government self-assessment means systemic shortcomings may be underemphasized.")

add_body("ngCERT (2022) paints a parallel picture in Nigeria, critically acknowledging that the vast majority of cyber incidents go unreported due to victim stigma and the absence of user-friendly reporting platforms. This echoes findings from Tade (2022), Akanle et al. (2020), and INTERPOL (2021), creating a consistent evidence base across both academic and institutional sources. There is, however, an unresolved tension in the policy literature: while both ANTIC and ngCERT call for improved reporting mechanisms, neither articulates how such mechanisms should be designed, funded, or governed. The African Union (2022) documented that only 11 of 55 member states had enacted comprehensive cybersecurity legislation, suggesting that the institutional gap is continental rather than country-specific.")

add_body("A critical observation across these policy sources is their emphasis on state-level responses, including stronger legislation, greater institutional capacity, and cross-border cooperation, while giving less attention to grassroots, technology-driven interventions. Academic sources like Akanle et al. (2020) and Adeniran (2008), by contrast, foreground the socioeconomic dimensions of cybercrime and the need for community-based approaches. Few sources bridge these institutional and grassroots perspectives, which is precisely the role the proposed application aims to fill: a technology platform that empowers individual citizens while generating structured data that strengthens institutional capacity. Together, these government reports confirm that both target countries formally recognize the need for the kind of platform this project proposes, lending institutional legitimacy to the research while also highlighting the space for innovation that official channels have not yet filled.")


# ─── 2.3 Research Gap ─────────────────────────────────────────────────────
add_subheading("2.3 Research Gap")

add_body("Despite the breadth of the reviewed literature, several significant gaps emerge. First, there is a pronounced geographic imbalance: while cybercrime in Nigeria has received considerable scholarly attention, Cameroon remains severely understudied. Ndifon and Achou (2023) represent one of the few peer-reviewed analyses of Cameroonian cybercrime, and the broader Francophone West African context is largely absent from the English-language academic literature. This project\u2019s inclusion of Douala, Cameroon, as a pilot location is a deliberate attempt to address this gap.")

add_body("Second, the AI and machine learning literature on cybercrime detection is overwhelmingly Western-centric. The NLP models reviewed by Salloum et al. (2022) and Hajomer and Yunus (2023) were trained on English-language datasets from North American and European sources. There is almost no published research on training scam detection models using West African cybercrime data, multilingual text in English, French, and Pidgin English, or scam typologies specific to the region such as mobile money fraud and SIM swap attacks. Chieloka and Ugwu (2021) represent a notable exception, but their small dataset underscores the difficulty of obtaining labeled cybercrime data in West Africa.")

add_body("Third, there is a tension in the literature between institutional and grassroots perspectives on combating cybercrime. Government and institutional sources frame the problem primarily through a law enforcement and policy lens, while academic sources emphasize the socioeconomic roots of cybercrime and the need for community-based interventions. Few sources bridge these perspectives to consider how technology-driven platforms can complement both institutional responses and community empowerment, which is precisely the role the proposed application aims to play.")

add_body("Fourth, the mobile crime reporting literature demonstrates proof of concept for citizen reporting systems in developing countries but does not address the specific requirements of cybercrime reporting, such as handling sensitive digital evidence, classifying scam types, or providing real-time AI-driven risk assessments. While the constituent elements of the proposed solution, AI-based scam detection, mobile crime reporting, and cybersecurity education, have each been validated independently, no existing study or platform combines all three into a single, mobile-first application tailored for the West African context.")


# ─── 2.4 Conceptual Framework ─────────────────────────────────────────────
add_subheading("2.4 Conceptual Framework")

add_body("Two theoretical frameworks inform the design and evaluation of the proposed application. The first is the Technology Acceptance Model (TAM), originally proposed by Davis (1989) and extended in subsequent research, which provides a lens for understanding how perceived usefulness and perceived ease of use influence adoption of new technologies. Ogbanufe and Gerhart (2020) build on this tradition by demonstrating that technology identity mediates deeper engagement, suggesting that the application must be designed to be perceived as a trusted, integral part of users\u2019 digital safety practices. This is especially important in contexts where Kshetri (2019) and Okuku et al. (2015) have documented low levels of digital trust and cybersecurity awareness.")

add_body("The second framework is the Routine Activity Theory (Cohen & Felson, 1979), which posits that crime occurs when a motivated offender, a suitable target, and the absence of a capable guardian converge. In the West African context, the \u201cmotivated offenders\u201d are well-documented by Adeniran (2008) and Tade (2022), while the \u201csuitable targets\u201d are the digitally vulnerable populations identified by Akanle et al. (2020) and Kshetri (2019). The \u201cabsence of a capable guardian\u201d is precisely the institutional gap described by INTERPOL (2021), ANTIC (2022), and ngCERT (2022). The proposed application aims to serve as a form of digital guardianship, empowering users to recognize, report, and respond to cybercrime threats.")

add_body("From a methodological standpoint, the reviewed literature employs diverse approaches: qualitative and mixed-methods research in the cybercrime landscape studies, quantitative experimental methodologies in the AI and machine learning studies, and design science research in the mobile application studies. This methodological diversity strengthens the evidence base for the project by triangulating insights from social science, computer science, and information systems research.")


# ─── 2.5 Conclusion ───────────────────────────────────────────────────────
add_subheading("2.5 Conclusion")

add_body("This literature review reveals that the proposed AI-powered mobile application addresses a clearly documented and urgent gap in West Africa\u2019s cybersecurity infrastructure. The reviewed sources collectively establish that West Africa faces a severe and growing cybercrime problem characterized by diverse scam types, disproportionate victimization of vulnerable populations, and a pervasive underreporting crisis. AI and machine learning techniques can effectively detect and classify cybercrime patterns, though region-specific training data remains a critical challenge. Mobile crime reporting platforms have demonstrated viability in Sub-Saharan Africa, and the Flutter-Firebase technical stack can meet the security and usability requirements of the proposed application.")

add_body("The key finding from this review is that while each constituent element of the proposed solution has been validated independently in the literature, no existing study or platform combines AI-based scam detection, mobile crime reporting, and cybersecurity education into a single application tailored for the West African context. This integration represents the project\u2019s primary contribution. The project will also contribute to closing the data gap by collecting and labeling cybercrime incident reports from Nigerian and Cameroonian users, generating empirical evidence on mobile cybercrime reporting in both Anglophone and Francophone West African contexts, and testing whether combining detection and prevention in a single platform can improve cybercrime resilience among vulnerable populations.")


# ═══════════════════════════════════════════════════════════════════════════
# CHAPTER THREE: RESEARCH METHODOLOGY
# ═══════════════════════════════════════════════════════════════════════════
add_page_break()
add_heading_text("CHAPTER THREE: RESEARCH METHODOLOGY", size=14)

# ─── 3.1 Introduction ─────────────────────────────────────────────────────
add_subheading("3.1 Introduction")

add_body("This chapter presents the research methodology for the AI-powered mobile application for cybercrime reporting and scam detection in West Africa. It covers the research design and development model, system analysis of existing and proposed systems, system architecture and design, development tools, and ethical considerations. The study uses a mixed-methods design that combines quantitative evaluation of the machine learning scam detection model and app usability with qualitative exploration of user experiences through interviews and open-ended survey questions. Additionally, this chapter details the software engineering approach, including the system\u2019s functional and non-functional requirements, architecture, and UML design diagrams.")


# ─── 3.2 Research Design ──────────────────────────────────────────────────
add_subheading("3.2 Research Design")

add_body("This study uses a mixed-methods research design that brings together quantitative and qualitative approaches. A mixed-methods approach is appropriate here because the project has two distinct but connected sides: building and evaluating an AI-powered mobile application, which requires quantitative measurement; and understanding how users perceive and interact with the app, which requires qualitative exploration (Creswell & Creswell, 2018).")

add_body("The quantitative side combines an experimental component with a descriptive survey. The experimental part involves training and testing machine learning models on West African cybercrime data to classify scam reports into five categories: advance-fee fraud, mobile money fraud, phishing, romance scams, and identity theft. Model performance will be measured by accuracy, precision, recall, and F1-score, with a target accuracy of at least 85%. The survey part involves giving pilot users the System Usability Scale (SUS), a 10-item questionnaire that produces a numerical usability score (Brooke, 1996).")

add_body("The qualitative side uses semi-structured interviews and open-ended survey questions to collect detailed feedback from users. Topics include how intuitive the reporting process feels, whether the AI-generated risk assessments are clear and useful, how relevant the educational resources are, and what users would change. This qualitative data adds context to the numbers and helps explain why users rate certain features the way they do.")

add_body("The study follows a convergent parallel design, meaning quantitative and qualitative data will be collected at the same time during pilot testing and then brought together during analysis (Creswell & Plano Clark, 2017). This allows checking whether what users say in interviews lines up with how they score the app on the SUS, and where those two data sources diverge is usually where the most useful insights are found.")

add_figure("fig6_mixed_methods.png", "Figure 6: Convergent Parallel Mixed-Methods Design (Creswell & Plano Clark, 2017)")

add_subheading("3.3 Population and Sampling Strategy")

add_body("The study targets adult mobile phone users (18 and older) living in Lagos, Nigeria, and Douala, Cameroon, who have either been victims of cybercrime or face a high risk of it. Lagos is Nigeria\u2019s commercial capital, with over 20 million residents and the country\u2019s highest concentration of internet users, and it has the highest rates of cybercrime victimization in the region (NCC, 2023). Douala is Cameroon\u2019s economic centre, home to roughly 4 million people, and was chosen because its mobile money market is expanding fast, digital fraud is on the rise, and it gives the study a Francophone perspective (ANTIC, 2022).")

add_body("The pilot will recruit 50 to 100 participants across both cities: roughly 30 to 50 from Lagos and 20 to 50 from Douala. Nielsen (2000) found that as few as 5 to 15 users can uncover most usability issues, but samples of 30 or more are needed for statistically meaningful SUS scores. Participants will be recruited through a combination of purposive and snowball sampling. Purposive sampling is the primary method because the study needs people who use mobile phones and have some exposure to cybercrime, so random sampling would be inefficient (Etikan et al., 2016). Snowball sampling will supplement this, especially in Douala, where early participants will refer others from their networks. Snowball sampling is particularly useful for reaching cybercrime victims, who are often reluctant to come forward on their own (Goodman, 1961).")

add_body("To participate, individuals must be at least 18 years old, own and actively use a smartphone (Android or iOS), live in Lagos or Douala, and have either experienced a cybercrime incident or face elevated risk because of frequent online transactions, heavy mobile money use, or limited digital literacy. They must also speak English, French, or Pidgin English. Anyone under 18, without a smartphone, living outside the two pilot cities, or unable to give informed consent will be excluded.")

add_subheading("3.4 Data Collection and Analysis")

add_subheading("3.4.1 Data Collection Methods", size=12)

add_body("Data collection draws on both primary and secondary sources, in line with the mixed-methods design. Four primary data collection methods will be used during the pilot:")

add_body("First, surveys and questionnaires. After using the app for a minimum of two weeks, every pilot participant will complete a questionnaire built around the System Usability Scale (SUS). The SUS is a 10-item Likert scale that yields a usability score from 0 to 100 (Brooke, 1996). Beyond the SUS items, the questionnaire will ask about demographics (age, gender, education, self-rated digital literacy), prior cybercrime experience, how useful participants found the risk assessment feature, how satisfied they were with the educational content, and whether they would use the app to report future incidents.")

add_body("Second, semi-structured interviews. Between 10 and 15 participants (5 to 8 from Lagos, 5 to 7 from Douala) will sit for one-on-one interviews lasting 20 to 30 minutes. The interviews will explore how easy or difficult the reporting process was, whether the AI risk assessments made sense, how relevant the educational materials felt, what barriers came up during use, and what participants would improve. Each interview will be conducted in the participant\u2019s preferred language and audio-recorded with their consent, then transcribed.")

add_body("Third, app usage analytics. Firebase Analytics will be embedded in the app to track behavioural data during the pilot. This includes the number of scam reports submitted, how long it takes to complete a report, how often users access educational resources, whether risk assessments are returned within the 30-second target, session length, and retention rates. These numbers give an objective counterpart to what users self-report in surveys and interviews.")

add_body("Fourth, machine learning model testing. The scam detection model will be evaluated on a labelled test set of West African cybercrime records. Classification performance will be measured for each of the five scam categories using accuracy, precision, recall, and F1-score. During the pilot, real-world accuracy will also be checked by having an expert manually review a random sample of user-submitted reports and comparing those labels to the model\u2019s predictions.")

add_body("Secondary data will come from three sources: the annotated bibliography and literature review completed earlier in this proposal; publicly available cybercrime datasets from the Anti-Phishing Working Group (APWG), the UCI Machine Learning Repository, and Kaggle fraud detection collections, supplemented with region-specific records from EFCC annual reports, ANTIC cybercrime logs, and ngCERT advisories; and institutional and policy documents from INTERPOL, the African Union, the World Bank, and national cybersecurity bodies.")

add_subheading("3.4.2 Data Collection Instruments", size=12)

add_body("Five instruments will be used: (1) the System Usability Scale (SUS) questionnaire, a 10-item questionnaire scored on a 5-point Likert scale with strong reliability at a Cronbach\u2019s alpha of 0.91 (Bangor et al., 2008); (2) a custom demographic and experience questionnaire capturing participant demographics, prior cybercrime experience, digital literacy, and opinions on specific app features; (3) a semi-structured interview guide with open-ended questions grouped by theme; (4) a Firebase Analytics dashboard that automatically tracks report submissions, feature usage, session length, and risk assessment response times; and (5) a machine learning evaluation framework using Python, Scikit-learn, and TensorFlow to compute accuracy, precision, recall, F1-score, and confusion matrices.")

add_body("The SUS needs no additional validation as it has been used and tested extensively in human-computer interaction research (Brooke, 1996). The custom questionnaire will be reviewed by the project supervisor for content validity and pre-tested with 5 to 10 people who will not be part of the main pilot. The interview guide will also be reviewed by the supervisor and piloted with 2 to 3 individuals. A French translation will be prepared for Douala participants and checked by a bilingual colleague. For the machine learning evaluation, stratified 5-fold cross-validation will produce robust performance estimates across all five scam categories (Kohavi, 1995).")

add_subheading("3.4.3 Data Analysis Methods", size=12)

add_body("Analysis will mirror the mixed-methods design, with separate quantitative and qualitative tracks that converge during interpretation. On the quantitative side, accuracy, precision, recall, and F1-score will be calculated for each scam category and for the model as a whole. A confusion matrix will show where misclassifications happen, with a benchmark of 85% overall accuracy. Each participant\u2019s SUS score will be calculated using Brooke\u2019s (1996) standard method, and a mean SUS score above 68 is generally considered above-average usability (Sauro, 2011). Descriptive statistics will also summarise the Likert-scale items on perceived usefulness and satisfaction. Firebase Analytics data will be summarised with descriptive statistics: total reports submitted, average time to finish a report, how often educational resources were accessed, and what percentage of risk assessments were delivered within the 30-second window.")

add_body("On the qualitative side, interview transcripts and open-ended survey responses will be analysed through thematic analysis, following Braun and Clarke\u2019s (2006) six phases: familiarisation with the data, generating initial codes, searching for themes, reviewing themes, defining and naming themes, and producing the report. Coding will be done manually, and the supervisor will review the codes and themes to reduce bias. During interpretation, qualitative themes will be placed alongside quantitative results in a side-by-side comparison (Creswell & Plano Clark, 2017).")

add_subheading("3.5 System Development and Design")

add_subheading("3.5.1 Development Model", size=12)

add_body("The application will be developed using an Agile development methodology, specifically the Scrum framework adapted for a single-developer research project. Agile was chosen because the project requires iterative development with continuous user feedback, flexibility to adapt features based on pilot testing results, and the ability to deliver working increments of the application at regular intervals (Schwaber & Sutherland, 2020).")

add_body("The development process is organized into six two-week sprints over the three-month development phase. Sprint 1 focuses on setting up the Flutter project structure, Firebase backend configuration, and user authentication module. Sprint 2 covers the scam reporting interface and form validation. Sprint 3 involves training and integrating the machine learning scam classification model. Sprint 4 develops the real-time risk assessment feature and push notification system. Sprint 5 builds the educational resource module with multilingual content. Sprint 6 is dedicated to integration testing, performance optimization, and preparation for pilot deployment. Each sprint concludes with a review against the research objectives and adjustments to the product backlog based on testing outcomes.")

add_figure("fig1_agile_model.png", "Figure 1: Agile Development Model Adapted for the Project")


# ─── 3.3 System Analysis ──────────────────────────────────────────────────
add_subheading("3.5.2 System Analysis", size=12)

add_body("A review of existing cybercrime reporting and scam detection solutions reveals significant limitations that the proposed system aims to address. Two systems are analyzed below.")

add_body("ScamAdviser is a web-based platform that allows users to check the trustworthiness of websites and online vendors by analyzing domain age, hosting location, and user reviews. Its strengths include a large database of reviewed websites and an easy-to-use interface. However, it is primarily web-based with no dedicated mobile reporting function, does not support victim incident reporting, was not designed for West African scam patterns such as mobile money fraud and SIM swap attacks, lacks multilingual support for African languages, and does not provide AI-driven risk assessments or educational resources (ScamAdviser, 2023).")

add_body("Trend Micro Fraud Buster is a mobile application that uses AI to scan messages and URLs for potential phishing and fraud threats. Its strengths include real-time threat detection and integration with messaging platforms. However, it does not support victim reporting or connect users to recovery resources, does not draw on region-specific scam data from West Africa, lacks support for local languages such as Pidgin English, Yoruba, Igbo, or French, and focuses exclusively on threat detection without educational or community-building features (Trend Micro, 2023).")

# Comparison table
add_body_no_indent("Table 1: Comparison of Existing Cybercrime Reporting Solutions")
table = doc.add_table(rows=8, cols=4)
table.style = 'Table Grid'

headers = ["Feature", "ScamAdviser", "Trend Micro Fraud Buster", "Proposed System"]
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.name = 'Times New Roman'
            r.font.size = Pt(10)

data = [
    ["Mobile-first design", "No (web-based)", "Yes", "Yes (Flutter cross-platform)"],
    ["Victim incident reporting", "No", "No", "Yes"],
    ["AI scam classification", "No", "Partial (URL/message scan)", "Yes (5-category ML model)"],
    ["Real-time risk assessment", "No", "Yes (threat detection)", "Yes (within 30 seconds)"],
    ["West African scam data", "No", "No", "Yes (region-specific training)"],
    ["Multilingual support", "Limited", "No", "Yes (English, French, Pidgin)"],
    ["Educational resources", "No", "No", "Yes (20+ articles/guides/videos)"],
]

for row_idx, row_data in enumerate(data):
    for col_idx, cell_text in enumerate(row_data):
        cell = table.rows[row_idx + 1].cells[col_idx]
        cell.text = cell_text
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(10)

doc.add_paragraph()  # spacing

add_subheading("3.5.3 Description of the Proposed System", size=12)

add_body("The proposed system is an AI-powered mobile application called CyberGuard West Africa, designed specifically for cybercrime victims in Nigeria and Cameroon. The application consists of five core modules:")

add_bullet("Scam Reporting Module: Allows users to submit detailed incident reports including scam type, description, evidence (screenshots, messages), date, and financial impact. The reporting form supports English, French, and Pidgin English and is designed for completion in under five minutes.")
add_bullet("AI Scam Classification Engine: A machine learning backend that automatically classifies submitted reports into one of five categories (advance-fee fraud, mobile money fraud, phishing, romance scams, identity theft) using NLP and ensemble classification methods trained on West African cybercrime datasets.")
add_bullet("Real-Time Risk Assessment Module: Powered by Firebase Cloud Functions, this module analyzes submitted reports and returns a risk assessment (Low, Medium, High, Critical) within 30 seconds, including a plain-language explanation of the threat and recommended next steps.")
add_bullet("Educational Resource Hub: A curated library of at least 20 articles, guides, and videos on scam identification, prevention strategies, and recovery steps, available in English, French, and Pidgin English. Content is organized by scam type and updated regularly.")
add_bullet("User Dashboard and Analytics: Provides users with a history of their submitted reports, risk assessment results, and personalized safety recommendations. Aggregated anonymized data feeds into a backend dashboard for researchers and law enforcement.")

add_body("The system is configured for deployment on both Android and iOS platforms through Flutter\u2019s single-codebase architecture. The backend infrastructure uses Firebase Authentication for secure user login, Cloud Firestore for real-time data storage, Firebase Cloud Functions for serverless ML model inference, and Firebase Cloud Messaging for push notifications.")

add_subheading("3.5.4 Functional and Non-Functional Requirements", size=12)

add_body_no_indent("Table 2: Functional Requirements of the Proposed System")
table = doc.add_table(rows=9, cols=3)
table.style = 'Table Grid'

fr_headers = ["ID", "Requirement", "Description"]
for i, h in enumerate(fr_headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.name = 'Times New Roman'
            r.font.size = Pt(10)

fr_data = [
    ["FR-01", "User Registration and Login", "Users can register using email or phone number and authenticate via Firebase Authentication with email verification or OTP."],
    ["FR-02", "Scam Incident Reporting", "Users can submit scam reports with category selection, free-text description, file attachments (screenshots, messages), date, and estimated financial loss."],
    ["FR-03", "AI Scam Classification", "The system automatically classifies submitted reports into one of five scam categories using the trained ML model."],
    ["FR-04", "Real-Time Risk Assessment", "Upon report submission, the system returns a risk level (Low/Medium/High/Critical) with a plain-language explanation within 30 seconds."],
    ["FR-05", "Multilingual Support", "The app interface and educational content are available in English, French, and Pidgin English."],
    ["FR-06", "Educational Resource Access", "Users can browse, search, and filter educational articles, guides, and videos organized by scam type."],
    ["FR-07", "Report History and Dashboard", "Users can view their submitted reports, risk assessment results, and personalized safety tips."],
    ["FR-08", "Push Notifications", "The system sends alerts about emerging scam trends, safety tips, and report status updates via Firebase Cloud Messaging."],
]

for row_idx, row_data in enumerate(fr_data):
    for col_idx, cell_text in enumerate(row_data):
        cell = table.rows[row_idx + 1].cells[col_idx]
        cell.text = cell_text
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(10)

doc.add_paragraph()

add_body_no_indent("Table 3: Non-Functional Requirements of the Proposed System")
table = doc.add_table(rows=7, cols=3)
table.style = 'Table Grid'

nfr_headers = ["ID", "Requirement", "Description"]
for i, h in enumerate(nfr_headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.name = 'Times New Roman'
            r.font.size = Pt(10)

nfr_data = [
    ["NFR-01", "Performance", "Risk assessment results must be returned within 30 seconds of report submission. App screens must load within 3 seconds on a 3G connection."],
    ["NFR-02", "Security", "All user data must be encrypted in transit (TLS 1.2+) and at rest. Firebase security rules must prevent unauthorized data access."],
    ["NFR-03", "Scalability", "The Firebase backend must support at least 500 concurrent users during pilot testing without performance degradation."],
    ["NFR-04", "Usability", "The app must achieve a mean System Usability Scale score of 68 or above, indicating above-average usability."],
    ["NFR-05", "Availability", "The system must maintain 99% uptime during the pilot testing period, excluding scheduled maintenance."],
    ["NFR-06", "Compatibility", "The app must run on Android 8.0+ and iOS 13.0+ devices, covering over 90% of smartphones in the target markets."],
]

for row_idx, row_data in enumerate(nfr_data):
    for col_idx, cell_text in enumerate(row_data):
        cell = table.rows[row_idx + 1].cells[col_idx]
        cell.text = cell_text
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(10)


# ─── 3.4 System Architecture ──────────────────────────────────────────────
doc.add_paragraph()
add_subheading("3.5.5 System Architecture", size=12)

add_body("The system follows a three-tier client-server architecture consisting of a presentation layer, an application logic layer, and a data layer. This architecture separates concerns, allowing independent development and testing of each tier while enabling horizontal scaling of backend services.")

add_body("The presentation layer is the Flutter mobile application running on Android and iOS devices. It handles user interaction, form input validation, and display of risk assessments and educational content. The application communicates with the backend through RESTful API calls and real-time Firestore listeners.")

add_body("The application logic layer consists of Firebase Cloud Functions that handle business logic, including routing scam reports to the ML classification engine, computing risk assessments, managing push notifications, and serving educational content. The ML model is deployed as a TensorFlow Lite model accessible via a Cloud Function endpoint, enabling serverless inference without maintaining dedicated ML infrastructure.")

add_body("The data layer uses Cloud Firestore as the primary database for storing user profiles, scam reports, risk assessments, and educational content metadata. Firebase Storage holds uploaded evidence files (screenshots, message logs). Firebase Authentication manages user identity and access control. All data is encrypted at rest and in transit, and Firestore security rules enforce role-based access control.")

add_figure("fig2_system_architecture.png", "Figure 2: System Architecture Diagram (Three-Tier Client-Server)")


# ─── 3.5 System Design ────────────────────────────────────────────────────
add_subheading("3.5.6 System Design", size=12)

add_body("This section presents the UML design diagrams for the proposed system. Three mandatory diagrams are provided: the Entity Relationship Diagram (ERD), the Class Diagram, and the Use Case Diagram.")

add_subheading("Entity Relationship Diagram (ERD)", size=12)

add_body("The ERD models the data structure of the application and the relationships between key entities. The primary entities are:")

add_bullet("User: Stores user profile information including user_id (PK), name, email, phone, language_preference, city, digital_literacy_level, and registration_date.")
add_bullet("ScamReport: Contains report_id (PK), user_id (FK), scam_category, description, evidence_url, date_of_incident, financial_loss, submission_date, and status.")
add_bullet("RiskAssessment: Stores assessment_id (PK), report_id (FK), risk_level (Low/Medium/High/Critical), confidence_score, explanation, and assessment_date.")
add_bullet("EducationalResource: Contains resource_id (PK), title, content_type (article/guide/video), scam_category, language, url, and publication_date.")
add_bullet("Notification: Stores notification_id (PK), user_id (FK), message, type (alert/tip/status_update), sent_date, and read_status.")

add_body("Key relationships: A User can submit many ScamReports (one-to-many). Each ScamReport generates exactly one RiskAssessment (one-to-one). EducationalResources are associated with scam categories but are independent of specific users. Notifications are sent to Users based on their report activity and scam trends (one-to-many).")

add_figure("fig3_erd.png", "Figure 3: Entity Relationship Diagram (ERD)")

add_subheading("Class Diagram", size=12)

add_body("The class diagram models the object-oriented structure of the application\u2019s codebase. The main classes are:")

add_bullet("UserService: Handles user registration, authentication, and profile management. Methods include registerUser(), loginUser(), updateProfile(), and getLanguagePreference().")
add_bullet("ReportService: Manages the scam reporting workflow. Methods include createReport(), uploadEvidence(), getReportHistory(), and updateReportStatus().")
add_bullet("ClassificationEngine: Interfaces with the ML model for scam classification. Methods include classifyReport(), loadModel(), preprocessText(), and getConfidenceScore().")
add_bullet("RiskAssessmentService: Computes and delivers risk assessments. Methods include generateAssessment(), calculateRiskLevel(), and formatExplanation().")
add_bullet("EducationService: Manages educational content retrieval and filtering. Methods include getResources(), filterByCategory(), filterByLanguage(), and searchResources().")
add_bullet("NotificationService: Handles push notification logic. Methods include sendAlert(), sendTip(), sendStatusUpdate(), and scheduleNotification().")

add_body("The ClassificationEngine depends on a ScamModel class that encapsulates the TensorFlow Lite model, a TextPreprocessor class for NLP pipeline operations, and a CategoryMapper class that maps model outputs to human-readable scam categories.")

add_figure("fig4_class_diagram.png", "Figure 4: Class Diagram")

add_subheading("Use Case Diagram", size=12)

add_body("The use case diagram identifies the system\u2019s actors and their interactions with the application. Three actors are identified:")

add_bullet("Victim/User: The primary actor who registers, submits scam reports, views risk assessments, accesses educational resources, manages their report history, and receives push notifications.")
add_bullet("ML Classification System: An internal system actor that receives scam reports, performs NLP preprocessing, classifies reports into scam categories, and returns confidence scores to the Risk Assessment module.")
add_bullet("Administrator/Researcher: A secondary actor who accesses the backend dashboard, reviews aggregated anonymized data, monitors system performance, and manages educational content.")

add_body("The primary use cases for the Victim/User actor are: Register/Login, Submit Scam Report, View Risk Assessment, Browse Educational Resources, View Report History, Update Language Preference, and Receive Notifications. The Submit Scam Report use case includes the Classify Scam (performed by the ML system) and Generate Risk Assessment sub-use cases. The Browse Educational Resources use case extends to Filter by Scam Category and Filter by Language.")

add_figure("fig5_use_case.png", "Figure 5: Use Case Diagram")


# ─── 3.6 Development Tools ────────────────────────────────────────────────
add_subheading("3.5.7 Development Tools", size=12)

add_body("The following tools and technologies will be used for developing the application:")

add_body_no_indent("Table 4: Development Tools and Technologies")
table = doc.add_table(rows=10, cols=3)
table.style = 'Table Grid'

dt_headers = ["Tool/Technology", "Purpose", "Justification"]
for i, h in enumerate(dt_headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.name = 'Times New Roman'
            r.font.size = Pt(10)

dt_data = [
    ["Flutter (Dart)", "Cross-platform mobile app development", "Single codebase for Android and iOS; high performance; growing community (Google Developers, 2023)"],
    ["Firebase Authentication", "User registration and login", "Supports email, phone (OTP), and social login; integrates natively with Flutter"],
    ["Cloud Firestore", "Real-time NoSQL database", "Real-time data synchronization; offline support; scalable serverless architecture"],
    ["Firebase Cloud Functions", "Serverless backend logic", "Handles ML inference, risk assessment computation, and notification dispatch without dedicated servers"],
    ["Firebase Cloud Storage", "File storage for evidence uploads", "Secure storage for screenshots and message logs with fine-grained access control"],
    ["Firebase Cloud Messaging", "Push notifications", "Reliable cross-platform notification delivery for scam alerts and safety tips"],
    ["TensorFlow / TensorFlow Lite", "ML model training and deployment", "Industry-standard ML framework; TFLite enables on-device or cloud inference for scam classification"],
    ["Scikit-learn", "ML model evaluation and preprocessing", "Standard library for classification metrics, cross-validation, and data preprocessing (Kohavi, 1995)"],
    ["Visual Studio Code", "Integrated development environment", "Lightweight IDE with excellent Flutter, Dart, and Python extension support"],
]

for row_idx, row_data in enumerate(dt_data):
    for col_idx, cell_text in enumerate(row_data):
        cell = table.rows[row_idx + 1].cells[col_idx]
        cell.text = cell_text
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(10)


# ─── 3.7 Ethical Considerations ───────────────────────────────────────────
doc.add_paragraph()
add_subheading("3.6 Ethical Considerations")

add_body("The study follows the Belmont Report\u2019s principles of respect for persons, beneficence, and justice (National Commission for the Protection of Human Subjects, 1979).")

add_body("Institutional approval will be obtained from the African Leadership University\u2019s Institutional Review Board (or its equivalent) before any pilot testing begins. No data involving human participants will be collected until that approval has been granted. Informed consent will be obtained from every participant through a consent form that explains the purpose of the study, what participation involves, what data will be collected, how it will be stored and used, and their right to withdraw at any point without consequences. In Douala, the informed consent form will be available in both English and French. Written informed consent is required before any data collection starts.")

add_body("All participant data will be kept confidential. Names and other personal identifiers will be stripped from datasets and replaced with codes during analysis and reporting. Surveys, interview transcripts, and app usage logs will be stored in encrypted files that only the researcher and supervisor can access. On the app side, Firebase Authentication and Firestore security rules will prevent unauthorized access. Participants will be told exactly what data the app collects and how it is processed.")

add_body("Some participants may find it uncomfortable to discuss their cybercrime experiences. Participation is entirely voluntary, and participants can skip any question or leave the study at any time. Where appropriate, information about local victim support services will be provided. The study does not involve any intervention that could cause physical or psychological harm; the app is a support tool, and using it during the pilot carries no risks beyond those of normal smartphone use.")

add_body("Research data will be kept for two years after the study ends, in case verification or follow-up analysis is needed. After that period, digital files will be permanently deleted and physical consent forms will be shredded. Participants will be recruited fairly through community technology hubs, universities, cybercaf\u00e9 networks, and local NGOs, without excluding anyone on the basis of gender, ethnicity, socioeconomic status, or language. The app itself supports English, French, and Pidgin English to keep access as broad as possible.")


# ═══════════════════════════════════════════════════════════════════════════
# REFERENCES
# ═══════════════════════════════════════════════════════════════════════════
add_page_break()
add_heading_text("References", size=14)

references = [
    "Adeniran, A. I. (2008). The internet and emergence of Yahoo-boys sub-culture in Nigeria. International Journal of Cyber Criminology, 2(2), 368\u2013381.",
    "African Union. (2022). African Union cybersecurity report: The state of cybersecurity in Africa. African Union Commission.",
    "Agushaka, J. O., Oyefolahan, I. O., Abisoye, O. A., & Umar, M. B. (2021). An intelligent crime management system for reporting, investigation, and detection. Fudma Journal of Sciences, 5(2), 47\u201356.",
    "Akanle, O., Adesina, J. O., & Akarah, E. P. (2020). Towards human dignity and the internet: The cybercrime (Yahoo Yahoo) phenomenon in Nigeria. African Journal of Science, Technology, Innovation and Development, 12(6), 683\u2013691. https://doi.org/10.1080/20421338.2019.1694993",
    "ANTIC. (2022). Rapport annuel sur la cybercriminalit\u00e9 au Cameroun 2021\u20132022. Agence Nationale des Technologies de l\u2019Information et de la Communication.",
    "Bangor, A., Kortum, P. T., & Miller, J. T. (2008). An empirical evaluation of the System Usability Scale. International Journal of Human-Computer Interaction, 24(6), 574\u2013594. https://doi.org/10.1080/10447310802205776",
    "Bott, M., & Young, G. (2012). The role of crowdsourcing for better governance in international development. PRAXIS: The Fletcher Journal of Human Security, 27, 47\u201370.",
    "Braun, V., & Clarke, V. (2006). Using thematic analysis in psychology. Qualitative Research in Psychology, 3(2), 77\u2013101. https://doi.org/10.1191/1478088706qp063oa",
    "Brooke, J. (1996). SUS: A \u201cquick and dirty\u201d usability scale. In P. W. Jordan, B. Thomas, B. A. Weerdmeester, & I. L. McClelland (Eds.), Usability evaluation in industry (pp. 189\u2013194). Taylor & Francis.",
    "Buczak, A. L., & Guven, E. (2016). A survey of data mining and machine learning methods for cyber security intrusion detection. IEEE Communications Surveys & Tutorials, 18(2), 1153\u20131176. https://doi.org/10.1109/COMST.2015.2494502",
    "Chieloka, O. E., & Ugwu, C. (2021). Application of machine learning in cybercrime detection in Nigeria. Journal of Computer Science and Its Application, 28(1), 14\u201327.",
    "Cohen, L. E., & Felson, M. (1979). Social change and crime rate trends: A routine activity approach. American Sociological Review, 44(4), 588\u2013608.",
    "Creswell, J. W., & Creswell, J. D. (2018). Research design: Qualitative, quantitative, and mixed methods approaches (5th ed.). SAGE Publications.",
    "Creswell, J. W., & Plano Clark, V. L. (2017). Designing and conducting mixed methods research (3rd ed.). SAGE Publications.",
    "Davis, F. D. (1989). Perceived usefulness, perceived ease of use, and user acceptance of information technology. MIS Quarterly, 13(3), 319\u2013340. https://doi.org/10.2307/249008",
    "EFCC. (2023). Economic and Financial Crimes Commission annual report 2022. Federal Republic of Nigeria.",
    "Etikan, I., Musa, S. A., & Alkassim, R. S. (2016). Comparison of convenience sampling and purposive sampling. American Journal of Theoretical and Applied Statistics, 5(1), 1\u20134. https://doi.org/10.11648/j.ajtas.20160501.11",
    "Glenn, J. C., & Gordon, T. J. (2009). Futures research methodology version 3.0. The Millennium Project.",
    "Goodman, L. A. (1961). Snowball sampling. The Annals of Mathematical Statistics, 32(1), 148\u2013170.",
    "Google Developers. (2023). Flutter documentation: Build apps for any screen. https://flutter.dev/docs",
    "Hajomer, A. A. E., & Yunus, F. (2023). A phishing-attack-detection model using natural language processing and deep learning. Applied Sciences, 13(9), 5275. https://doi.org/10.3390/app13095275",
    "Iliescu, D., & Olariu, C. (2021). Improving the tourists\u2019 experiences: Application of Firebase and Flutter technologies in mobile applications development process. In Proceedings of the 2021 International Conference on Engineering Technologies and Computer Science (pp. 45\u201351). IEEE. https://doi.org/10.1109/EnT52731.2021.9623025",
    "Internet World Stats. (2023). Africa internet users, 2023 population and Facebook statistics. https://www.internetworldstats.com/stats1.htm",
    "INTERPOL. (2021). African cyberthreat assessment report 2021. INTERPOL.",
    "Kohavi, R. (1995). A study of cross-validation and bootstrap for accuracy estimation and model selection. Proceedings of the 14th International Joint Conference on Artificial Intelligence, 2, 1137\u20131143.",
    "Kshetri, N. (2019). Cybercrime and cybersecurity in Africa. Journal of Global Information Technology Management, 22(2), 77\u201381. https://doi.org/10.1080/1097198X.2019.1603527",
    "Mamoun, R., Alqudah, S., & Jaradat, A. (2021). Design and development of a mobile healthcare application prototype using Flutter. International Journal of Interactive Mobile Technologies, 15(17), 4\u201320.",
    "Morgan, S. (2022). Cybercrime to cost the world $10.5 trillion annually by 2025. Cybersecurity Ventures. https://cybersecurityventures.com/cybercrime-damage-costs-10-trillion-by-2025/",
    "Mwiya, B., & Phiri, J. (2020). Public crime reporting and monitoring system model using GSM and GIS technologies: A case of Zambia Police Service. International Journal of Advanced Computer Science and Applications, 11(4), 356\u2013363.",
    "National Commission for the Protection of Human Subjects. (1979). The Belmont Report: Ethical principles and guidelines for the protection of human subjects of research. U.S. Department of Health and Human Services.",
    "NCC. (2023). Subscriber statistics: Annual report 2022. Nigerian Communications Commission.",
    "Ndifon, R. A., & Achou, B. F. (2023). Cybercrime and mobile money fraud in Cameroon: Trends, challenges, and policy implications. African Journal of Criminology and Justice Studies, 16(1), 45\u201363.",
    "ngCERT. (2022). Nigeria Computer Emergency Response Team annual report 2021\u20132022. Federal Republic of Nigeria.",
    "Nielsen, J. (2000). Why you only need to test with 5 users. Nielsen Norman Group. https://www.nngroup.com/articles/why-you-only-need-to-test-with-5-users/",
    "Ogbanufe, O., & Gerhart, N. (2020). The mediating influence of smartwatch identity on deep use and innovative individual use. Information Systems Frontiers, 22(4), 897\u2013912. https://doi.org/10.1007/s10796-019-09910-w",
    "Okuku, A., Renaud, K., & Vasileiou, I. (2015). Cybersecurity awareness in developing nations: A study of university students in Kenya. In Proceedings of the 2015 Information Security for South Africa Conference (pp. 1\u20137). IEEE.",
    "Salloum, S., Gaber, T., Vadera, S., & Shaalan, K. (2022). A systematic literature review on phishing email detection using natural language processing techniques. IEEE Access, 10, 65703\u201365727.",
    "Sauro, J. (2011). A practical guide to the System Usability Scale: Background, benchmarks, and best practices. Measuring Usability LLC.",
    "ScamAdviser. (2023). About ScamAdviser. https://www.scamadviser.com/about",
    "Schwaber, K., & Sutherland, J. (2020). The Scrum guide: The definitive guide to Scrum: The rules of the game. Scrum.org.",
    "Tade, O. (2022). Internet fraud and the challenges of cybercrime in Nigeria. In Handbook of Research on Cybercrime and Information Technology (pp. 112\u2013130). IGI Global.",
    "Trend Micro. (2023). Trend Micro Fraud Buster. https://www.trendmicro.com/en_us/forHome/products/fraud-buster.html",
    "Warner, J. (2011). Understanding cyber-crime in Ghana: A view from below. International Journal of Cyber Criminology, 5(1), 736\u2013749.",
]

for ref in references:
    add_reference(ref)


# ═══════════════════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════════════════
output_path = r"C:\Users\LENOVO\Desktop\Capstone-Project\Major-assesment\WilsonsNavidWadoTiwa_Pre-Capstone_Research_Proposal.docx"
doc.save(output_path)
print(f"Document saved to: {output_path}")
print("Done!")
