import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 2.0

# Set margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('AI-Powered Mobile Application for Cybercrime Reporting and Scam Detection in West Africa')
run.bold = True
run.font.size = Pt(14)
run.font.name = 'Arial'

doc.add_paragraph()

# Author info
for line in ['Wilsons Navid Wado Tiwa', 'African Leadership University', 'Pre-Capstone Project', 'Supervisor: Unknown']:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

doc.add_paragraph()

# Chapter heading
h = doc.add_paragraph()
h.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = h.add_run('CHAPTER FOUR: RESEARCH METHODOLOGY')
run.bold = True
run.font.size = Pt(14)
run.font.name = 'Arial'

doc.add_paragraph()

def add_heading_text(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(14)
        run.font.name = 'Arial'
    else:
        run.font.size = Pt(12)
        run.font.name = 'Arial'
    return p

def add_body(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.paragraph_format.line_spacing = 2.0
    return p

# ============================================================
# 4.0 Introduction
# ============================================================
add_heading_text(doc, '4.0 Introduction')
add_body(doc,
    'This chapter describes the research methodology that will guide the development and evaluation '
    'of the AI-powered mobile application for cybercrime reporting and scam detection in West Africa. '
    'The methodology outlines the research design, target population and sampling strategy, data '
    'collection methods and instruments, data analysis techniques, and ethical considerations. The '
    'study adopts a mixed-methods research design, combining quantitative and qualitative approaches '
    'to address the research objectives comprehensively. The quantitative component focuses on '
    'measuring the technical performance of the machine learning scam detection model and evaluating '
    'user experience through standardized metrics, while the qualitative component captures in-depth '
    'user perspectives on the application\u2019s usability, relevance, and potential impact on cybercrime '
    'reporting behaviour in Nigeria and Cameroon.')

# ============================================================
# 4.1 Research Design
# ============================================================
add_heading_text(doc, '4.1 Research Design')
add_body(doc,
    'This study employs a mixed-methods research design, integrating both quantitative and '
    'qualitative approaches within a single research framework. The mixed-methods approach was '
    'selected because the project involves two distinct but interconnected dimensions: (1) the '
    'technical development and performance evaluation of an AI-powered mobile application, which '
    'requires quantitative measurement, and (2) the assessment of user experience, perceived '
    'usefulness, and behavioural intent, which requires qualitative exploration (Creswell & '
    'Creswell, 2018).')

add_body(doc,
    'The quantitative component of the study adopts an experimental and descriptive survey design. '
    'The experimental element involves training and testing machine learning models on West African '
    'cybercrime datasets to classify scam reports into five categories: advance-fee fraud, mobile '
    'money fraud, phishing, romance scams, and identity theft. Model performance will be evaluated '
    'using standard classification metrics, including accuracy, precision, recall, and F1-score, '
    'with a target accuracy of at least 85%. The descriptive survey element involves administering '
    'the System Usability Scale (SUS) questionnaire to pilot users to generate numerical scores on '
    'the application\u2019s usability (Brooke, 1996).')

add_body(doc,
    'The qualitative component employs an exploratory design, using semi-structured interviews and '
    'open-ended survey questions to gather detailed user feedback on their experience with the '
    'application. This includes perceptions of the reporting process, the clarity and usefulness of '
    'AI-generated risk assessments, the relevance of educational resources, and suggestions for '
    'improvement. The qualitative data will provide context and depth to the quantitative findings, '
    'enabling a more comprehensive understanding of the application\u2019s strengths and limitations.')

add_body(doc,
    'The convergent parallel mixed-methods design will be used, meaning that quantitative and '
    'qualitative data will be collected concurrently during the pilot testing phase and then '
    'integrated during the analysis and interpretation stages (Creswell & Plano Clark, 2017). '
    'This design is appropriate because the study seeks to validate quantitative usability scores '
    'with qualitative user narratives, providing a richer and more reliable evaluation of the '
    'application.')

# ============================================================
# 4.2 Population and Sample
# ============================================================
add_heading_text(doc, '4.2 Population and Sample')

add_heading_text(doc, '4.2.1 Target Population', level=2)
add_body(doc,
    'The target population for this study consists of adult mobile phone users (aged 18 years and '
    'older) residing in Lagos, Nigeria, and Douala, Cameroon, who have either been victims of '
    'cybercrime or are considered at elevated risk of cyber victimization. Lagos was selected as it '
    'is Nigeria\u2019s commercial capital, home to over 20 million residents and the highest '
    'concentration of internet users in the country, making it the epicentre of cybercrime '
    'perpetration and victimization in West Africa (NCC, 2023). Douala, Cameroon\u2019s economic hub '
    'with approximately 4 million residents, was selected for its rapidly expanding mobile money '
    'ecosystem, rising rates of digital fraud, and its value in representing the Francophone West '
    'African context (ANTIC, 2022).')

add_heading_text(doc, '4.2.2 Sample Size', level=2)
add_body(doc,
    'The pilot study will involve a sample of 50 to 100 participants, distributed across both '
    'cities. Approximately 30 to 50 participants will be recruited from Lagos, and 20 to 50 '
    'participants from Douala. This sample size is appropriate for a pilot usability study, as '
    'recommended by Nielsen (2000), who suggests that 5 to 15 users can identify the majority of '
    'usability issues, while larger samples of 30 or more are suitable for generating statistically '
    'meaningful SUS scores. The sample size of 50 to 100 also aligns with the project scope defined '
    'in the research proposal.')

add_heading_text(doc, '4.2.3 Sampling Technique', level=2)
add_body(doc,
    'The study will employ a combination of purposive sampling and snowball sampling. Purposive '
    'sampling will be used as the primary technique to select participants who meet specific criteria '
    'relevant to the study objectives. This non-probability sampling method is appropriate because '
    'the study targets a specific population\u2014mobile phone users who have experienced or are at risk '
    'of cybercrime\u2014rather than seeking to generalise findings to the entire population (Etikan et '
    'al., 2016). Snowball sampling will be used as a supplementary technique, particularly in '
    'Douala, where initial participants will be asked to refer other eligible individuals from their '
    'networks. This approach is effective for reaching populations that may be difficult to access '
    'through conventional recruitment methods, especially cybercrime victims who may be reluctant to '
    'self-identify (Goodman, 1961).')

add_heading_text(doc, '4.2.4 Inclusion and Exclusion Criteria', level=2)
add_body(doc,
    'Inclusion Criteria: Participants must be aged 18 years or older, must be active mobile phone '
    'users (smartphone with Android or iOS), must reside in Lagos, Nigeria, or Douala, Cameroon, '
    'and must have either experienced at least one cybercrime incident (such as phishing, online '
    'fraud, mobile money scam, or identity theft) or be identified as at elevated risk due to '
    'frequent online transactions, mobile money usage, or limited digital literacy. Participants '
    'must also be able to communicate in English, French, or Pidgin English.')

add_body(doc,
    'Exclusion Criteria: Individuals under the age of 18, individuals who do not use smartphones, '
    'individuals residing outside the two pilot cities, and individuals who are unable to provide '
    'informed consent will be excluded from the study.')

# ============================================================
# 4.3 Data Collection Methods
# ============================================================
add_heading_text(doc, '4.3 Data Collection Methods')
add_body(doc,
    'Data collection for this study will involve both primary and secondary methods, aligned with '
    'the mixed-methods research design.')

add_heading_text(doc, '4.3.1 Primary Data Collection', level=2)
add_body(doc, 'Primary data will be collected through four main methods during the pilot testing phase:')

add_body(doc,
    'Surveys and Questionnaires: A structured questionnaire incorporating the System Usability '
    'Scale (SUS) will be administered to all pilot participants after they have used the application '
    'for a minimum of two weeks. The SUS is a widely validated 10-item Likert scale instrument that '
    'produces a usability score ranging from 0 to 100 (Brooke, 1996). In addition to the SUS, the '
    'questionnaire will include demographic questions (age, gender, education level, digital literacy '
    'self-assessment), questions about prior cybercrime experience, and Likert-scale items measuring '
    'perceived usefulness of the risk assessment feature, satisfaction with educational resources, '
    'and willingness to use the app for future incident reporting.')

add_body(doc,
    'Semi-Structured Interviews: A subset of 10 to 15 participants (5 to 8 from Lagos, 5 to 7 from '
    'Douala) will be selected for semi-structured interviews to explore their experiences with the '
    'application in greater depth. Interview questions will focus on the ease of the reporting '
    'process, the clarity and helpfulness of AI-generated risk assessments, the relevance and '
    'accessibility of educational content, barriers encountered during use, and recommendations for '
    'improvement. Interviews will be conducted in the participant\u2019s preferred language (English, '
    'French, or Pidgin English) and will last approximately 20 to 30 minutes each. All interviews '
    'will be audio-recorded with participant consent and subsequently transcribed for analysis.')

add_body(doc,
    'App Usage Analytics: The application will be instrumented with Firebase Analytics to capture '
    'quantitative usage data during the pilot period. This includes metrics such as the number of '
    'scam reports submitted, average time to complete a report, frequency of access to educational '
    'resources, risk assessment response time (measured against the 30-second target), session '
    'duration, and user retention rates. These analytics will provide objective behavioural data to '
    'complement self-reported survey and interview findings.')

add_body(doc,
    'Machine Learning Model Testing: The scam detection model will be evaluated using a labelled '
    'test dataset comprising West African cybercrime records. The model\u2019s classification performance '
    'will be assessed across the five scam categories (advance-fee fraud, mobile money fraud, '
    'phishing, romance scams, and identity theft) using accuracy, precision, recall, and F1-score. '
    'Additionally, the model\u2019s real-world performance will be monitored during the pilot by '
    'comparing its automated classifications with manual expert reviews of a random sample of '
    'user-submitted reports.')

add_heading_text(doc, '4.3.2 Secondary Data Collection', level=2)
add_body(doc,
    'Secondary data will be gathered from the following sources to support model training, '
    'contextual analysis, and benchmarking:')

add_body(doc,
    'Literature Review: An extensive review of existing scholarly literature on cybercrime in West '
    'Africa, AI-based fraud detection, mobile application development for public safety, and '
    'ethical AI integration in developing contexts. This review, largely completed during the '
    'annotated bibliography phase (Unit 3), provides the theoretical and empirical foundation for '
    'the study.')

add_body(doc,
    'Cybercrime Datasets: Publicly available cybercrime datasets will be used for training and '
    'validating the machine learning model. These include datasets from the Anti-Phishing Working '
    'Group (APWG), the UCI Machine Learning Repository (email spam and phishing datasets), and the '
    'Kaggle fraud detection datasets. These will be supplemented with incident records sourced from '
    'EFCC annual reports, ANTIC cybercrime logs, and ngCERT advisories, which provide region-specific '
    'scam patterns essential for training the model to recognise West African cybercrime typologies.')

add_body(doc,
    'Institutional and Policy Documents: Reports from INTERPOL, the African Union, the World Bank, '
    'and national cybersecurity agencies (EFCC, ANTIC, ngCERT) will be reviewed to contextualise '
    'the findings within the broader cybercrime landscape in West Africa.')

# ============================================================
# 4.4 Data Collection Tools
# ============================================================
add_heading_text(doc, '4.4 Data Collection Tools')

add_heading_text(doc, '4.4.1 Identification of Data Collection Instruments', level=2)
add_body(doc, 'The following instruments will be used for data collection:')

add_body(doc,
    '1. System Usability Scale (SUS) Questionnaire: A standardised 10-item questionnaire scored on '
    'a 5-point Likert scale, widely used to evaluate the usability of software systems. The SUS has '
    'demonstrated strong reliability (Cronbach\u2019s alpha of 0.91) and is considered an industry '
    'standard for usability assessment (Bangor et al., 2008).')

add_body(doc,
    '2. Custom Demographic and Experience Questionnaire: A researcher-designed questionnaire '
    'capturing participant demographics, prior cybercrime experience, digital literacy levels, and '
    'perceptions of the application\u2019s specific features (risk assessment, educational resources, '
    'reporting interface).')

add_body(doc,
    '3. Semi-Structured Interview Guide: A set of open-ended questions organised around key themes: '
    'reporting experience, risk assessment feedback, educational content relevance, barriers to use, '
    'and suggestions for improvement. The guide will be flexible enough to allow follow-up probing '
    'based on participant responses.')

add_body(doc,
    '4. Firebase Analytics Dashboard: An automated data collection tool embedded within the '
    'application to track user behaviour metrics including report submissions, feature usage '
    'frequency, session duration, and risk assessment response times.')

add_body(doc,
    '5. Machine Learning Evaluation Framework: A Python-based testing framework using Scikit-learn '
    'and TensorFlow evaluation modules to compute classification metrics (accuracy, precision, '
    'recall, F1-score, confusion matrix) on the scam detection model.')

add_heading_text(doc, '4.4.2 Development and Validation of Instruments', level=2)
add_body(doc,
    'The SUS questionnaire requires no further validation, as it is an established and widely '
    'validated instrument in the field of human-computer interaction (Brooke, 1996). The custom '
    'demographic and experience questionnaire will be developed by the researcher and reviewed by '
    'the project supervisor for content validity. A small-scale pre-test will be conducted with 5 '
    'to 10 individuals (not included in the main pilot sample) to assess the clarity, relevance, '
    'and completeness of the questionnaire items. Feedback from the pre-test will be used to refine '
    'question wording and response options before the main pilot deployment.')

add_body(doc,
    'The semi-structured interview guide will be developed based on the research questions and '
    'reviewed by the supervisor. It will be piloted with 2 to 3 individuals to ensure that questions '
    'are clear, culturally appropriate, and elicit meaningful responses. The interview guide will be '
    'translated into French for use with Francophone participants in Douala, and the translation '
    'will be reviewed by a bilingual colleague to ensure accuracy.')

add_body(doc,
    'The machine learning evaluation framework will follow standard practices in the field, using '
    'stratified k-fold cross-validation (k=5) to ensure robust and generalisable performance '
    'estimates across all five scam categories (Kohavi, 1995).')

# ============================================================
# 4.5 Data Analysis Methods
# ============================================================
add_heading_text(doc, '4.5 Data Analysis Methods')
add_body(doc,
    'Data analysis will employ both quantitative and qualitative techniques, consistent with the '
    'mixed-methods research design.')

add_heading_text(doc, '4.5.1 Quantitative Analysis', level=2)
add_body(doc,
    'Machine Learning Model Performance: Classification metrics including accuracy, precision, '
    'recall, and F1-score will be computed for each of the five scam categories and for the model '
    'overall. A confusion matrix will be generated to identify patterns of misclassification. The '
    'target benchmark is an overall classification accuracy of at least 85%. Performance will be '
    'analysed using Python libraries, specifically Scikit-learn for metric computation and '
    'Matplotlib/Seaborn for visualisation.')

add_body(doc,
    'Usability Analysis: SUS scores will be calculated for each participant following the standard '
    'scoring methodology (Brooke, 1996). The mean SUS score and standard deviation will be computed '
    'for the overall sample as well as for subgroups (by city, age group, gender, and digital '
    'literacy level). A SUS score above 68 is considered above average usability (Sauro, 2011). '
    'Descriptive statistics (mean, median, standard deviation, frequency distributions) will be '
    'used to summarise demographic data, Likert-scale responses on perceived usefulness, and '
    'satisfaction ratings.')

add_body(doc,
    'App Usage Analytics: Firebase Analytics data will be analysed using descriptive statistics to '
    'summarise usage patterns, including the total number of reports submitted, average report '
    'completion time, educational resource access frequency, and risk assessment response times. '
    'The 30-second response time target will be evaluated by computing the percentage of risk '
    'assessments delivered within the threshold.')

add_heading_text(doc, '4.5.2 Qualitative Analysis', level=2)
add_body(doc,
    'Interview and open-ended survey responses will be analysed using thematic analysis, following '
    'the six-phase framework outlined by Braun and Clarke (2006): (1) familiarisation with the data '
    'through reading and re-reading transcripts, (2) generating initial codes, (3) searching for '
    'themes, (4) reviewing themes, (5) defining and naming themes, and (6) producing the final '
    'report. The analysis will identify recurring patterns in user experiences, perceptions, and '
    'recommendations. Coding will be conducted manually, with codes and themes reviewed by the '
    'project supervisor to enhance analytical rigour and reduce researcher bias.')

add_body(doc,
    'The qualitative findings will be integrated with the quantitative results during the '
    'interpretation phase using a side-by-side comparison approach, where qualitative themes are '
    'mapped against quantitative metrics to identify areas of convergence and divergence (Creswell '
    '& Plano Clark, 2017). For example, if a particular feature receives a low quantitative '
    'satisfaction score, the qualitative data will be examined to identify specific reasons and user '
    'suggestions for improvement.')

# ============================================================
# 4.6 Procedures (Ethical Considerations)
# ============================================================
add_heading_text(doc, '4.6 Procedures and Ethical Considerations')
add_body(doc,
    'This study will adhere to established ethical principles for research involving human '
    'participants, guided by the Belmont Report\u2019s core principles of respect for persons, '
    'beneficence, and justice (National Commission for the Protection of Human Subjects, 1979).')

add_body(doc,
    'Institutional Approval: Ethical clearance will be sought from the African Leadership '
    'University\u2019s Institutional Review Board (IRB) or equivalent ethics committee prior to the '
    'commencement of pilot testing. No data collection involving human participants will begin until '
    'formal ethical approval has been obtained.')

add_body(doc,
    'Informed Consent: All participants will be provided with a detailed informed consent form '
    'prior to their involvement in the study. The consent form will explain the purpose of the '
    'research, the nature of their participation, the types of data that will be collected, how '
    'their data will be used and stored, their right to withdraw at any time without penalty, and '
    'the contact details of the researcher and supervisor. For participants in Douala, the consent '
    'form will be available in both English and French. Consent will be obtained in writing before '
    'any data collection begins.')

add_body(doc,
    'Confidentiality and Data Privacy: All participant data will be treated as strictly '
    'confidential. Personal identifiers will be removed from datasets and replaced with unique '
    'participant codes during analysis and reporting. Survey responses, interview transcripts, and '
    'app usage data will be stored in encrypted digital files accessible only to the researcher and '
    'supervisor. Firebase data will be secured using Firebase Authentication and Firestore security '
    'rules to prevent unauthorised access. The application will comply with relevant data protection '
    'principles, and participants will be informed about what data the app collects and how it is '
    'processed.')

add_body(doc,
    'Sensitive Subject Matter: Given that the study involves cybercrime victims, the researcher '
    'recognises that participants may experience discomfort when recounting their experiences. '
    'Participation will be entirely voluntary, and participants will be reminded that they may skip '
    'any question or withdraw from the study at any time. Where appropriate, participants will be '
    'provided with information about available victim support services in their city.')

add_body(doc,
    'Data Retention and Disposal: Research data will be retained for a period of two years after '
    'the completion of the study for verification and potential follow-up analysis. After this '
    'period, all data will be permanently deleted from digital storage. Physical consent forms will '
    'be stored in a locked cabinet and shredded after the retention period.')

add_body(doc,
    'No Harm Principle: The study does not involve any experimental interventions that could cause '
    'physical or psychological harm. The application is designed as a support tool for cybercrime '
    'victims, and its use during the pilot testing phase poses no foreseeable risks to participants '
    'beyond those encountered in normal smartphone usage.')

add_body(doc,
    'Fair Participant Selection: Recruitment will be conducted through community technology hubs, '
    'universities (University of Lagos, Yaba College of Technology, University of Douala), '
    'cybercaf\u00e9 networks, and local NGOs. The researcher will ensure that recruitment practices do '
    'not exclude individuals on the basis of gender, ethnicity, socioeconomic status, or language. '
    'The multilingual design of the application (English, French, and Pidgin English) is intended '
    'to promote inclusive access.')

# ============================================================
# References
# ============================================================
doc.add_page_break()
add_heading_text(doc, 'References')

refs = [
    'ANTIC. (2022). Rapport annuel sur la cybercriminalit\u00e9 au Cameroun 2021\u20132022. Agence Nationale des Technologies de l\u2019Information et de la Communication.',
    'Bangor, A., Kortum, P. T., & Miller, J. T. (2008). An empirical evaluation of the System Usability Scale. International Journal of Human-Computer Interaction, 24(6), 574\u2013594. https://doi.org/10.1080/10447310802205776',
    'Braun, V., & Clarke, V. (2006). Using thematic analysis in psychology. Qualitative Research in Psychology, 3(2), 77\u2013101. https://doi.org/10.1191/1478088706qp063oa',
    'Brooke, J. (1996). SUS: A \u201cquick and dirty\u201d usability scale. In P. W. Jordan, B. Thomas, B. A. Weerdmeester, & I. L. McClelland (Eds.), Usability evaluation in industry (pp. 189\u2013194). Taylor & Francis.',
    'Creswell, J. W., & Creswell, J. D. (2018). Research design: Qualitative, quantitative, and mixed methods approaches (5th ed.). SAGE Publications.',
    'Creswell, J. W., & Plano Clark, V. L. (2017). Designing and conducting mixed methods research (3rd ed.). SAGE Publications.',
    'Etikan, I., Musa, S. A., & Alkassim, R. S. (2016). Comparison of convenience sampling and purposive sampling. American Journal of Theoretical and Applied Statistics, 5(1), 1\u20134. https://doi.org/10.11648/j.ajtas.20160501.11',
    'Goodman, L. A. (1961). Snowball sampling. The Annals of Mathematical Statistics, 32(1), 148\u2013170.',
    'Kohavi, R. (1995). A study of cross-validation and bootstrap for accuracy estimation and model selection. Proceedings of the 14th International Joint Conference on Artificial Intelligence, 2, 1137\u20131143.',
    'National Commission for the Protection of Human Subjects. (1979). The Belmont Report: Ethical principles and guidelines for the protection of human subjects of research. U.S. Department of Health and Human Services.',
    'NCC. (2023). Subscriber statistics: Annual report 2022. Nigerian Communications Commission.',
    'Nielsen, J. (2000). Why you only need to test with 5 users. Nielsen Norman Group. https://www.nngroup.com/articles/why-you-only-need-to-test-with-5-users/',
    'Sauro, J. (2011). A practical guide to the System Usability Scale: Background, benchmarks, and best practices. Measuring Usability LLC.',
]

for ref in refs:
    p = doc.add_paragraph()
    run = p.add_run(ref)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    p.paragraph_format.line_spacing = 2.0

doc.save('WilsonsNavidWadoTiwa_Unit Four Assignment.docx')
print('Document saved successfully!')
