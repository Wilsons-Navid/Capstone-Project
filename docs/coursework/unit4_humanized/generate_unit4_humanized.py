import sys, io, re
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = docx.Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 2.0

# Set margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('AI-Powered Mobile Application for Cybercrime Reporting and Scam Detection in West Africa')
run.bold = True
run.font.size = Pt(14)
run.font.name = 'Arial'

doc.add_paragraph()

# Author info
for line in ['Wilsons Navid Wado Tiwa', 'African Leadership University', 'Pre-Capstone Project', 'Supervisor: Unknown']:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

doc.add_paragraph()

# Chapter heading
h = doc.add_paragraph()
h.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = h.add_run('CHAPTER FOUR: RESEARCH METHODOLOGY')
run.bold = True
run.font.size = Pt(14)
run.font.name = 'Arial'

doc.add_paragraph()

def add_heading_text(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(14)
        run.font.name = 'Arial'
    else:
        run.font.size = Pt(12)
        run.font.name = 'Arial'
    return p

def add_body(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.paragraph_format.line_spacing = 2.0
    return p

# ============================================================
# 4.0 Introduction
# ============================================================
add_heading_text(doc, '4.0 Introduction')
add_body(doc,
    'This chapter lays out the research methodology for the AI-powered mobile application '
    'for cybercrime reporting and scam detection in West Africa. It covers the research design, '
    'the target population and how participants will be sampled, what data will be collected and '
    'with which instruments, how that data will be analysed, and the ethical safeguards in place. '
    'The study uses a mixed-methods design. On the quantitative side, the machine learning scam '
    'detection model will be tested against standard classification benchmarks, and user experience '
    'will be scored through the System Usability Scale. On the qualitative side, semi-structured '
    'interviews and open-ended survey questions will capture how users in Nigeria and Cameroon '
    'actually experience the app, what they find useful, and where they see room for improvement.')

# ============================================================
# 4.1 Research Design
# ============================================================
add_heading_text(doc, '4.1 Research design')
add_body(doc,
    'This study uses a mixed-methods research design that brings together quantitative and '
    'qualitative approaches. A mixed-methods approach makes sense here because the project has '
    'two distinct but connected sides: first, building and evaluating an AI-powered mobile '
    'application, which requires quantitative measurement; and second, understanding how users '
    'perceive and interact with the app, which requires qualitative exploration (Creswell & Creswell, 2018).')

add_body(doc,
    'The quantitative side combines an experimental component with a descriptive survey. The '
    'experimental part involves training and testing machine learning models on West African '
    'cybercrime data to classify scam reports into five categories: advance-fee fraud, mobile money '
    'fraud, phishing, romance scams, and identity theft. Model performance will be measured by '
    'accuracy, precision, recall, and F1-score, with a target accuracy of at least 85%. The '
    'survey part involves giving pilot users the System Usability Scale (SUS), a 10-item '
    'questionnaire that produces a numerical usability score (Brooke, 1996).')

add_body(doc,
    'The qualitative side uses semi-structured interviews and open-ended survey questions to '
    'collect detailed feedback from users. Topics include how intuitive the reporting process '
    'feels, whether the AI-generated risk assessments are clear and useful, how relevant the '
    'educational resources are, and what users would change. This qualitative data adds context '
    'to the numbers and helps explain why users rate certain features the way they do.')

add_body(doc,
    'The study follows a convergent parallel design, meaning quantitative and qualitative data '
    'will be collected at the same time during pilot testing and then brought together during '
    'analysis (Creswell & Plano Clark, 2017). This is a good fit because I want to check '
    'whether what users say in interviews lines up with how they score the app on the SUS, and '
    'where those two data sources diverge, that is usually where the most useful insights are.')

# ============================================================
# 4.2 Population and Sample
# ============================================================
add_heading_text(doc, '4.2 Population and sample')

add_heading_text(doc, '4.2.1 Target population', level=2)
add_body(doc,
    'The study targets adult mobile phone users (18 and older) living in Lagos, Nigeria, and '
    'Douala, Cameroon, who have either been victims of cybercrime or face a high risk of it. '
    'Lagos is Nigeria\'s commercial capital, with over 20 million residents and the country\'s '
    'highest concentration of internet users. It also has the highest rates of cybercrime '
    'victimization in the region (NCC, 2023). Douala is Cameroon\'s economic centre, home to '
    'roughly 4 million people, and was chosen because its mobile money market is expanding fast, '
    'digital fraud is on the rise, and it gives the study a Francophone perspective that Lagos '
    'alone cannot provide (ANTIC, 2022).')

add_heading_text(doc, '4.2.2 Sample size', level=2)
add_body(doc,
    'The pilot will recruit 50 to 100 participants across both cities: roughly 30 to 50 from '
    'Lagos and 20 to 50 from Douala. Nielsen (2000) found that as few as 5 to 15 users can '
    'uncover most usability issues, but samples of 30 or more are needed for statistically '
    'meaningful SUS scores. A target of 50 to 100 balances these considerations and matches the '
    'scope laid out in the research proposal.')

add_heading_text(doc, '4.2.3 Sampling technique', level=2)
add_body(doc,
    'Participants will be recruited through a combination of purposive and snowball sampling. '
    'Purposive sampling is the primary method: the study needs people who use mobile phones and '
    'have some exposure to cybercrime, so random sampling would be inefficient (Etikan et al., '
    '2016). Snowball sampling will supplement this, especially in Douala, where early participants '
    'will refer others from their networks. Snowball sampling is particularly useful for reaching '
    'cybercrime victims, who are often reluctant to come forward on their own (Goodman, 1961).')

add_heading_text(doc, '4.2.4 Inclusion and exclusion criteria', level=2)
add_body(doc,
    'To participate, individuals must be at least 18 years old, own and actively use a smartphone '
    '(Android or iOS), live in Lagos or Douala, and have either experienced a cybercrime incident '
    '(phishing, fraud, mobile money scam, identity theft, etc.) or face elevated risk because of '
    'frequent online transactions, heavy mobile money use, or limited digital literacy. They must '
    'also speak English, French, or Pidgin English.')

add_body(doc,
    'Anyone under 18, anyone without a smartphone, anyone living outside the two pilot cities, '
    'and anyone unable to give informed consent will be excluded.')

# ============================================================
# 4.3 Data Collection Methods
# ============================================================
add_heading_text(doc, '4.3 Data collection methods')
add_body(doc,
    'Data collection draws on both primary and secondary sources, in line with the mixed-methods '
    'design.')

add_heading_text(doc, '4.3.1 Primary data collection', level=2)
add_body(doc, 'Four primary data collection methods will be used during the pilot:')

add_body(doc,
    'Surveys and questionnaires. After using the app for a minimum of two weeks, every pilot '
    'participant will complete a questionnaire built around the System Usability Scale (SUS). '
    'The SUS is a 10-item Likert scale that yields a usability score from 0 to 100 (Brooke, '
    '1996). Beyond the SUS items, the questionnaire will ask about demographics (age, gender, '
    'education, self-rated digital literacy), prior cybercrime experience, how useful participants '
    'found the risk assessment feature, how satisfied they were with the educational content, and '
    'whether they would use the app to report future incidents.')

add_body(doc,
    'Semi-structured interviews. Between 10 and 15 participants (5 to 8 from Lagos, 5 to 7 from '
    'Douala) will sit for one-on-one interviews lasting 20 to 30 minutes. The interviews will explore '
    'how easy or difficult the reporting process was, whether the AI risk assessments made '
    'sense, how relevant the educational materials felt, what barriers came up during use, and '
    'what participants would improve. Each interview will be conducted in the participant\'s '
    'preferred language and audio-recorded with their consent, then transcribed.')

add_body(doc,
    'App usage analytics. Firebase Analytics will be embedded in the app to track behavioural '
    'data during the pilot. This includes the number of scam reports submitted, how long it takes '
    'to complete a report, how often users access educational resources, whether risk assessments '
    'are returned within the 30-second target, session length, and retention rates. These numbers '
    'give an objective counterpart to what users self-report in surveys and interviews.')

add_body(doc,
    'Machine learning model testing. The scam detection model will be evaluated on a labelled '
    'test set of West African cybercrime records. Classification performance will be measured for '
    'each of the five scam categories using accuracy, precision, recall, and F1-score. During the '
    'pilot, real-world accuracy will also be checked by having an expert manually review a random '
    'sample of user-submitted reports and comparing those labels to the model\'s predictions.')

add_heading_text(doc, '4.3.2 Secondary data collection', level=2)
add_body(doc,
    'Secondary data will come from three sources:')

add_body(doc,
    'Literature review. The annotated bibliography completed in Unit 3 surveyed existing research '
    'on cybercrime in West Africa, AI-based fraud detection, mobile apps for public safety, and '
    'ethical AI deployment in developing countries. That review supplies the theoretical and '
    'empirical grounding for this study.')

add_body(doc,
    'Cybercrime datasets. Publicly available datasets from the Anti-Phishing Working Group (APWG), '
    'the UCI Machine Learning Repository (email spam and phishing sets), and Kaggle fraud detection '
    'collections will be used to train and validate the model. These will be supplemented with '
    'region-specific records from EFCC annual reports, ANTIC cybercrime logs, and ngCERT advisories, '
    'which capture scam patterns particular to West Africa.')

add_body(doc,
    'Institutional and policy documents. Reports from INTERPOL, the African Union, the World Bank, '
    'and national cybersecurity bodies (EFCC, ANTIC, ngCERT) will provide broader context for '
    'interpreting the study\'s findings.')

# ============================================================
# 4.4 Data Collection Tools
# ============================================================
add_heading_text(doc, '4.4 Data collection tools')

add_heading_text(doc, '4.4.1 Instruments', level=2)
add_body(doc, 'Five instruments will be used:')

add_body(doc,
    '1. System Usability Scale (SUS) questionnaire. A 10-item questionnaire scored on a 5-point '
    'Likert scale. The SUS has strong reliability (Cronbach\'s alpha of 0.91) and is widely '
    'accepted as a standard usability measure (Bangor et al., 2008).')

add_body(doc,
    '2. Custom demographic and experience questionnaire. A researcher-designed form that captures '
    'participant demographics, prior cybercrime experience, digital literacy, and opinions on '
    'specific app features like the risk assessment tool, educational resources, and reporting '
    'interface.')

add_body(doc,
    '3. Semi-structured interview guide. A set of open-ended questions grouped by theme: the '
    'reporting experience, risk assessment feedback, educational content, barriers to use, and '
    'ideas for improvement. The guide leaves room for follow-up questions based on what '
    'participants say.')

add_body(doc,
    '4. Firebase Analytics dashboard. An automated tracking layer inside the app that logs '
    'report submissions, feature usage, session length, and risk assessment response times.')

add_body(doc,
    '5. Machine learning evaluation framework. A Python-based pipeline using Scikit-learn and '
    'TensorFlow to compute accuracy, precision, recall, F1-score, and confusion matrices for '
    'the scam detection model.')

add_heading_text(doc, '4.4.2 Instrument validation', level=2)
add_body(doc,
    'The SUS needs no additional validation; it has been used and tested extensively in '
    'human-computer interaction research (Brooke, 1996). The custom questionnaire will be '
    'drafted by the researcher and reviewed by the project supervisor for content validity, '
    'then pre-tested with 5 to 10 people (who will not be part of the main pilot) to check '
    'whether the questions are clear and the response options make sense. Any confusing items '
    'will be revised before deployment.')

add_body(doc,
    'The interview guide will also be reviewed by the supervisor and piloted with 2 to 3 '
    'individuals to make sure the questions are clear, culturally appropriate, and generate '
    'useful responses. A French translation will be prepared for Douala participants and checked '
    'by a bilingual colleague.')

add_body(doc,
    'For the machine learning evaluation, standard practice will be followed: stratified 5-fold '
    'cross-validation to produce robust performance estimates across all five scam categories '
    '(Kohavi, 1995).')

# ============================================================
# 4.5 Data Analysis Methods
# ============================================================
add_heading_text(doc, '4.5 Data analysis methods')
add_body(doc,
    'Analysis will mirror the mixed-methods design, with separate quantitative and qualitative '
    'tracks that converge during interpretation.')

add_heading_text(doc, '4.5.1 Quantitative analysis', level=2)
add_body(doc,
    'Machine learning performance. Accuracy, precision, recall, and F1-score will be calculated '
    'for each scam category and for the model as a whole. A confusion matrix will show where '
    'misclassifications happen. The benchmark is 85% overall accuracy. All metrics will be '
    'computed in Python using Scikit-learn, with Matplotlib and Seaborn for visualisation.')

add_body(doc,
    'Usability scores. Each participant\'s SUS score will be calculated using Brooke\'s (1996) '
    'standard method. I will compute the mean and standard deviation for the full sample and '
    'for subgroups split by city, age, gender, and digital literacy. A mean SUS score above 68 '
    'is generally considered above-average usability (Sauro, 2011). Descriptive statistics '
    '(mean, median, standard deviation, frequency counts) will also summarise the Likert-scale '
    'items on perceived usefulness and satisfaction.')

add_body(doc,
    'App usage data. Firebase Analytics data will be summarised with descriptive statistics: '
    'total reports submitted, average time to finish a report, how often educational resources '
    'were accessed, and what percentage of risk assessments were delivered within the 30-second '
    'window.')

add_heading_text(doc, '4.5.2 Qualitative analysis', level=2)
add_body(doc,
    'Interview transcripts and open-ended survey responses will be analysed through thematic '
    'analysis, following Braun and Clarke\'s (2006) six phases: getting familiar with the data '
    'by reading transcripts multiple times, generating initial codes, looking for themes across '
    'codes, reviewing those themes, defining and naming them, and writing up the findings. '
    'Coding will be done manually, and the supervisor will review the codes and themes to '
    'reduce bias.')

add_body(doc,
    'During interpretation, qualitative themes will be placed alongside quantitative results in '
    'a side-by-side comparison (Creswell & Plano Clark, 2017). If a feature gets a low score '
    'on the survey, for instance, the interview data can help explain why and point toward '
    'specific fixes.')

# ============================================================
# 4.6 Procedures (Ethical Considerations)
# ============================================================
add_heading_text(doc, '4.6 Procedures and ethical considerations')
add_body(doc,
    'The study follows the Belmont Report\'s principles of respect for persons, beneficence, '
    'and justice (National Commission for the Protection of Human Subjects, 1979).')

add_body(doc,
    'Institutional approval. Ethical clearance will be obtained from the African Leadership '
    'University\'s Institutional Review Board (or its equivalent) before any pilot testing '
    'begins. No data involving human participants will be collected until that approval has '
    'been granted.')

add_body(doc,
    'Informed consent. Every participant will receive a consent form that explains the purpose '
    'of the study, what participation involves, what data will be collected, how it will be '
    'stored and used, and their right to withdraw at any point without consequences. In Douala, '
    'the form will be available in English and French. Written consent is required before any '
    'data collection starts.')

add_body(doc,
    'Confidentiality and data privacy. All participant data will be kept confidential. Names '
    'and other personal identifiers will be stripped from datasets and replaced with codes '
    'during analysis and reporting. Surveys, interview transcripts, and app usage logs will be '
    'stored in encrypted files that only the researcher and supervisor can access. On the app '
    'side, Firebase Authentication and Firestore security rules will prevent unauthorised access. '
    'Participants will be told exactly what data the app collects and how it is processed.')

add_body(doc,
    'Sensitive subject matter. Some participants may find it uncomfortable to talk about their '
    'cybercrime experiences. Participation is voluntary, and participants can skip any question '
    'or leave the study at any time. Where appropriate, information about local victim support '
    'services will be provided.')

add_body(doc,
    'Data retention and disposal. Research data will be kept for two years after the study ends, '
    'in case verification or follow-up analysis is needed. After that, digital files will be '
    'permanently deleted and physical consent forms will be shredded.')

add_body(doc,
    'No harm. The study does not involve any intervention that could cause physical or '
    'psychological harm. The app is a support tool, and using it during the pilot carries no '
    'risks beyond those of normal smartphone use.')

add_body(doc,
    'Fair recruitment. Participants will be recruited through community technology hubs, '
    'universities (University of Lagos, Yaba College of Technology, University of Douala), '
    'cybercafe networks, and local NGOs. Recruitment will not exclude anyone on the basis of '
    'gender, ethnicity, socioeconomic status, or language. The app itself supports English, '
    'French, and Pidgin English to keep access as broad as possible.')

# ============================================================
# References
# ============================================================
doc.add_page_break()
add_heading_text(doc, 'References')

refs = [
    'ANTIC. (2022). Rapport annuel sur la cybercriminalit\u00e9 au Cameroun 2021\u20132022. Agence Nationale des Technologies de l\u2019Information et de la Communication.',
    'Bangor, A., Kortum, P. T., & Miller, J. T. (2008). An empirical evaluation of the System Usability Scale. International Journal of Human-Computer Interaction, 24(6), 574\u2013594. https://doi.org/10.1080/10447310802205776',
    'Braun, V., & Clarke, V. (2006). Using thematic analysis in psychology. Qualitative Research in Psychology, 3(2), 77\u2013101. https://doi.org/10.1191/1478088706qp063oa',
    'Brooke, J. (1996). SUS: A "quick and dirty" usability scale. In P. W. Jordan, B. Thomas, B. A. Weerdmeester, & I. L. McClelland (Eds.), Usability evaluation in industry (pp. 189\u2013194). Taylor & Francis.',
    'Creswell, J. W., & Creswell, J. D. (2018). Research design: Qualitative, quantitative, and mixed methods approaches (5th ed.). SAGE Publications.',
    'Creswell, J. W., & Plano Clark, V. L. (2017). Designing and conducting mixed methods research (3rd ed.). SAGE Publications.',
    'Etikan, I., Musa, S. A., & Alkassim, R. S. (2016). Comparison of convenience sampling and purposive sampling. American Journal of Theoretical and Applied Statistics, 5(1), 1\u20134. https://doi.org/10.11648/j.ajtas.20160501.11',
    'Goodman, L. A. (1961). Snowball sampling. The Annals of Mathematical Statistics, 32(1), 148\u2013170.',
    'Kohavi, R. (1995). A study of cross-validation and bootstrap for accuracy estimation and model selection. Proceedings of the 14th International Joint Conference on Artificial Intelligence, 2, 1137\u20131143.',
    'National Commission for the Protection of Human Subjects. (1979). The Belmont Report: Ethical principles and guidelines for the protection of human subjects of research. U.S. Department of Health and Human Services.',
    'NCC. (2023). Subscriber statistics: Annual report 2022. Nigerian Communications Commission.',
    'Nielsen, J. (2000). Why you only need to test with 5 users. Nielsen Norman Group. https://www.nngroup.com/articles/why-you-only-need-to-test-with-5-users/',
    'Sauro, J. (2011). A practical guide to the System Usability Scale: Background, benchmarks, and best practices. Measuring Usability LLC.',
]

def add_hyperlink(paragraph, url, text):
    """Add a clickable hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0563C1')
    rPr.append(c)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '24')  # 12pt = 24 half-points
    rPr.append(sz)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

url_pattern = re.compile(r'(https?://\S+)')

for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    p.paragraph_format.line_spacing = 2.0
    # Split reference text around URLs
    parts = url_pattern.split(ref)
    for part in parts:
        if url_pattern.match(part):
            add_hyperlink(p, part, part)
        else:
            run = p.add_run(part)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

doc.save('WilsonsNavidWadoTiwa_Unit Four Assignment.docx')
print('Document saved successfully!')
