"""Generates the README submission checklist as a Microsoft Word .docx."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

BROWN = RGBColor(0x2D, 0x1B, 0x14)
AMBER = RGBColor(0x7A, 0x5C, 0x00)
GREEN = RGBColor(0x2E, 0x7D, 0x32)
GREY = RGBColor(0x49, 0x45, 0x4F)
HDR = "EFE7DE"

doc = Document()
nn = doc.styles["Normal"]; nn.font.name = "Calibri"; nn.font.size = Pt(11)


def shade(c, h):
    p = c._tc.get_or_add_tcPr(); s = OxmlElement("w:shd")
    s.set(qn("w:val"), "clear"); s.set(qn("w:fill"), h); p.append(s)


def head(t, lvl=1, color=BROWN):
    p = doc.add_heading(level=lvl); r = p.add_run(t); r.font.color.rgb = color; r.font.name = "Calibri"


def para(t, bold=False, italic=False, color=None, size=11):
    p = doc.add_paragraph(); r = p.add_run(t)
    r.bold = bold; r.italic = italic; r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    return p


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, ht in enumerate(headers):
        shade(t.rows[0].cells[i], HDR)
        r = t.rows[0].cells[i].paragraphs[0].add_run(ht); r.bold = True; r.font.color.rgb = BROWN; r.font.size = Pt(10)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            r = cells[i].paragraphs[0].add_run(str(v)); r.font.size = Pt(10)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()


tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run("RethicsAI — README Submission Checklist"); r.bold = True; r.font.size = Pt(20); r.font.color.rgb = BROWN
sp = doc.add_paragraph(); sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sp.add_run("What's left to fill in the README before submitting on Canvas"); sr.italic = True; sr.font.color.rgb = AMBER
doc.add_paragraph()
para("Everything else in the README is done (deployed app, build steps, testing strategies, proposal "
     "comparison, ML figures, architecture, model API). Only the items below need you: 2 text edits and "
     "7 screenshots. The README already references the screenshot filenames, so images appear "
     "automatically once you drop them into docs/assets/ — no editing needed for those.", color=GREY)

head("A. Two text edits (type these into README.md)", 1)
table(
    ["#", "Where", "What to do"],
    [
        ["1", "Section 10 — Demo video", "Replace “_add link here_” with your real 5-minute video URL (YouTube unlisted / Drive)."],
        ["2", "Section 4.3 — Performance table", "Replace the “e.g. …” placeholders with the real device, Android version, RAM and result for the 2 configs you test on."],
    ],
    widths=[0.4, 2.2, 4.4],
)

head("B. Seven screenshots to drop into docs/assets/", 1)
table(
    ["File name (exact)", "What to capture"],
    [
        ["test_run.png", "The `flutter test` terminal output showing “All 36 passed!”"],
        ["scan_advance_fee.png", "Scanner verdict for the advance-fee message"],
        ["scan_momo.png", "Scanner verdict for the mobile-money message"],
        ["scan_phishing.png", "Scanner verdict for the phishing message"],
        ["scan_safe.png", "Scanner verdict for the safe message"],
        ["perf_phone.png", "The app running on device / config #1"],
        ["perf_emulator.png", "The app running on device / config #2"],
    ],
    widths=[2.3, 4.7],
)

head("The 4 scanner inputs (paste these into the app)", 1)
table(
    ["Type", "Message"],
    [
        ["Advance-fee", "Congratulations! You have won 2,000,000 in the MTN promo. Send your BVN and a 5,000 activation fee to claim your prize now."],
        ["Mobile-money", "Your MoMo account will be blocked today. Dial *123*PIN# now to verify your wallet and avoid suspension."],
        ["Phishing", "Dear customer, your bank account has been suspended. Click http://bit.ly/secure-restore to reactivate immediately."],
        ["Safe", "Hi, are we still meeting at 3pm tomorrow at the office?"],
    ],
    widths=[1.3, 5.7],
)

head("When you're done", 1)
para("Tip: once the video is recorded and the 7 images are in docs/assets/, Claude can make the 2 text "
     "edits (video link + performance table), confirm every image renders, and commit/push it for you.")
para("Reminder before recording: warm up the model first (run one scan, wait ~30–60s) so the scanner "
     "shows the AI model verdict, not the keyword fallback. See DEMO_VIDEO_SCRIPT.docx for the full shot list.",
     italic=True, color=GREY)

out = os.path.join(os.path.dirname(__file__), "README_Submission_Checklist.docx")
doc.save(out)
print("Saved:", out)
