"""Generate the supervisor-facing Reconciliation Memo as a DOCX.

Mirrors the markdown source at docs/RECONCILIATION_MEMO.md but laid out as a
formal memorandum with TO/FROM/DATE/RE header block, structured tables for the
gap analysis and decisions, and a references section.

Output: docs/WilsonsNavidWadoTiwa_Reconciliation_Memo.docx
"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

if sys.stdout.encoding and sys.stdout.encoding.lower() != "utf-8":
    try:
        sys.stdout.reconfigure(encoding="utf-8")
    except Exception:
        pass

OUTPUT = Path(__file__).resolve().parent / "WilsonsNavidWadoTiwa_Reconciliation_Memo.docx"

doc = Document()

# ============================================================
# Page setup + default style
# ============================================================
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(12)
normal.paragraph_format.line_spacing = 1.15
normal.paragraph_format.space_after = Pt(6)


# ============================================================
# Helpers
# ============================================================
def add_page_number(section):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._element.append(fld_begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run._element.append(instr)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._element.append(fld_end)
    for r in p.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(10)


def add_header(section, text):
    h = section.header
    p = h.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    run.italic = True


def heading(text, level=1):
    sizes = {1: 14, 2: 12, 3: 11}
    spaces_before = {1: 16, 2: 10, 3: 8}
    size = sizes.get(level, 12)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(spaces_before.get(level, 8))
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = True
    return p


def para(text, bold=False, italic=False, size=12, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def bullet(text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    p.paragraph_format.space_after = Pt(3)
    for r in p.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)
    return p


def numbered(text):
    p = doc.add_paragraph(text, style="List Number")
    p.paragraph_format.space_after = Pt(3)
    for r in p.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)
    return p


def callout(label, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.2)
    r1 = p.add_run(label + " ")
    r1.font.name = "Arial"
    r1.font.size = Pt(11)
    r1.bold = True
    r1.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)
    r2 = p.add_run(text)
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(12)
    return p


def table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    try:
        t.style = "Light Grid Accent 1"
    except KeyError:
        t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.font.name = "Arial"
        run.font.size = Pt(11)
        run.bold = True
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            cell = t.rows[ri].cells[ci]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
    doc.add_paragraph()
    return t


# ============================================================
# Footer + header on all sections
# ============================================================
for sec in doc.sections:
    add_page_number(sec)
    add_header(sec, "Reconciliation Memo - Capstone")


# ============================================================
# Memo header block
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("MEMORANDUM")
r.font.name = "Arial"
r.font.size = Pt(20)
r.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Phase 0 Reconciliation - Capstone Scope vs. Built Reality")
r.font.name = "Arial"
r.font.size = Pt(13)
r.italic = True
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()

# TO / FROM / DATE / RE block
header_rows = [
    ("To:", "Capstone Supervisor"),
    ("From:", "Wilsons Navid Wado Tiwa"),
    ("Date:", "5 May 2026"),
    ("Re:", "Reconciliation of Pre-Capstone proposal with built application; five decisions required"),
    ("Window:", "2 months (target completion ~5 July 2026, 56 days)"),
    ("Decision required by:", "End of Sprint 1 (Day 7)"),
    ("Read time:", "~10 minutes"),
]

t = doc.add_table(rows=len(header_rows), cols=2)
t.autofit = False
t.columns[0].width = Inches(1.6)
t.columns[1].width = Inches(5.4)
for i, (label, value) in enumerate(header_rows):
    cell_l = t.rows[i].cells[0]
    cell_l.text = ""
    r = cell_l.paragraphs[0].add_run(label)
    r.font.name = "Arial"
    r.font.size = Pt(11)
    r.bold = True
    cell_r = t.rows[i].cells[1]
    cell_r.text = ""
    r = cell_r.paragraphs[0].add_run(value)
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)

doc.add_paragraph()

# Horizontal rule (visual separator)
hr = doc.add_paragraph()
hr_run = hr.add_run("___________________________________________________________________________")
hr_run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
hr_run.font.size = Pt(10)


# ============================================================
# 1. Purpose
# ============================================================
heading("1. Purpose", level=1)

para(
    "The Pre-Capstone proposal (docs/coursework/Major-assesment/) and the working application "
    "diverge in three load-bearing ways. Two further threads you raised - "
    "deepfakes and the CMU Africa Upanzi Network research on mobile money "
    "scams - are not yet incorporated. This memo:"
)
numbered("States the gap honestly.")
numbered("Proposes a single reframing that resolves all three divergences and absorbs both new threads.")
numbered("Asks for sign-off on five specific decisions before Sprint 2 begins.")

para(
    "I am asking for sign-off so that I can stop revisiting the same scoping "
    "questions every week and spend the remainder of the two-month window on "
    "research, build, and pilot work."
)


# ============================================================
# 2. What the proposal locked in
# ============================================================
heading("2. What the Pre-Capstone Proposal Locked In", level=1)

table(
    ["Element", "Proposal commitment"],
    [
        ["Pilot scope", "Lagos, Nigeria + Douala, Cameroon (50 to 100 users)"],
        ["Languages", "English, French, Pidgin English (3)"],
        ["Scam categories", "5 - advance-fee fraud, mobile money fraud, phishing, romance scams, identity theft"],
        ["ML approach", "Custom Scikit-learn / TensorFlow classifier trained on West African data"],
        ["ML target", "At least 85% classification accuracy"],
        ["Risk-assessment SLA", "30 seconds via Firebase Cloud Function"],
        ["Educational content", "20 or more articles in EN / FR / Pidgin English"],
        ["Sample size", "50 to 100 users; 10 to 15 semi-structured interviews; SUS questionnaire"],
    ],
)


# ============================================================
# 3. What the app actually is
# ============================================================
heading("3. What the Application Actually Is, Today", level=1)

table(
    ["Element", "Built reality"],
    [
        ["Name", "Rethicssec (folder still says rethicsai), v1.0.2+3"],
        ["Positioning", "Cybersecurity Platform for Africa (pan-African)"],
        ["Languages", "11 - en, sw, fr, ar, ha, yo, ig, zu, xh, af, sawa (Duala)"],
        ["Features built", "auth, dashboard, ai_assistant, incidents, cases, education, scanner, emergency, notifications, settings, admin"],
        ["AI/ML layer", "Vertex AI Gemini 1.5 Flash + OpenAI GPT-4o-mini via Firebase Cloud Functions"],
        ["Custom-trained model", "None"],
        ["Measured accuracy", "None"],
        ["Fixed scam taxonomy in code", "None - open-ended LLM threat analysis"],
        ["Production state", "Privacy Policy finalised; Play Store assets ready"],
    ],
)


# ============================================================
# 4. The gaps
# ============================================================
heading("4. The Gaps, Scored by Risk", level=1)

table(
    ["#", "Gap", "Severity", "Why it matters for the defense"],
    [
        ["G1", "Custom ML model promised but absent", "Critical",
         "Without it, \"AI-powered\" in the title is delivered entirely by third-party APIs. No evaluation, no accuracy claim, no contribution to the West African ML literature."],
        ["G2", "Pilot scope (Lagos+Douala) vs. app positioning (pan-African)", "High",
         "Pilot data will only validate two cities; broader claims will be unsupported by evidence."],
        ["G3", "3 languages promised, 11 in app", "Medium",
         "More features than I can rigorously evaluate; risk of shallow multilingual claims."],
    ],
)


# ============================================================
# 5. New threads from supervisor
# ============================================================
heading("5. Two New Threads You Raised", level=1)

heading("5.1 Deepfakes", level=2)

para(
    "Operationally, AI-generated voice (vishing 2.0), video (romance and "
    "sextortion), and image (fake KYC documents) attacks are now in the wild "
    "against African mobile money users. Voice cloning targeting M-Pesa and "
    "MTN MoMo confirmations has been reported across East Africa."
)

para(
    "For this capstone, however, detecting deepfakes is a different ML problem "
    "from classifying scam text. It would require audio or vision models, which "
    "is a separate research thread, not a tweak. The recommended posture is to "
    "treat deepfakes as a modular extension, not a pivot:"
)
bullet(
    "Add a sixth category, synthetic_media_fraud, to the text-only classification "
    "taxonomy. Even text-only classifiers can pick up the giveaway language "
    "(\"listen to this voice note from your son\")."
)
bullet(
    "Add one focused chapter section on the deepfake threat landscape in Africa, "
    "with audio/visual deepfake detection framed as Future Work."
)
bullet(
    "Optionally, prototype one deepfake-aware UI feature (a \"this audio looks "
    "AI-generated\" warning) using an off-the-shelf detector."
)

heading("5.2 CMU Africa Upanzi Network - mobile money scam research", level=2)

para(
    "This is the highest-leverage external thread in the entire memo. The "
    "Upanzi Network at CMU-Africa (Kigali) runs research that maps directly "
    "onto this project's needs:"
)
bullet(
    "Lamptey, Gueye, Luhanga, Seidu, Sowon (2024) - \"A honeynet infrastructure to "
    "battle SMS scammers.\" Inexpensive honeynet collecting smishing messages "
    "from Rwanda, Botswana, Ghana, Kenya, Uganda."
)
bullet(
    "Sowon et al. (2024) - \"The Role of User-Agent Interactions on Mobile Money "
    "Practices in Kenya and Tanzania.\" IEEE Symposium on Security & Privacy 2024. "
    "72 MoMo-user interviews."
)
bullet(
    "\"Mitigating Mobile Money Services Frauds in Rwanda\" (2022). CMU-Africa authors."
)
bullet(
    "\"Security Gaps in the Mobile Money System in Rwanda: Challenges, Risks "
    "and Mitigation\" (Springer 2024). Found the dominant attack types are "
    "SMishing, identity theft, phishing, and authentication attacks - "
    "corroborating four of our five categories with peer-reviewed evidence."
)

para("Implications for this capstone:")
numbered(
    "The Upanzi honeynet is, in effect, a working solution to my single biggest "
    "risk: labelled multilingual scam data."
)
numbered(
    "I will email Assane Gueye and Karen Sowon in Sprint 1 with a one-paragraph "
    "academic data-sharing collaboration request - subject to Decision D4 below. "
    "Worst case: no reply. Best case: real-world data and a reference."
)
numbered(
    "The Sowon S&P 2024 paper supersedes parts of my literature review and will "
    "be cited in the dissertation."
)
numbered(
    "Position my contribution as victim-facing application + comparative ML "
    "evaluation, complementing the Upanzi Network's scam-collection infrastructure."
)
numbered(
    "Given the 2-month window, the Upanzi data is best-effort, NOT on the critical "
    "path. The plan assumes their reply will not arrive in time."
)


# ============================================================
# 6. Reframing
# ============================================================
heading("6. Proposed Reframing - One Sentence", level=1)

callout("From:", "Build a custom ML classifier for West African scam detection.")
callout(
    "To:",
    "Compare API-based LLMs against classical ML baselines for multilingual scam "
    "classification across the five categories in West Africa, deploy the chosen "
    "approach in a pilot mobile app in one West African city, and treat "
    "deepfake-modality detection as a modular extension.",
)

para(
    "This single reframing resolves all three gaps (G1 to G3) and absorbs both "
    "new threads:"
)

table(
    ["Element", "How the reframing handles it"],
    [
        ["G1 - missing custom ML",
         "Classical baselines are the contribution; LLMs become the comparison arm."],
        ["G2 - scope creep",
         "Pilot stays in one West African city only; pan-African app is acknowledged as deployment infrastructure but not the unit of evaluation."],
        ["G3 - 11 langs vs 3",
         "Evaluation focuses on EN/FR/PCM; other 8 languages remain in the app as production features but are out of scope for accuracy claims."],
        ["Deepfakes",
         "Added as 6th category for text-only detection of \"voice note\" / \"video call\" lures; full audio/visual detection deferred to future work."],
        ["CMU Africa",
         "Upanzi work cited in lit review; honeynet dataset pursued in Sprint 1 best-effort; project positioned as complementary."],
    ],
)


# ============================================================
# 7. Five decisions
# ============================================================
heading("7. Five Decisions I Need from You", level=1)

table(
    ["#", "Decision", "My recommendation", "Tick"],
    [
        ["D1", "Adopt the LLM-vs-classical comparative reframing as the core ML contribution?", "Yes", ""],
        ["D2", "Lock pilot to ONE West African city (15 to 25 users) given the 2-month window; second city moves to Future Work?", "Yes", ""],
        ["D3", "Add synthetic_media_fraud as a sixth text-only category; defer audio/visual deepfake detection to Future Work?", "Yes", ""],
        ["D4", "Approve outreach to CMU-Africa Upanzi Network for academic data-sharing (best-effort, not on critical path)?", "Yes", ""],
        ["D5", "Update Chapter One Objective 2 to match the comparative-evaluation framing - replacing \"at least 85% accuracy with custom classifier\" with a comparative-evaluation objective?", "Yes", ""],
    ],
)

para(
    "If any answer is no, the rest of the two-month plan changes materially. I "
    "would prefer to resolve all five in this first meeting.",
    italic=True,
)


# ============================================================
# 8. Compressed sprint plan
# ============================================================
heading("8. Compressed Two-Month Sprint Plan (Post-Reframing)", level=1)

para(
    "The original proposal assumed a six-month plan. We have two. The app is "
    "already built, so most of that compresses cleanly. The pilot is the main "
    "scope cut."
)

table(
    ["Sprint", "Days", "Focus", "Exit gate"],
    [
        ["S1 - Foundation", "1-7",
         "This memo, supervisor sign-off, CMU outreach email sent, codebase recovery from 6-month dormancy, ML virtual environment, public datasets pulled (APWG, UCI SMS Spam, Kaggle phishing), PII masking, first labelled set (~500 examples).",
         "TF-IDF + LR baseline running end-to-end; D1-D5 signed off."],
        ["S2 - ML core", "8-14",
         "Random Forest baseline; LLM zero-shot evaluation for Gemini 1.5 Flash and GPT-4o-mini; per-language confusion matrices on EN/FR/PCM; error analysis; dataset card and model cards.",
         "Three-approach comparison table on held-out test set in ml/reports/."],
        ["S3 - App + pilot prep", "15-21",
         "Lock 6-category taxonomy in analyzeSuspiciousContent; log every classification (model, prompt, latency, output) to Firestore; deepfake-aware UI warning; consent forms, interview guide, recruitment plan, IRB submission if needed.",
         "Scanner emits structured incidents tagged with category + confidence + model. Pilot kit ready."],
        ["S4 - Pilot launch", "22-35",
         "Pilot in ONE city only, 15 to 25 users (down from 50 to 100 across 2 cities); SUS + Firebase Analytics; 5 to 8 mid-pilot interviews.",
         "Pilot wraps with usable data."],
        ["S5 - Analysis + writing", "36-49",
         "SUS scoring (Brooke 1996); thematic analysis (Braun and Clarke 2006); on-pilot ML accuracy via expert relabelling; chapters 4 (Results) and 5 (Discussion) drafted in parallel.",
         "All analysis done; chapters 4 and 5 drafted."],
        ["S6 - Submit + defend", "50-56",
         "Polish, demo video, defense rehearsal.",
         "Submitted, defended."],
    ],
)

heading("8.1 What This Compression Cuts vs. the Original 6-Month Proposal", level=2)

table(
    ["Element", "Original", "Compressed"],
    [
        ["Pilot users", "50-100", "15-25"],
        ["Pilot cities", "Lagos + Douala", "One only (whichever is accessible)"],
        ["Pilot duration", "8-10 weeks", "About 2 weeks"],
        ["ML baselines", "4 (TF-IDF x3 + transformer)", "2 (TF-IDF + LR; TF-IDF + RF)"],
        ["Transformer fine-tuning", "Yes (XLM-R / AfroXLMR)", "Cut - too expensive in time"],
        ["CMU-Africa data dependency", "Plan assumed available", "Send the email; assume no reply in time"],
        ["Interviews", "10-15", "5-8"],
    ],
)

heading("8.2 Implication for D2 (Pilot Scope Decision)", level=2)
para(
    "D2 in section 7 originally read \"lock pilot to Lagos and Douala only\" in "
    "earlier drafts of this memo. With the two-month constraint, D2 is revised: "
    "pilot in one city only. I will confirm the city in Sprint 1 based on what "
    "is physically accessible. The other city is documented as future work in "
    "the dissertation."
)


# ============================================================
# 9. What this memo does NOT propose
# ============================================================
heading("9. What This Memo Does NOT Propose to Change", level=1)

bullet("The mixed-methods research design in Unit 4 (still appropriate).")
bullet(
    "The annotated bibliography (will be extended, not redone - adding ~5 "
    "CMU-Africa citations and ~3 deepfake citations)."
)
bullet("The Pre-Capstone proposal's literature review themes.")
bullet("The Flutter + Firebase tech stack.")
bullet(
    "The Pre-Capstone proposal as a whole - it remains the foundation document; "
    "this memo is an addendum."
)


# ============================================================
# 10. References
# ============================================================
heading("10. Key References Introduced by This Memo", level=1)

para(
    "Lamptey, B. O., Gueye, A., Luhanga, E., Seidu, M., & Sowon, K. (2024). "
    "A honeynet infrastructure to battle SMS scammers. CMU-Africa Upanzi "
    "Network. https://www.africa.engineering.cmu.edu/news/2024/10/10-smishing.html"
)
para(
    "Sowon, K., et al. (2024). The role of user-agent interactions on mobile "
    "money practices in Kenya and Tanzania. IEEE Symposium on Security & "
    "Privacy (S&P 2024). https://www.cylab.cmu.edu/news/2024/07/10-navigating-"
    "digital-financial-inclusion-in-africa.html"
)
para(
    "Mitigating mobile money services frauds in Rwanda (2022). CMU-Africa. "
    "https://www.researchgate.net/publication/364959459_Mitigating_Mobile_Money_"
    "Services_frauds_in_Rwanda"
)
para(
    "Security gaps in the mobile money system in Rwanda: Challenges, risks "
    "and mitigation (2024). Springer. https://link.springer.com/chapter/10.1007/"
    "978-3-031-62277-9_42"
)
para(
    "CyLab-Africa initiative - joint CMU CyLab + CMU-Africa. "
    "https://africa.engineering.cmu.edu/research/cybersecurity/cylab/index.html"
)
para(
    "Upanzi Network at CMU-Africa. https://www.africa.engineering.cmu.edu/"
    "research/upanzi/index.html"
)


# ============================================================
# Save
# ============================================================
doc.save(str(OUTPUT))
print("Wrote:", OUTPUT)
