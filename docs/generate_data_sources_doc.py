"""Generate docs/DATA_SOURCES.docx — citable list of training-dataset sources.

Mirrors docs/DATA_SOURCES.md for school submission.
Run: python docs/generate_data_sources_doc.py
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path(__file__).resolve().parent / "DATA_SOURCES.docx"

ALU_BLUE = RGBColor(0x1A, 0x3C, 0x5E)
LINK_BLUE = RGBColor(0x0B, 0x57, 0xD0)

doc = Document()

# base style
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)


def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = ALU_BLUE
    return h


def link(label, url):
    """A bullet line: label + url rendered as a clickable-looking blue line."""
    p = doc.add_paragraph()
    p.add_run(label + ": ").bold = True
    r = p.add_run(url)
    r.font.color.rgb = LINK_BLUE
    r.underline = True
    return p


def body(text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    return p


def bullet(text):
    return doc.add_paragraph(text, style="List Bullet")


# ---- Title ----
title = doc.add_heading("Dataset Sources — Rethicsec Scam-Detection Model", level=0)
for run in title.runs:
    run.font.color.rgb = ALU_BLUE

intro = doc.add_paragraph()
intro.add_run(
    "This document lists every raw dataset used to train the project's "
    "scam-detection classifier (classes: advance_fee_fraud, mobile_money_fraud, "
    "phishing, not_a_scam), with citable source links. Provenance tags shown "
    "below are stored verbatim in each record's source_url field in ml/data/raw/."
)

# ---- Datasets used ----
heading("Datasets Used in Training", level=1)

heading("1. UCI SMS Spam Collection v.1", level=2)
body(
    "English SMS messages (5,574) labelled ham/spam. Provided English smishing "
    "and legitimate-SMS examples."
)
link("Dataset page", "https://archive.ics.uci.edu/dataset/228/sms+spam+collection")
link(
    "Direct download",
    "https://archive.ics.uci.edu/ml/machine-learning-databases/00228/smsspamcollection.zip",
)
body(
    "Citation: Almeida, T.A., Gomez Hidalgo, J.M., Yamakami, A. (2011). "
    "Contributions to the Study of SMS Spam Filtering: New Collection and "
    "Results. ACM DOCENG'11."
)

heading("2. Nazario Phishing Corpus", level=2)
body("Phishing emails (.mbox archives). Provided the phishing class.")
link("Project page", "https://monkey.org/~jose/phishing/")
link("Landing/wiki", "https://monkey.org/~jose/wiki/doku.php?id=PhishingCorpus")
body("Note: distribution is gated behind a manual download form.")

heading("3. SMS Phishing Dataset (Dataset_5971) — Mendeley Data", level=2)
body("Smishing/phishing SMS with URL/email/phone indicator flags.")
link("Mendeley Data", "https://data.mendeley.com/datasets/f45bkkt8pr")
body("DOI: 10.17632/f45bkkt8pr")
body("Authors: Mishra, Sandhya; Soni, Devpriya.")
body("Provenance tag: mendeley:f45bkkt8pr")

heading("4. MOZ-Smishing (Mozambique mobile-money SMS)", level=2)
body(
    "Portuguese M-Pesa mobile-money SMS. Provided the mobile-money-fraud class "
    "and African/regional coverage."
)
link("Hugging Face", "https://huggingface.co/datasets/MOZNLP/MOZ-Smishing")
body("Provenance tag: hf:MOZNLP/MOZ-Smishing")

heading("5. Premium Times Regional News Stream (supplementary)", level=2)
body(
    "West African scam/fraud news context, harvested via the public WordPress "
    "REST API (/wp-json/wp/v2/posts?search=<query>). Provided regional "
    "advance-fee / scam context."
)
link("Site", "https://www.premiumtimesng.com")

# ---- Datasets to be added ----
heading("Datasets Being Added to the Corpus", level=1)
body(
    "The following real, African scam/SMS datasets are being incorporated to "
    "strengthen African-language and African-English coverage. Both require "
    "relabelling from their native binary labels into the project's four-class "
    "taxonomy (advance_fee_fraud / mobile_money_fraud / phishing / not_a_scam), "
    "and their Kaggle licence must be confirmed before redistribution."
)

heading("6. Swahili SMS Detection Dataset (BongoScam)", level=2)
body(
    "1,508 real Tanzanian Swahili SMS labelled scam / trust. Adds Swahili "
    "(East African) mobile-money smishing coverage to a currently "
    "Portuguese-only mobile-money class. Native labels: scam / trust."
)
link("Kaggle", "https://www.kaggle.com/datasets/henrydioniz/swahili-sms-detection-dataset")
link("Repo", "https://github.com/Henryle-hd/BongoScamDetection")
body(
    "Supporting literature: arXiv:2502.16947 (Using Machine Learning to Detect "
    "Fraudulent SMSs in Swahili); Uncovering SMS Spam in Swahili Text Using "
    "Deep Learning Approaches."
)
body("Status: adopted for the corpus; licence verification + relabel pending.")

heading("7. ExAIS_SMS Spam Dataset", level=2)
body(
    "5,240 real African-English SMS (2,350 spam / 2,890 ham) collected from 20 "
    "consenting members of the Federal University of Agriculture, Abeokuta, "
    "Nigeria; sensitive details masked. African-English context is directly "
    "relevant to the English-language smishing gap. Native labels: Spam / Ham."
)
link("Kaggle", "https://www.kaggle.com/datasets/ysfbil/exais-sms-dataset")
link("Original repo", "https://github.com/AbayomiAlli/SMS-Spam-Dataset")
body(
    "Citation: Onashoga, A.S., Abayomi-Alli, O.O., Sodiya, A.S., Ojo, D.A. "
    "(2015). An Adaptive and Collaborative Server-Side SMS Spam Filtering Scheme "
    "Using Artificial Immune System. Information Security Journal: A Global "
    "Perspective, 24(4-6), 133-145."
)
body("Status: candidate; licence verification + relabel pending.")

# ---- Future dataset ----
heading("Future / Requested Dataset (Not Yet in Training)", level=1)
heading("CMU-Africa Upanzi Smishing-Honeynet — English Mobile-Money Smishing", level=2)
body(
    "Real English-language mobile-money smishing SMS, intended to close the "
    "cross-lingual MoMo gap (the mobile_money_fraud class is currently "
    "Lusophone-only and misclassifies English MoMo messages as phishing). The "
    "dataset is gated and not publicly downloadable."
)
bullet(
    "Source paper: Lamptey, Gueye, Seidu, Luhanga, Sowon. Smishing honeynet "
    "study. COMPASS '24. DOI: 10.1145/3674829.3675080"
)
bullet("Honeynet coverage: Rwanda, Botswana, Ghana, Kenya, Uganda.")
p = doc.add_paragraph(style="List Bullet")
r = p.add_run("Data-access request status: SENT 2026-06-22, awaiting reply. ")
r.bold = True
p.add_run("Request directed to Edith Luhanga (CMU-Africa), cc Karen Sowon.")
bullet("Request letter: docs/outreach/CMU_Upanzi_data_request.md")

# ---- Evaluated and rejected ----
heading("Datasets Evaluated and Rejected (Out of Scope)", level=1)
body(
    "These were assessed but are out of scope for a text classifier: they "
    "contain transaction records, not scam-message text, so the model cannot "
    "train on them. Listed for transparency / due diligence."
)
p = doc.add_paragraph(style="List Bullet")
p.add_run("PaySim (Kaggle, ealaxi/paysim1): ").bold = True
p.add_run(
    "~6.36M synthetic mobile-money transaction records (amounts, balances, "
    "transfer types, isFraud) seeded from real African logs; CC BY-SA 4.0. No "
    "message text. Would only be relevant if scope extended to a "
    "transaction-level fraud model."
)
p = doc.add_paragraph(style="List Bullet")
p.add_run("mobile-money-fraud-detection (github.com/antann07): ").bold = True
p.add_run(
    "synthetic mobile-money transaction records / engineered behavioural "
    "features (SIM-swap, device/location, velocity). No message text."
)

# ---- Attempted but not used ----
heading("Sources Attempted but NOT Used", level=1)
body(
    "The following regional government advisory sources were probed but yielded "
    "no training data (blocked), and are not counted among the sources used:"
)
bullet("ngCERT (cert.gov.ng) — Cloudflare anti-bot, HTTP 403")
bullet("ANTIC (antic.cm) — host unreachable (connect timeout)")
bullet("EFCC (efccnigeria.org) — empty shell / 404")

doc.save(OUT)
print(f"Wrote {OUT}")
