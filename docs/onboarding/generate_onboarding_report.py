"""Generate the Capstone Project Onboarding Report.

A single comprehensive Word document for a fresh-eyes reader: problem context,
vision, current state, gaps, reframing strategy, phase plan, risks, and
concrete next steps. Written so a stranger could read it once and understand
the whole project.

Output: docs/onboarding/Capstone_Onboarding_Report.docx
"""
from __future__ import annotations

import os
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

if sys.stdout.encoding and sys.stdout.encoding.lower() != "utf-8":
    try:
        sys.stdout.reconfigure(encoding="utf-8")
    except Exception:
        pass

OUTPUT = Path(__file__).resolve().parent / "Capstone_Onboarding_Report.docx"
OUTPUT.parent.mkdir(parents=True, exist_ok=True)

doc = Document()

# ============================================================
# Page setup + default style
# ============================================================
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(12)
normal.paragraph_format.line_spacing = 1.15
normal.paragraph_format.space_after = Pt(6)


# ============================================================
# Helpers
# ============================================================
def add_page_number(section):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._element.append(fld_begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run._element.append(instr)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._element.append(fld_end)
    for r in p.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(10)


def add_header(section, text):
    h = section.header
    p = h.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    run.italic = True


def heading(text, level=1):
    sizes = {1: 16, 2: 13, 3: 12}
    spaces_before = {1: 18, 2: 12, 3: 8}
    size = sizes.get(level, 12)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(spaces_before.get(level, 8))
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = True
    return p


def para(text, bold=False, italic=False, size=12, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def bullet(text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    p.paragraph_format.space_after = Pt(3)
    for r in p.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)
    return p


def numbered(text):
    p = doc.add_paragraph(text, style="List Number")
    p.paragraph_format.space_after = Pt(3)
    for r in p.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)
    return p


def callout(label, text):
    """Visually distinct labelled paragraph (e.g. 'Why this matters: ...')."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.2)
    r1 = p.add_run(label + " ")
    r1.font.name = "Arial"
    r1.font.size = Pt(11)
    r1.bold = True
    r1.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)
    r2 = p.add_run(text)
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(12)
    return p


def table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    try:
        t.style = "Light Grid Accent 1"
    except KeyError:
        t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.font.name = "Arial"
        run.font.size = Pt(11)
        run.bold = True
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            cell = t.rows[ri].cells[ci]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
    doc.add_paragraph()  # spacing after table
    return t


# ============================================================
# Footer + header on all sections
# ============================================================
for sec in doc.sections:
    add_page_number(sec)
    add_header(sec, "Capstone Onboarding Report")


# ============================================================
# Title page
# ============================================================
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("CAPSTONE PROJECT")
r.font.name = "Arial"
r.font.size = Pt(22)
r.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Onboarding Report")
r.font.name = "Arial"
r.font.size = Pt(18)
r.bold = True
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("AI-Powered Mobile Application for Cybercrime Reporting\nand Scam Detection in West Africa")
r.font.name = "Times New Roman"
r.font.size = Pt(14)
r.italic = True

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("A complete walkthrough for a fresh-eyes reader")
r.font.name = "Times New Roman"
r.font.size = Pt(12)

for _ in range(5):
    doc.add_paragraph()

for label, value in [
    ("Author", "Wilsons Navid Wado Tiwa"),
    ("Programme", "BSc Software Engineering, Final Year Capstone"),
    ("Institution", "African Leadership University"),
    ("Date", "5 May 2026"),
    ("Window", "2 months (target ~5 July 2026)"),
    ("Status", "Sprint 1 of 6 - Foundation"),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(label + ": ")
    r1.font.name = "Arial"
    r1.font.size = Pt(11)
    r1.bold = True
    r2 = p.add_run(value)
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(11)

doc.add_page_break()


# ============================================================
# 1. Executive Summary
# ============================================================
heading("1. Executive Summary", level=1)

para(
    "This is a final-year capstone in software engineering with a "
    "**two-month execution window** (target completion around 5 July 2026). The "
    "project is an AI-powered mobile application that helps people in West Africa "
    "report cybercrime, get instant guidance about how serious it looks, and "
    "learn how to stay safer online. The pilot will run with 15 to 25 users in "
    "one West African city (down from the proposal's original 50 to 100 across "
    "Lagos and Douala — see section 8 for why). The work is positioned against "
    "the Millennium Project's Global Challenge 12, which asks how the world can "
    "stop transnational organised crime networks from growing in power and "
    "sophistication."
)

para(
    "All of the planning is done. Units 1 through 4 of the academic sequence are "
    "complete, the annotated bibliography has 22 sources, the methodology chapter "
    "is locked, and the Pre-Capstone Major Assessment proposal has been written and "
    "submitted with full UML and architecture diagrams. The mobile app itself is "
    "also already built, branded as Rethicssec (folder name still says rethicsai), "
    "currently at version 1.0.2 build 3, with a finalised Privacy Policy and "
    "Play Store assets."
)

para(
    "What is not yet done is the academic centrepiece of the year: the machine "
    "learning research. The proposal promised a custom-trained classifier on West "
    "African scam data, hitting at least 85% accuracy across five categories. The "
    "app instead routes everything through Vertex AI Gemini Flash and OpenAI "
    "GPT-4o-mini. There is no measured accuracy and no fixed taxonomy in code. The "
    "supervisor will spot this in five minutes, so it has to be addressed head-on."
)

para(
    "The recommended response is a single reframing: treat the deployed LLMs as "
    "one arm of a comparative study, build classical ML baselines as the second "
    "arm, and add a sixth category for synthetic-media (deepfake) fraud as a "
    "modular extension. Reach out to the Upanzi Network at Carnegie Mellon Africa "
    "for collaboration on smishing data; their published work covers exactly this "
    "problem space. This reframing rescues the proposal-versus-built-app gap, "
    "absorbs the supervisor's two new asks (deepfakes and CMU Africa research), "
    "and gives the project a research artefact strong enough to defend."
)

para(
    "The remainder of this report walks through the problem we are solving, what "
    "exactly is in the codebase today, where the gaps are, what the year looks "
    "like phase by phase, what could go wrong, and what to do in the first week.",
    italic=True,
)

doc.add_page_break()


# ============================================================
# 2. The Problem
# ============================================================
heading("2. The Problem We Are Solving", level=1)

heading("2.1 The Bigger Picture", level=2)

para(
    "Africa is in the middle of a digital growth surge. Mobile phone ownership has "
    "spread faster than internet infrastructure could mature, financial services "
    "have moved onto SIM cards instead of bank branches, and hundreds of millions "
    "of people who never went online before are now spending hours a day on "
    "WhatsApp, Facebook, and mobile money. That growth is not slowing down."
)

para(
    "Cybercrime has scaled with it. The African Union estimates that cybercrime "
    "costs African economies around four billion US dollars every year (African "
    "Union, 2022). Globally, the figure is now in the trillions and projected to "
    "exceed ten trillion dollars annually by the end of 2025 (Morgan, 2022). Most "
    "of those numbers come from data that does not even include the African "
    "majority, because most African incidents are never formally reported."
)

para(
    "INTERPOL's 2021 African Cyberthreat Assessment found that fewer than 20% of "
    "cybercrime incidents on the continent are formally reported. In countries "
    "with limited digital reporting infrastructure, the real figure is almost "
    "certainly lower. That underreporting is not just a measurement problem. It "
    "directly enables the criminals: when no one reports, no one investigates, no "
    "one prosecutes, and the networks grow."
)

heading("2.2 Why West Africa Specifically", level=2)

para(
    "Nigeria and Cameroon, our two pilot countries, sit at the centre of this "
    "crisis. Nigeria is Africa's largest economy and most populous country, with "
    "over 100 million internet users (Internet World Stats, 2023). It has been "
    "associated for decades with advance-fee fraud, the so-called \"419 scams\" "
    "named after the relevant section of the Nigerian Criminal Code. Today the "
    "playbook has evolved into business email compromise, romance scams, "
    "cryptocurrency fraud, and phishing campaigns delivered through WhatsApp and "
    "SMS. Nigeria's Economic and Financial Crimes Commission reported over $500 "
    "million in cybercrime-related losses in 2022 alone (EFCC, 2023)."
)

para(
    "Cameroon is a different story but a parallel one. With about 15 million "
    "internet users and a fast-growing mobile money ecosystem, the country has "
    "seen a surge in mobile money fraud, SIM swap attacks, and social media "
    "scams. Cameroon's bilingual nature, split between French-speaking and "
    "English-speaking populations, adds a real complication for awareness "
    "campaigns and reporting infrastructure. The National Agency for Information "
    "and Communication Technologies (ANTIC) has documented the rise but has no "
    "public-facing digital reporting platform of its own."
)

heading("2.3 Who Gets Hurt", level=2)

para(
    "The victims are not random. Women, the elderly, and people going online for "
    "the first time are disproportionately targeted (Akanle et al., 2020). These "
    "are also the people least likely to come forward and report what happened, "
    "because of stigma, distrust of authorities, and the simple fact that they "
    "do not know where to turn. The result is a cycle: those who suffer most "
    "are also least visible to anyone who could help, which makes them more "
    "attractive targets, and the cycle deepens."
)

heading("2.4 The Five Scam Categories", level=2)

para(
    "The project narrows its focus to five scam types that together account for "
    "the bulk of cybercrime cases in Nigeria and Cameroon:"
)

bullet(
    "Advance-fee fraud (\"419 scams\"). The classic Nigerian fraud. Victim is "
    "promised a large payout in exchange for a small upfront fee, then strung "
    "along through fake legal documents, fake bank communications, and fake "
    "intermediaries until they stop paying."
)
bullet(
    "Mobile money fraud. SIM swap attacks, fake reversal requests, agent fraud, "
    "PIN harvesting. Targets M-Pesa, MTN Mobile Money, Airtel Money. Fast-growing "
    "category as more financial life moves onto SIM cards."
)
bullet(
    "Phishing (and smishing). Fake messages or websites designed to capture "
    "credentials. SMS phishing (smishing) is particularly common in Africa "
    "because mobile-first usage means SMS still carries a lot of trust."
)
bullet(
    "Romance scams. Long-cultivated fake relationships that end in financial "
    "exploitation. Often run from West Africa against diaspora and overseas "
    "victims, but increasingly also against local victims."
)
bullet(
    "Identity theft. Stealing personal information (BVN, ID numbers, biometrics) "
    "for impersonation, account takeover, or fraudulent registration."
)

para(
    "We are also adding a sixth category, synthetic media fraud, to capture the "
    "deepfake-enabled attacks that are now appearing in the region: AI-generated "
    "voice notes used to authorise mobile money transfers, fabricated video calls "
    "in romance scams, AI-generated photo evidence in advance-fee fraud. More on "
    "this in section 7."
)

heading("2.5 Why Existing Solutions Are Not Enough", level=2)

para(
    "Two existing tools have tried to solve parts of this problem and fall short:"
)

bullet(
    "ScamAdviser is a web-based tool that lets users check whether a website is "
    "trustworthy. It is useful for desktop checks but has no dedicated mobile "
    "reporting feature, no support for African languages, and was not designed "
    "around the scam patterns common in West African markets."
)
bullet(
    "Trend Micro Fraud Buster is a mobile app that uses AI to scan messages and "
    "URLs. Detection works, but the app does not support victim reporting, does "
    "not draw on West African data, and has no African-language coverage."
)

para(
    "Local institutional channels have their own gaps. Nigeria's EFCC and the "
    "Nigeria Police Force Cybercrime Unit require in-person visits or formal "
    "written complaints, processes that are slow, intimidating, and effectively "
    "out of reach for many citizens (Adeniran, 2008). In Cameroon, ANTIC is the "
    "designated national cybersecurity body but has no public-facing digital "
    "reporting platform, and most citizens have never heard of it. Cameroon's "
    "cybercrime legislation is also incomplete, leaving victims with limited legal "
    "recourse even when they do come forward (ANTIC, 2022)."
)

para(
    "What is missing across the board is an integrated, mobile-first, multilingual "
    "platform that combines reporting, AI-driven risk assessment, and educational "
    "resources tailored for the West African context. That is the gap this project "
    "is built to fill."
)

doc.add_page_break()


# ============================================================
# 3. The Vision and Goal
# ============================================================
heading("3. The Vision: What Success Looks Like", level=1)

para(
    "By the end of the two-month window, the project should have produced six things:"
)

numbered(
    "A working, deployed mobile app that 15 to 25 pilot users in one West African "
    "city actually use over a roughly two-week testing period."
)
numbered(
    "Quantitative evidence of the app working: machine learning model accuracy on "
    "the five (or six) scam categories, latency measurements showing the 30-second "
    "risk-assessment SLA is met, and System Usability Scale (SUS) scores from "
    "every pilot user."
)
numbered(
    "Qualitative evidence of the app being useful: thematic analysis of 5 to 8 "
    "semi-structured user interviews covering whether the reporting felt safe, "
    "whether the AI guidance made sense, whether the educational content helped, "
    "and what users wanted to change."
)
numbered(
    "A research artefact strong enough to defend: a comparative ML report "
    "evaluating classical models versus API-based LLMs on the same data, with "
    "per-language breakdowns, error analysis, dataset card, and model cards."
)
numbered(
    "A defended dissertation drawing all of the above into one cohesive narrative."
)
numbered(
    "Ideally: an institutional collaboration with the Upanzi Network at "
    "Carnegie Mellon Africa, both for data sharing and for academic visibility."
)

para(
    "The pilot is not the whole project. It is a feedback mechanism. The app and "
    "the ML research are independently valuable, and the pilot tells us whether "
    "what we built is actually usable by the people we built it for."
)

heading("3.1 The Five Original Objectives", level=2)
para("From Chapter One of the Pre-Capstone proposal, slightly trimmed:")

t = table(
    ["#", "Objective"],
    [
        ["1", "Build a cross-platform Flutter mobile application with multilingual support."],
        ["2", "Train an ML model classifying reports into five scam categories with at least 85% accuracy."],
        ["3", "Integrate a Firebase-powered real-time risk assessment with a 30-second response SLA."],
        ["4", "Build an in-app educational resource module of at least 20 articles, guides, and videos."],
        ["5", "Pilot with 50 to 100 users in Lagos and Douala over a two-month testing window."],
    ],
)

para(
    "Objective 2 is the one being changed. The reframing replaces \"custom "
    "classifier with at least 85% accuracy\" with \"comparative evaluation of "
    "classical ML versus API-based LLMs across the five (or six) categories.\" "
    "All other objectives stand. See section 7 for the full reframing argument.",
    italic=True,
)

doc.add_page_break()


# ============================================================
# 4. What Already Exists
# ============================================================
heading("4. What I Already Have", level=1)

heading("4.1 The Academic Foundation", level=2)
para(
    "Everything in the academic sequence has been completed and submitted. All "
    "files now live under the docs/ folder of the workspace."
)

table(
    ["Deliverable", "File", "What it covers"],
    [
        ["Unit 1", "docs/coursework/Unit_One_Draft - Copy.docx",
         "Chapter One: introduction, background, problem statement, five SMART objectives, five research questions, scope, and significance."],
        ["Unit 2", "docs/coursework/WilsonsNavidWadoTiwa-Unit Two Assignment.docx",
         "Refined Unit 1, polished and resubmitted."],
        ["Unit 3", "docs/coursework/unit3/WilsonsNavidWadoTiwa_Unit_Three Assignment.docx",
         "Annotated bibliography of 22 scholarly sources covering cybercrime in Africa, ML for cybercrime detection, mobile crime reporting, digital literacy, transnational crime, and Flutter/Firebase development."],
        ["Unit 4", "docs/coursework/unit4/WilsonsNavidWadoTiwa_Unit Four Assignment.docx",
         "Methodology chapter: mixed-methods design (quantitative ML metrics + SUS scores; qualitative thematic analysis of interviews), sampling strategy, data collection instruments, ethical framework. A humanised version sits in docs/coursework/unit4_humanized/."],
        ["Pre-Capstone Major Assessment", "docs/coursework/Major-assesment/WilsonsNavidWadoTiwa_Pre-Capstone_Research_Proposal_v2.docx",
         "Full proposal: literature review across seven themes, system analysis, three-tier architecture, all six UML and architecture diagrams (Agile model, system architecture, ERD, class diagram, use case, mixed-methods diagram), Agile sprint plan, ethical considerations."],
    ],
)

heading("4.2 The Mobile Application (Rethicssec)", level=2)

para(
    "The app lives in-tree at mobile/rethicsai/ as of 2026-05-05. It was "
    "previously at C:\\Users\\LENOVO\\Desktop\\rethicsai\\rethicsai\\ and was "
    "consolidated into the workspace; the original location is preserved as a "
    "read-only backup and should not be edited. The Privacy Policy and Play "
    "Store assets sit one level up at mobile/. The app is more mature than a "
    "typical capstone build:"
)

table(
    ["Aspect", "Detail"],
    [
        ["Display name", "Rethicssec (renamed from RethicsAI; folder still uses old name)"],
        ["Version", "1.0.2+3"],
        ["Frontend", "Flutter 3.24+, Dart 3+, Material 3"],
        ["Architecture", "Clean Architecture per feature (data / domain / presentation)"],
        ["State management", "BLoC + Provider, GetIt for DI, Freezed for models"],
        ["Local storage", "Hive (offline-first) + SharedPreferences + flutter_secure_storage"],
        ["Backend", "Firebase: Auth, Firestore, Storage, Cloud Functions, Messaging, Analytics, Crashlytics"],
        ["Languages supported", "11 - en, sw, fr, ar, ha, yo, ig, zu, xh, af, sawa (Duala)"],
        ["Privacy Policy", "Finalised, in PRIVACY_POLICY.md and PRIVACY_POLICY.docx"],
        ["Play Store", "Descriptions + readiness guide written; ready to submit"],
    ],
)

para("The app has eleven feature modules:")
bullet("auth - sign-in, sign-up, password reset, social login")
bullet("dashboard - home screen with feature grid, quick stats, recent activity")
bullet("ai_assistant - the Wilson AI chat interface (named after the author)")
bullet("incidents - structured incident reporting with evidence upload")
bullet("cases - case tracking and history")
bullet("education - tips, guides, and video content for digital safety")
bullet("scanner - threat scanning for URLs, messages, and files")
bullet("emergency - one-tap access to local cybersecurity contacts")
bullet("notifications - push and in-app alerts")
bullet("settings - profile, language, preferences")
bullet("admin - back-office dashboard for content management and case review")

heading("4.3 The Backend (Cloud Functions)", level=2)

para(
    "The backend is a TypeScript Firebase Cloud Functions project at "
    "mobile/rethicsai/functions/. Six HTTPS endpoints are deployed:"
)

table(
    ["Function", "Backend model", "Purpose"],
    [
        ["wilsonChat", "OpenAI GPT-4o-mini", "Main chat endpoint for the Wilson AI assistant"],
        ["analyzeSuspiciousContent", "OpenAI GPT-4o-mini", "Threat analysis for URLs / messages submitted via the scanner"],
        ["getCyberInsights", "OpenAI GPT-4o-mini", "Daily auto-generated tips for the dashboard"],
        ["wilsonAIVertex", "Vertex AI Gemini 1.5 Flash", "Enhanced chat with African cybersecurity context priming"],
        ["getAfricanThreatIntelligence", "Vertex AI Gemini 1.5 Flash", "Region-specific threat updates"],
        ["generateSecurityTraining", "Vertex AI Gemini 1.5 Flash", "On-demand training content generation"],
    ],
)

callout(
    "Important:",
    "All six functions route to third-party LLMs (Gemini and OpenAI). There is "
    "no custom-trained ML model anywhere in the codebase. No Scikit-learn, no "
    "TensorFlow, no TFLite. The proposal promised a custom classifier with at "
    "least 85% accuracy across five categories. None of that exists. This is the "
    "single most important gap to address.",
)

heading("4.4 The Workspace", level=2)
para(
    "The Capstone-Project workspace at C:\\Users\\LENOVO\\Desktop\\Capstone-Project\\ "
    "was scaffolded on 4 May 2026. It is the place where everything other than "
    "the app itself lives: planning docs, ML research code, supervisor meeting "
    "notes, pilot artefacts, and the dissertation. Layout:"
)

bullet("README.md - status board, single source of truth")
bullet(".gitignore - tuned for Python, ML data, and Word lock files")
bullet("docs/ - all planning deliverables (Units 1-4, Pre-Capstone proposal, this report)")
bullet("mobile/ - README pointing to the existing Flutter project")
bullet("backend/ - README pointing to the existing Cloud Functions")
bullet("ml/ - the machine learning research workspace (the missing piece)")
bullet("ml/src/ - skeleton code: dataset.py, preprocessing.py, baselines.py, llm_baselines.py, eval.py")
bullet("ml/data/, ml/notebooks/, ml/models/, ml/reports/ - data and outputs")
bullet("meetings/ - one markdown file per supervisor meeting")
bullet("pilot/ - placeholder for consent forms, interview guides, transcripts, analysis")
bullet("dissertation/ - placeholder for the final report")

para(
    "There is also a Stop hook configured at .claude/hooks/memory-nudge.ps1 that "
    "ensures progress is saved to long-term memory at the end of every Claude "
    "session, so a fresh session can pick up cold without re-explanation."
)

doc.add_page_break()


# ============================================================
# 5. The Critical Gap
# ============================================================
heading("5. The Gap Between Proposal and Reality", level=1)

para(
    "The Pre-Capstone proposal and the built app diverge in three load-bearing "
    "ways. None of them are bad in isolation. Together they are the central "
    "problem of Phase 0:"
)

table(
    ["Element", "Proposal commitment", "Built reality"],
    [
        ["Pilot scope", "Lagos + Douala only", "Pan-African positioning"],
        ["Languages", "English, French, Pidgin English (3)",
         "11 languages including Sawa (Duala)"],
        ["ML approach", "Custom Scikit-learn / TensorFlow classifier",
         "Vertex Gemini Flash + GPT-4o-mini API calls"],
        ["ML target", "At least 85% accuracy on 5 categories",
         "No measured accuracy, no fixed taxonomy in code"],
        ["Risk SLA", "30 seconds via Firebase Cloud Function",
         "Latency depends on third-party API"],
    ],
)

para(
    "Why this matters: a supervisor reading the proposal first and then opening "
    "the app will spot the gap in five minutes. Capstones are penalised hard for "
    "\"you said X, you did Y, you did not acknowledge it.\" They reward \"you said "
    "X, you did Y, here is exactly why and what changed.\" The reconciliation memo "
    "(docs/RECONCILIATION_MEMO.md) is the tool for the second framing."
)

para(
    "The biggest gap by severity is the missing custom ML. Without it, the word "
    "\"AI-powered\" in the project title is delivered entirely by Google and "
    "OpenAI. There is no contribution to the West African ML literature, no "
    "evaluation that could be defended in a viva, and no answer to a basic "
    "\"how good is your model?\" question. The fix is in section 7."
)

doc.add_page_break()


# ============================================================
# 6. New Threads from the Supervisor
# ============================================================
heading("6. Two New Threads from the Supervisor", level=1)

heading("6.1 Deepfakes", level=2)

para(
    "The supervisor flagged deepfakes as a current threat vector. Two facts to "
    "hold simultaneously:"
)

para(
    "Operationally, AI-generated voice (voice cloning), AI-generated video (face-"
    "swapped video calls in romance scams), and AI-generated images (fake KYC "
    "documents) are now in the wild against African mobile money users. Voice "
    "cloning in particular has been used to defeat the social trust that mobile "
    "money confirmations rely on: a victim hears what sounds like their child or "
    "their boss authorising a transaction, and complies."
)

para(
    "For this capstone, however, detecting deepfakes is a different ML problem "
    "from classifying scam text. It would require audio models for voice spoofing "
    "(such as Wav2Vec2), or vision models for face-swap detection. That is a "
    "separate research thread, not a tweak. Trying to do both within one capstone "
    "year would dilute everything."
)

para("The recommended posture is to treat deepfakes as a modular extension:")

bullet(
    "Add a sixth category, synthetic_media_fraud, to the text-only classification "
    "taxonomy. Even text-only classifiers can pick up the giveaway language "
    "(\"listen to this voice note from your son and confirm,\" \"video call me "
    "now to verify\"). This is cheap and earns a 2026-current dimension in the "
    "research."
)
bullet(
    "Add a focused chapter section in the dissertation on the deepfake threat "
    "landscape in Africa, citing recent voice-cloning and face-swap incidents. "
    "Frame full audio/visual deepfake detection as future work."
)
bullet(
    "Optionally, prototype a single deepfake-aware feature in the app: a \"this "
    "audio looks AI-generated\" warning using an off-the-shelf detector. One "
    "screen, one feature, no new research thread."
)

heading("6.2 CMU Africa Research on Mobile Money Scams", level=2)

para(
    "This is the highest-leverage external thread in the entire project. The "
    "Upanzi Network at Carnegie Mellon University Africa, based at the Kigali "
    "campus, conducts directly relevant research on mobile money scams. Several "
    "publications matter:"
)

bullet(
    "Lamptey, Gueye, Luhanga, Seidu, and Sowon (2024) built an inexpensive "
    "honeynet that actively collects smishing messages from scammers across "
    "Rwanda, Botswana, Ghana, Kenya, and Uganda. This is, in effect, a working "
    "solution to the project's single biggest risk: labelled multilingual scam "
    "data from the right region. (See https://www.africa.engineering.cmu.edu/news/2024/10/10-smishing.html)"
)
bullet(
    "Sowon and colleagues (2024) presented \"The Role of User-Agent Interactions "
    "on Mobile Money Practices in Kenya and Tanzania\" at IEEE Symposium on "
    "Security & Privacy 2024, drawing on 72 interviews with mobile money users. "
    "This paper supersedes parts of our literature review and should be cited in "
    "the dissertation."
)
bullet(
    "\"Mitigating Mobile Money Services Frauds in Rwanda\" (CMU Africa, 2022) "
    "and \"Security Gaps in the Mobile Money System in Rwanda: Challenges, Risks "
    "and Mitigation\" (Springer, 2024) both identify SMishing, identity theft, "
    "phishing, and authentication attacks as the dominant attack types. This "
    "corroborates four of our five categories with peer-reviewed evidence."
)

para("Implications for this project, in order of priority:")

numbered(
    "Email Assane Gueye and Karen Sowon at CMU-Africa in week 1 of Phase 1 with "
    "a one-paragraph academic data-sharing collaboration request. Worst case: no "
    "reply. Best case: real-world West and East African smishing data and a "
    "reference letter."
)
numbered(
    "Cite the Upanzi work in any update to the lit review. Their work becomes the "
    "strongest precedent, replacing the much weaker Chieloka and Ugwu (2021)."
)
numbered(
    "Position this project as complementary to the Upanzi Network: \"victim-facing "
    "application + comparative ML evaluation, complementing the Upanzi Network's "
    "scam-collection infrastructure.\" Frames the work as part of a broader effort, "
    "not a competitor."
)

doc.add_page_break()


# ============================================================
# 7. The Reframing Strategy
# ============================================================
heading("7. The Reframing in One Sentence", level=1)

callout(
    "From",
    "Build a custom ML classifier for West African scam detection.",
)
callout(
    "To",
    "Compare API-based LLMs against classical ML baselines for multilingual scam "
    "classification across the five categories in West Africa, deploy the chosen "
    "approach in a pilot mobile app in Lagos and Douala, and treat deepfake-"
    "modality detection as a modular extension.",
)

para(
    "This single reframing solves four problems at once: it gives the existing "
    "LLM-based app a research role (comparison arm), it adds the missing "
    "classical ML rigor (the second arm), it absorbs deepfakes (sixth category), "
    "and it gives a natural opening for the CMU Africa data collaboration."
)

heading("7.1 Why It Is Better Research Than the Original Plan", level=2)

para(
    "The original plan, train a Random Forest on West African data and report "
    "85% accuracy, would have produced a respectable but not particularly "
    "interesting capstone. In 2026, every reviewer will ask: \"why did you not "
    "use an LLM?\" Now we have an answer baked into the research design."
)

para(
    "More than that, the comparative framing addresses a genuinely open research "
    "question that the African NLP community cares about: when do classical "
    "models still outperform foundation models on low-resource African languages, "
    "and at what cost? Sawa (Duala) is supported in our app and in almost no "
    "translation system. Per-language evaluation across 11 languages, even just "
    "on the three core ones (English, French, Pidgin English), gives evidence "
    "the field is missing."
)

heading("7.2 The Five Decisions Pending Supervisor Sign-off", level=2)

para(
    "All five decisions are in docs/RECONCILIATION_MEMO.md and need to be "
    "resolved in the first supervisor meeting:"
)

table(
    ["#", "Decision", "Recommended"],
    [
        ["D1", "Adopt the LLM-vs-classical comparative reframing as the core ML contribution?", "Yes"],
        ["D2", "Lock pilot to ONE city (15-25 users) given the 2-month window; second city moves to Future Work?", "Yes"],
        ["D3", "Add synthetic_media_fraud as a sixth text-only category, defer audio/visual deepfake detection to Future Work?", "Yes"],
        ["D4", "Approve outreach to CMU-Africa Upanzi Network for academic data-sharing (best-effort, not on the critical path)?", "Yes"],
        ["D5", "Update Chapter One Objective 2 to match the comparative-evaluation framing?", "Yes"],
    ],
)

doc.add_page_break()


# ============================================================
# 8. The Phase Plan
# ============================================================
heading("8. The Year, Phase by Phase", level=1)

para(
    "The capstone year is broken into five phases. Each phase has a clearly "
    "defined exit gate so the supervisor can sign off before the next phase "
    "begins. Phase numbering starts at 0 because Phase 0 is foundational rather "
    "than productive."
)

table(
    ["Phase", "Weeks", "Focus", "Exit gate"],
    [
        ["0. Reconcile", "2",
         "This onboarding report, the reconciliation memo, the supervisor meeting, the dataset acquisition plan, the CMU outreach email.",
         "All five D1-D5 decisions signed off in writing."],
        ["1. ML R&D", "8",
         "Dataset acquisition (own + Upanzi if granted). Build classical baselines: TF-IDF + Logistic Regression, Random Forest, Gradient Boosting. Add a transformer baseline (XLM-R or AfroXLMR fine-tuned). Run zero-shot and few-shot LLM evaluation for both Gemini and OpenAI. Augment dataset with deepfake-text examples.",
         "ml/reports/ contains a dataset card, model cards for at least four baselines, and a decision memo on which approach goes into the production app."],
        ["2. App refinement", "4",
         "Lock the six-category taxonomy in analyzeSuspiciousContent. Log every classification (model, prompt, latency, output) to Firestore for offline re-scoring. Add a deepfake-aware UI warning. Make sure incident reports emit one of the six categories.",
         "Scanner emits structured incidents tagged with category + confidence + model used."],
        ["3. Pilot", "8-10",
         "Deploy in Lagos + Douala. Recruit 50 to 100 users through community tech hubs, university campuses, NGOs. Run two-week minimum usage period. Collect SUS questionnaire from every user. Run 10 to 15 semi-structured interviews. Capture Firebase Analytics. Re-score user-submitted reports against ground truth.",
         "Cleaned pilot dataset, thematic analysis writeup, mean SUS score with breakdown by city, on-pilot ML accuracy."],
        ["4. Writing + defense", "6",
         "Write the dissertation chapters. Polish the ML report. Record the demo video. Run defense rehearsals.",
         "Dissertation submitted, capstone successfully defended."],
    ],
)

heading("8.1 Phase 0 in Detail (the next two weeks)", level=2)
para("Before anything else, do these things:")
numbered(
    "Read this report end to end. Read README.md and docs/RECONCILIATION_MEMO.md. "
    "Skim the Pre-Capstone proposal at docs/coursework/Major-assesment/. You should be able "
    "to explain the reframing in one paragraph from memory."
)
numbered(
    "Schedule the first supervisor meeting. Bring the reconciliation memo. Walk "
    "through D1 to D5 one by one. Get a written record (an email recap is fine) "
    "of what was agreed."
)
numbered(
    "Set up the Python virtual environment for the ML workspace: cd ml; python "
    "-m venv .venv; .venv\\Scripts\\activate; pip install -r requirements.txt."
)
numbered(
    "Draft the outreach email to Assane Gueye and Karen Sowon at CMU-Africa. "
    "Keep it under 200 words. Introduce the project, explain the comparative ML "
    "study, ask whether the Upanzi honeynet dataset is shareable for academic use "
    "with proper attribution. Show the email to the supervisor before sending."
)
numbered(
    "Begin dataset acquisition planning. Pull what is publicly available: APWG "
    "eCrime corpora, UCI SMS Spam Collection, Kaggle phishing URL datasets. "
    "Document each source in a SOURCE.md as you pull it. Note licensing terms."
)
numbered(
    "Initialise the workspace as a git repository (git init in Capstone-Project). "
    "Do NOT init inside mobile/rethicsai/ - that is already a git repo (its "
    "own .git/ travelled with the consolidation). From "
    "now on, commit incrementally with meaningful messages. Defense examiners "
    "love being able to see how thinking evolved."
)

heading("8.2 Cadence", level=2)
para(
    "Weekly supervisor meetings are non-negotiable. Use the template at "
    "meetings/_template.md. Every meeting covers four sections in order:"
)
bullet("Last week: what I committed to, what I delivered, evidence")
bullet("Blockers: specific questions I need help with")
bullet("Reading and learning: what was assigned, what I took away, follow-up questions")
bullet("This week: what I am committing to")

para(
    "One markdown file per meeting at meetings/YYYY-MM-DD.md. These files become "
    "the project's working memory and, eventually, the appendix of the "
    "dissertation describing how the work evolved."
)

doc.add_page_break()


# ============================================================
# 9. Risks
# ============================================================
heading("9. Risks and Honest Assessment", level=1)

heading("9.1 What Genuinely Worries Me", level=2)

para("Six concerns, in rough order of severity:")

para("9.1.1 Time itself.", bold=True)
para(
    "Two months for a final-year capstone with a working app, ML research, "
    "pilot, dissertation, and defense is extremely tight. Anything that slips "
    "more than a few days will eat into the pilot or the writing window, both "
    "of which are already compressed. Treat every sprint exit gate as a hard "
    "checkpoint. If a gate slips, the next sprint must absorb the slip - no "
    "stretching the calendar."
)

para("9.1.2 Data scarcity.", bold=True)
para(
    "Without West African labelled scam data, every accuracy claim in the "
    "dissertation is fragile. The Chieloka and Ugwu (2021) paper has the only "
    "directly relevant precedent and uses a small dataset. The CMU Africa "
    "Upanzi honeynet would be the answer, but academic email response times "
    "make it unrealistic to count on within two months. Plan as if the data is "
    "not coming: pull public corpora in Sprint 1, label a small in-house seed "
    "set fast, augment via paraphrase if needed."
)

para("9.1.3 Code dormancy.", bold=True)
para(
    "The last meaningful commit in the rethicsai repo was 3 November 2025. The "
    "codebase will be nearly six months stale by the time it is reopened. Budget "
    "two to three days of Sprint 1 for codebase recovery alone: dependency "
    "updates, breakage from Flutter SDK changes, possibly Firebase config drift. "
    "Do not treat the first half of Sprint 1 as productive in terms of new "
    "features."
)

para("9.1.4 Single squashed git commit.", bold=True)
para(
    "The rethicsai repo has exactly one commit: \"Initial commit - Rethics AI "
    "v1.0.2\". This is a defensive vulnerability. Examiners often ask \"show me "
    "how this evolved\" and there is currently nothing to show. Do not try to "
    "fake history backward. From Sprint 1 forward, commit incrementally, with "
    "meaningful messages, dated and tagged."
)

para("9.1.5 Scope creep.", bold=True)
para(
    "The app expanded to pan-African ambitions, but the academic deliverables "
    "and the pilot only support West Africa specifically. Either we narrow our "
    "claims (the recommended path) or we accept that some claims will be "
    "aspirational rather than evidenced. The reconciliation memo handles this "
    "explicitly."
)

para("9.1.6 LLM dependency.", bold=True)
para(
    "Relying on Gemini Flash and GPT-4o-mini means our \"AI\" is provided by "
    "Google and OpenAI, not by us. The classical baseline comparison saves the "
    "research story, but we still need to confront the cost question: if we are "
    "deploying LLMs in production, the dissertation has to defend that as a "
    "pragmatic choice with measured trade-offs, not as a hidden assumption."
)

heading("9.2 What I Think Is Genuinely Strong", level=2)

para("Five things to leverage:")

bullet(
    "The problem is real, well-documented, and matters. African Union and "
    "INTERPOL numbers carry weight. The lived reality of mobile money fraud "
    "and 419 scams is undeniable in West Africa."
)
bullet(
    "The academic foundation is unusually solid. Most capstone students arrive "
    "at the execution year having only sketched a problem statement. We have "
    "four polished units, an annotated bibliography, a complete proposal with "
    "diagrams, and a locked methodology."
)
bullet(
    "The app is a production-quality artefact already shipped. Most capstone "
    "students never get to Play Store readiness. We have a Privacy Policy and "
    "a v1.0.2 build."
)
bullet(
    "The multilingual coverage, especially Sawa (Duala), is itself a research "
    "artefact. Sawa is barely supported by any translation system. Per-language "
    "evaluation on Sawa, even just for awareness, is publishable as a curiosity."
)
bullet(
    "The tech stack is right for the market. Flutter on Android-first, Firebase "
    "for backend, LLM APIs for AI: pragmatic, fast, defensible. Not flashy, but "
    "not naive either."
)

heading("9.3 The Honest Net Assessment", level=2)

para(
    "This is a strong-to-very-strong capstone if the reframing is locked in early "
    "and Phase 1 ML R&D is treated as the academic centrepiece, not a checkbox. "
    "The app is already a defensible artefact. What is missing, and entirely "
    "buildable in eight weeks of focused work, is the research artefact to match "
    "it. Without the reframing, the project risks landing as \"good app, thin "
    "research.\" With it, the project lands as \"good app, real research, "
    "honest about its trade-offs.\" That is the difference between a clear "
    "distinction and a competent pass."
)

doc.add_page_break()


# ============================================================
# 10. How to Proceed (Week 1)
# ============================================================
heading("10. How to Proceed - Sprint 1 (Days 1-7)", level=1)

para(
    "Concrete actions for the first seven days. The exit gate at the end of "
    "Sprint 1 is non-negotiable: TF-IDF + LR baseline running end-to-end and "
    "supervisor sign-off on D1-D5. Everything else in Sprint 1 supports those "
    "two outputs."
)

table(
    ["Day", "Action", "Output"],
    [
        ["1", "Read this report end to end. Read README.md and RECONCILIATION_MEMO.md. Skim docs/coursework/Major-assesment/.",
         "Mental model in place."],
        ["1-2", "Email supervisor to schedule the first meeting. Attach the reconciliation memo.",
         "Meeting on the calendar."],
        ["2-3", "Recover the rethicsai codebase: flutter pub get, attempt to build, document any breakage. Same for functions/.",
         "Codebase builds again."],
        ["3", "Set up the ML Python venv: cd ml; python -m venv .venv; .venv\\Scripts\\activate; pip install -r requirements.txt. Try importing src/baselines.py.",
         "Working ML environment."],
        ["3-4", "First supervisor meeting. Walk through D1-D5. Get written sign-off.",
         "Five decisions resolved."],
        ["4-5", "Draft and send the CMU-Africa outreach email (Gueye + Sowon). Show to supervisor first.",
         "Email out (don't wait for reply)."],
        ["4-6", "Pull public datasets: APWG eCrime, UCI SMS Spam, Kaggle phishing. Document each in SOURCE.md. Apply PII masking via src/preprocessing.py.",
         "Raw + processed data in ml/data/."],
        ["5-7", "Hand-label or weakly-label about 500 examples across the six categories.",
         "First labelled set on disk."],
        ["6-7", "Train TF-IDF + LR baseline. Evaluate. Even if accuracy is poor, end-to-end pipeline must work.",
         "First model card draft in ml/reports/."],
        ["7", "git init the workspace. First meaningful commit. Push to GitHub or GitLab if you want backup.",
         "Workspace under version control."],
    ],
)

para(
    "By the end of Sprint 1 (day 7), you should be able to truthfully say: I "
    "have supervisor sign-off on the research direction, the codebase builds "
    "again, I have a working Python environment, I have raw data on disk, I "
    "have a reference institution that knows the project exists, I have version "
    "control running, and I have a baseline ML model trained end-to-end. Those "
    "seven things together unlock Sprint 2.",
    italic=True,
)

doc.add_page_break()


# ============================================================
# 11. Glossary
# ============================================================
heading("11. Glossary", level=1)

table(
    ["Term", "Meaning"],
    [
        ["419 scam", "Advance-fee fraud, named after section 419 of the Nigerian Criminal Code"],
        ["ANTIC", "Agence Nationale des Technologies de l'Information et de la Communication - Cameroon's cybersecurity agency"],
        ["BLoC", "Business Logic Component - Flutter's recommended state management pattern"],
        ["BVN", "Bank Verification Number - Nigeria's bank identity system, frequent target of identity theft"],
        ["CyberGuard", "Internal name once used for the AI assistant; replaced by \"Wilson\""],
        ["EFCC", "Economic and Financial Crimes Commission - Nigeria's main financial-crime body"],
        ["Flutter", "Google's cross-platform UI framework, single codebase for Android and iOS"],
        ["Gemini Flash", "Google's lower-cost Vertex AI LLM, currently used for Wilson chat"],
        ["GPT-4o-mini", "OpenAI's lower-cost model, currently used for the Wilson AI fallback"],
        ["MoMo / mobile money", "Phone-based money transfer (M-Pesa, MTN MoMo, Airtel Money, etc.)"],
        ["ngCERT", "Nigeria Computer Emergency Response Team"],
        ["Pidgin English (pcm)", "English-based creole spoken across West Africa, especially Nigeria"],
        ["Routine Activity Theory", "Cohen and Felson's 1979 framework: crime occurs when motivated offender + suitable target + absent guardian converge"],
        ["Sawa (Duala)", "Bantu language spoken around Douala, Cameroon; supported in Rethicssec"],
        ["SIM swap attack", "Attacker convinces a mobile carrier to port the victim's phone number to a SIM they control"],
        ["Smishing", "SMS phishing - phishing delivered through text messages"],
        ["SUS", "System Usability Scale - 10-item questionnaire scored 0 to 100"],
        ["Synthetic media fraud", "Scams using AI-generated voice, video, or images (deepfakes)"],
        ["TAM", "Technology Acceptance Model (Davis 1989) - explains adoption via perceived usefulness and ease of use"],
        ["TFLite", "TensorFlow Lite - mobile-optimised model format for on-device inference"],
        ["Upanzi Network", "Carnegie Mellon Africa's research network on digital systems and cybersecurity"],
        ["Vertex AI", "Google Cloud's managed AI platform (hosts Gemini)"],
        ["Wilson AI", "User-facing name of the in-app AI assistant; named after the author"],
        ["XLM-R / AfroXLMR", "Multilingual transformer models; AfroXLMR is fine-tuned for African languages"],
    ],
)


doc.add_page_break()


# ============================================================
# 12. Where Everything Lives
# ============================================================
heading("12. Where Everything Lives", level=1)

para("A quick reference for finding things:")

table(
    ["Looking for...", "Path"],
    [
        ["This report", "C:\\Users\\LENOVO\\Desktop\\Capstone-Project\\docs\\onboarding\\Capstone_Onboarding_Report.docx"],
        ["The reconciliation memo", "C:\\Users\\LENOVO\\Desktop\\Capstone-Project\\docs\\RECONCILIATION_MEMO.md"],
        ["Workspace status board", "C:\\Users\\LENOVO\\Desktop\\Capstone-Project\\README.md"],
        ["Pre-Capstone proposal", "C:\\Users\\LENOVO\\Desktop\\Capstone-Project\\docs\\coursework\\Major-assesment\\WilsonsNavidWadoTiwa_Pre-Capstone_Research_Proposal_v2.docx"],
        ["Methodology chapter", "C:\\Users\\LENOVO\\Desktop\\Capstone-Project\\docs\\coursework\\unit4\\WilsonsNavidWadoTiwa_Unit Four Assignment.docx"],
        ["Annotated bibliography", "C:\\Users\\LENOVO\\Desktop\\Capstone-Project\\docs\\coursework\\unit3\\WilsonsNavidWadoTiwa_Unit_Three Assignment.docx"],
        ["The Flutter app", "C:\\Users\\LENOVO\\Desktop\\Capstone-Project\\mobile\\rethicsai\\ (original backup at C:\\Users\\LENOVO\\Desktop\\rethicsai\\, read-only)"],
        ["The Cloud Functions", "C:\\Users\\LENOVO\\Desktop\\Capstone-Project\\mobile\\rethicsai\\functions\\"],
        ["ML research code", "C:\\Users\\LENOVO\\Desktop\\Capstone-Project\\ml\\src\\"],
        ["ML requirements", "C:\\Users\\LENOVO\\Desktop\\Capstone-Project\\ml\\requirements.txt"],
        ["Supervisor meeting template", "C:\\Users\\LENOVO\\Desktop\\Capstone-Project\\meetings\\_template.md"],
        ["Privacy Policy (final)", "C:\\Users\\LENOVO\\Desktop\\Capstone-Project\\mobile\\PRIVACY_POLICY.docx"],
        ["UML diagrams", "C:\\Users\\LENOVO\\Desktop\\Capstone-Project\\docs\\coursework\\Major-assesment\\diagrams\\"],
    ],
)


doc.add_page_break()


# ============================================================
# 13. Closing Note
# ============================================================
heading("13. Closing Note", level=1)

para(
    "The hardest thing about this kind of project is not the technology. It is "
    "remembering, every week, that the goal is to help people who are right now "
    "being defrauded out of money they cannot afford to lose. The 419 victim, "
    "the elderly woman who lost her pension to a romance scam, the small trader "
    "in Douala whose mobile money agent ran off with her float - these are not "
    "abstractions. They are the audience. Every design decision should be "
    "stress-tested against the question: would this actually help that person?"
)

para(
    "The technology is interesting and the research is interesting, but they are "
    "in service of that audience. Keep that in view through Sprint 2's late "
    "nights debugging Scikit-learn pipelines, through Sprint 4's frustrations of "
    "field recruitment, through Sprint 5's exhausted dissertation revisions. "
    "When the supervisor asks \"why does this matter?\" the right answer is not "
    "\"because the proposal said so.\" It is to describe one specific person and "
    "what the app would do for them."
)

para(
    "Everything else is implementation."
)


# ============================================================
# Save
# ============================================================
doc.save(str(OUTPUT))
print("Wrote:", OUTPUT)
