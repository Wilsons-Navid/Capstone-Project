"""Generates the Rethics product & brand report as a Microsoft Word .docx.
Author: Wilsons Navid Wado Tiwa. Run: python docs/generate_rethics_report.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---- Brand palette (the app's earth-tone identity) ----
BROWN = RGBColor(0x2D, 0x1B, 0x14)   # primary
AMBER = RGBColor(0x7A, 0x5C, 0x00)   # AA-safe amber (text)
GREEN = RGBColor(0x2E, 0x7D, 0x32)   # success / safe
DANGER = RGBColor(0xB3, 0x26, 0x1E)  # threat
GREY = RGBColor(0x49, 0x45, 0x4F)
LIGHTBROWN_HEX = "EFE7DE"
AMBER_HEX = "F3E9CC"

doc = Document()

# Base style
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor(0x22, 0x22, 0x22)


def shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hex_color)
    tcPr.append(sh)


def h(text, level=1, color=BROWN):
    p = doc.add_heading(level=level)
    run = p.add_run(text)
    run.font.color.rgb = color
    run.font.name = "Calibri"
    return p


def para(text, bold=False, italic=False, color=None, size=11, align=None, space_after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p


def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        r.font.color.rgb = BROWN
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def callout(title, body, accent=AMBER, fill=AMBER_HEX):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    shade(cell, fill)
    p = cell.paragraphs[0]
    tr = p.add_run(title + "  ")
    tr.bold = True
    tr.font.color.rgb = accent
    p.add_run(body)
    doc.add_paragraph()
    return tbl


def kv_table(headers, rows, header_fill=LIGHTBROWN_HEX, widths=None):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Light Grid Accent 1"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.rows[0].cells
    for i, htext in enumerate(headers):
        shade(hdr[i], header_fill)
        run = hdr[i].paragraphs[0].add_run(htext)
        run.bold = True
        run.font.color.rgb = BROWN
        run.font.size = Pt(10)
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            r = cells[i].paragraphs[0].add_run(str(val))
            r.font.size = Pt(10)
    if widths:
        for row in tbl.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return tbl


# ============================ TITLE ============================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run("RETHICS")
tr.bold = True
tr.font.size = Pt(40)
tr.font.color.rgb = BROWN
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run("Africa's scam-defence companion")
sr.italic = True
sr.font.size = Pt(15)
sr.font.color.rgb = AMBER
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
mr = meta.add_run("Product & Brand Report  •  v1.0.6  •  June 2026\nPrepared by Wilsons Navid Wado Tiwa")
mr.font.size = Pt(10)
mr.font.color.rgb = GREY
doc.add_paragraph()

# ============================ EXEC SUMMARY ============================
h("1. Executive Summary", 1)
para(
    "Rethics is a mobile-first cybersecurity companion built for the people most exposed to "
    "digital fraud and least served by existing tools: everyday mobile users across Africa. "
    "It turns a confusing, frightening moment — “Is this message a scam?” — into a clear, "
    "instant verdict and a concrete next action. Where most security products stop at detection, "
    "Rethics closes the loop: it explains the risk, lets the user act on it, routes a report to the "
    "right national authority, and teaches the user to recognise the next attack."
)
para(
    "The product pairs a custom-trained machine-learning classifier (not merely a wrapped large "
    "language model) with a polished Flutter application, a Firebase backend, an in-app AI assistant, "
    "an education hub, and an authority-reporting directory covering 14 African countries. This report "
    "documents what Rethics does, why each capability matters, how the AI was trained, where it falls "
    "short today, and the potential ahead.",
    space_after=10,
)
callout("THE ONE-LINE PITCH",
        "Rethics is the app that tells you a scam is a scam — in your language, in seconds — and helps "
        "you do something about it.")

# ============================ PROBLEM ============================
h("2. The Problem — and Why It Matters", 1)
para(
    "Digital fraud in Africa is not a fringe nuisance; it is a mass-market epidemic riding on the same "
    "rails that brought financial inclusion. Mobile money put a bank in every pocket — and a target on "
    "every SIM. The dominant attack vectors are deeply local:"
)
bullet("schemes that impersonate mobile-money operators, banks, telcos and government agencies.",
       bold_lead="Mobile-money fraud: ")
bullet("the classic “you've won / you've inherited / pay a small fee to release funds” trap, "
       "re-skinned for WhatsApp and SMS.", bold_lead="Advance-fee fraud: ")
bullet("look-alike links and credential-harvesting pages targeting bank and wallet logins.",
       bold_lead="Phishing & smishing: ")
para(
    "Why it matters: the victims are disproportionately first-time internet users, the elderly, and "
    "small traders for whom a single successful scam can wipe out weeks of income. Existing antivirus and "
    "spam tools are built for Western inboxes and English desktop email — they do not understand a Twi "
    "smishing text, a Swahili MoMo lure, or a Pidgin investment scam. Rethics exists to fill exactly that gap.",
    space_after=10,
)

# ============================ POSITIONING ============================
h("3. Product Vision & Brand Positioning", 1)
para("Rethics is positioned at the intersection of three jobs users are trying to get done:")
kv_table(
    ["User job", "What Rethics promises", "Brand proof"],
    [
        ["“Tell me if this is safe.”", "An instant, explainable verdict on any pasted message, link, email or number.",
         "Scanner + custom ML classifier"],
        ["“Help me do something.”", "One-tap actions: verify, block, and report to the real authority in my country.",
         "Action bar + 14-country authority directory"],
        ["“Help me not fall for it again.”", "Short lessons, threat intel, and an assistant that answers in my language.",
         "Education hub + Wilson AI assistant"],
    ],
    widths=[2.0, 3.2, 1.8],
)
para("Brand personality", bold=True, color=BROWN)
para(
    "Trustworthy, warm, and calm under pressure — never alarmist, never cold-corporate. The visual identity "
    "uses an African earth-tone palette (deep brown, sunset amber, acacia green) rather than the generic "
    "neon-on-black of typical security apps, signalling that this product is of and for its users, not imported."
)
callout("WHY THE BRAND CHOICE IS STRATEGIC",
        "Security software earns its right to exist on trust. An interface that feels local, legible and "
        "human lowers the barrier for exactly the non-technical users who need it most.", accent=GREEN, fill="E3EFE3")

# ============================ FEATURES ============================
h("4. The Product — Full Feature Inventory", 1)
para(
    "Rethics ships as a production Flutter application (v1.0.6) backed by Firebase and a Python ML service. "
    "The following eleven feature modules are live in the app today."
)
kv_table(
    ["Module", "What it does", "Why it matters"],
    [
        ["Scanner", "Paste an SMS, email, URL, phone number or social-media text; the AI returns a threat level, category and explanation.",
         "The core value: turns uncertainty into a verdict in seconds."],
        ["Report-to-Authorities", "Country-aware directory of police / cyber-crime / financial-crime units (14 countries); call, email or report online with a pre-filled message.",
         "Converts a verdict into real-world action — the step competitors skip."],
        ["Incident Reporting", "Structured report with evidence upload, geolocation, priority and currency.",
         "Builds an auditable record and, at scale, a regional threat dataset."],
        ["Case Tracking", "Lifecycle tracking of submitted cases with status timeline and filters.",
         "Gives victims visibility and closure instead of a dead-end form."],
        ["Education Hub", "Lessons (video + interactive) with gamification and certificates.",
         "Prevention: the cheapest scam to stop is the one the user spots first."],
        ["Wilson AI Assistant", "Conversational cyber-safety Q&A (Claude Haiku) with in-context guidance.",
         "A patient, always-on expert for users with nobody else to ask."],
        ["Emergency Contacts", "Personal + national emergency directory, Firebase-backed and admin-editable.",
         "Help is one tap away during an active fraud."],
        ["Notifications", "Real-time inbox with read/unread and batch actions; push delivery.",
         "Keeps users alert to new threats and case updates."],
        ["Dashboard", "Personal security overview: reports, resolved cases, threats blocked.",
         "Makes safety legible and rewarding — a reason to come back."],
        ["Admin Console", "Governance: contacts, content moderation, cases, analytics, threat & learning management.",
         "Lets the team curate authority data and content without shipping a release."],
        ["Localization & Settings", "11 locales (English, French, Swahili, Hausa, Yoruba, Igbo, Zulu, Xhosa, Afrikaans, Arabic, Duala).",
         "Meets users in their own language — a precondition for trust and reach."],
    ],
    widths=[1.5, 3.0, 2.5],
)

para("4.1  The scanner in action — worked examples", bold=True, color=BROWN, size=12)
para("The classifier returns one of four categories with a risk level. Representative behaviour:")
kv_table(
    ["Example message (paraphrased)", "Category", "Verdict"],
    [
        ["“Congrats! You won ₦2,000,000 in the MTN promo. Send your BVN + ₦5,000 activation fee to claim.”",
         "Advance-fee fraud", "HIGH RISK"],
        ["“Your MoMo account will be blocked. Dial *123*PIN# now to verify your wallet.”",
         "Mobile-money fraud", "HIGH RISK"],
        ["“Dear customer, your bank account is suspended. Click http://bit.ly/secure-restore to reactivate.”",
         "Phishing", "HIGH RISK"],
        ["“Hi, are we still meeting at 3pm tomorrow at the office?”",
         "Not a scam", "SAFE"],
    ],
    widths=[3.8, 1.6, 1.3],
)
callout("DESIGN DETAIL THAT MATTERS",
        "A verdict is never shown by colour alone — every result pairs a colour with an icon and a text "
        "label, and meets WCAG AA contrast. Security guidance is worthless if a user can't read it.",
        accent=GREEN, fill="E3EFE3")

# ============================ THE AI ============================
h("5. The AI Engine — Training, Drawbacks & Potential", 1)
para("5.1  Training approach & architecture", bold=True, color=BROWN, size=12)
para(
    "Rethics deliberately owns its core intelligence rather than renting it. The classifier is a custom, "
    "four-class scam detector (advance-fee fraud, mobile-money fraud, phishing, not-a-scam) trained on a "
    "hand-labelled corpus assembled from public scam datasets (Nazario phishing, an SMS smishing set) and "
    "African-context sources. The model is a soft-voting ensemble of two complementary views of the text:"
)
bullet("a classical TF-IDF + linear model that captures the tell-tale surface tokens of scams "
       "(“activation fee”, “verify your PIN”, shortened links).", bold_lead="Lexical signal: ")
bullet("multilingual e5-small sentence embeddings that capture meaning across languages, so a lure "
       "phrased in Swahili or Pidgin can still be recognised.", bold_lead="Semantic signal: ")
para(
    "On a held-out test split the ensemble reaches a macro-F1 of 0.955 — strong, balanced performance "
    "across the four classes in-distribution. The model is served behind an API and integrated into the "
    "app's scanner, with the app pre-warming the model when the scanner opens to avoid a cold-start delay."
)
kv_table(
    ["Metric", "Value", "Reading"],
    [
        ["Classes", "4 (advance-fee, mobile-money, phishing, not-a-scam)", "Scoped to Africa's dominant vectors"],
        ["Test macro-F1", "0.955", "Strong, balanced — in-distribution"],
        ["Architecture", "TF-IDF + LR ⊕ e5-small embeddings (soft-vote)", "Lexical + multilingual semantic"],
        ["Serving", "API + in-app pre-warm", "Real-time, avoids cold-start"],
    ],
    widths=[1.7, 3.1, 2.2],
)

para("5.2  Honest drawbacks", bold=True, color=DANGER, size=12)
para("A credible product names its limits. Today's model has four real weaknesses:")
bullet("the training corpus is skewed toward phishing, with mobile-money and advance-fee examples in "
       "the minority. The model inherits a majority-class (phishing) bias: some genuine mobile-money or "
       "inheritance scams are mislabelled as phishing, and even some legitimate messages drift toward "
       "“scam” at low confidence.", bold_lead="Class imbalance: ")
bullet("the headline 0.955 is in-distribution. On out-of-distribution messages (new lures, unusual "
       "phrasings) accuracy drops, and controlled experiments showed that re-balancing strategies alone "
       "did not fix the same hard cases — the real fix is more authentic minority-class data, not algorithmic tricks.",
       bold_lead="Out-of-distribution fragility: ")
bullet("if the served model is slow to respond on first call, the scanner falls back to heuristics, which "
       "are weaker than the model. Reliability engineering (warm-up, timeouts) is as important as model accuracy.",
       bold_lead="Cold-start fallback: ")
bullet("the model reads text. It cannot yet judge images, voice notes, QR codes or the reputation of a "
       "live URL — all growing scam channels.", bold_lead="Modality & coverage: ")
callout("THE CORE DATA TRUTH",
        "The single biggest lever on quality is not a better algorithm — it is more real, labelled, "
        "African scam data, especially for mobile-money and advance-fee fraud. Data is the moat.",
        accent=DANGER, fill="F6E0DE")

para("5.3  Potential", bold=True, color=GREEN, size=12)
bullet("every confirmed report enriches a proprietary, regionally-specific scam dataset that no "
       "competitor can replicate by adding a language pack. The product gets smarter as it is used.",
       bold_lead="A data flywheel: ")
bullet("surfacing the model's confidence and deferring uncertain cases to a clear “unsure — treat "
       "with caution” state would raise trust and cut false positives immediately.", bold_lead="Confidence-aware UX: ")
bullet("a quantised on-device model would let Rethics screen messages offline and privately, crucial in "
       "low-connectivity and data-sensitive contexts.", bold_lead="On-device inference: ")
bullet("with user consent, verified reports could feed tiered escalation — from helping the user act "
       "(live today), to consent-based partnerships with operators and regulators, to automated alerts.",
       bold_lead="Escalation tiers: ")
bullet("image, link-reputation and voice-note analysis would close the channels the current text model "
       "cannot see.", bold_lead="Multi-modal expansion: ")

# ============================ ARCHITECTURE ============================
h("6. Technical Architecture", 1)
kv_table(
    ["Layer", "Technology", "Role"],
    [
        ["Mobile app", "Flutter (Dart), Material 3, 11 locales", "Cross-platform client; one codebase, Android-first"],
        ["Backend", "Firebase (Auth, Firestore, Cloud Functions, FCM)", "Identity, data, serverless logic, push"],
        ["AI assistant", "Claude Haiku via Cloud Functions", "Conversational cyber-safety guidance"],
        ["ML service", "Python (TF-IDF + e5 ensemble) behind an API", "The custom scam classifier"],
        ["Design system", "Token-based earth-tone theme + MASTER.md contract", "Consistent, accessible, on-brand UI"],
    ],
    widths=[1.5, 3.0, 2.5],
)

# ============================ DIFFERENTIATION ============================
h("7. Durable Differentiation", 1)
para("Rethics is defensible for reasons a competitor cannot copy by translating their UI:")
bullet("trained on and tuned for the continent's actual scam vectors and languages.", bold_lead="Africa-first intelligence: ")
bullet("detection plus action plus authority reporting — not a verdict that leaves the user stranded.",
       bold_lead="Closes the loop: ")
bullet("a custom classifier the team controls end-to-end, not a thin wrapper over someone else's API.",
       bold_lead="Owned ML: ")
bullet("every report compounds into a regional dataset that widens the moat over time.",
       bold_lead="Data network effect: ")

# ============================ ROADMAP ============================
h("8. Roadmap Snapshot", 1)
kv_table(
    ["Horizon", "Focus"],
    [
        ["Now (shipped)", "Scanner, 4-class model (0.955), 14-country reporting, education, assistant, admin console, 11 locales"],
        ["Next", "Confidence-aware verdicts, minority-class data collection, reliability hardening, broader push delivery"],
        ["Later", "On-device model, multi-modal (image/URL/voice) detection, consent-based authority partnerships"],
    ],
    widths=[1.8, 5.2],
)

# ============================ RISKS ============================
h("9. Risks & Limitations", 1)
bullet("scam quality scales with authentic minority-class data; acquiring it is the critical path.",
       bold_lead="Data scarcity: ")
bullet("SMS-reading permissions and unknown-developer sideloading are real publication hurdles to manage.",
       bold_lead="Distribution & permissions: ")
bullet("a security tool must never over-promise; clear confidence and caution states protect user trust.",
       bold_lead="Trust & false positives: ")

# ============================ CONCLUSION ============================
h("10. Conclusion", 1)
para(
    "Rethics is more than a scam detector; it is a trust layer for African digital life. It is technically "
    "credible (a custom, balanced classifier rather than an LLM veneer), strategically defensible (Africa-first "
    "data and an action-closing loop), and human in its design. Its honest limitations — class imbalance, "
    "out-of-distribution fragility, and the hunger for real local data — are precisely the problems that, once "
    "solved through usage, become the product's deepest competitive moat.",
    space_after=10,
)
para("Rethics — know the scam. Stop the scam.", bold=True, color=BROWN, align=WD_ALIGN_PARAGRAPH.CENTER, size=13)

import os
out = os.path.join(os.path.dirname(__file__), "Rethics_Product_Brand_Report.docx")
doc.save(out)
print("Saved:", out)
