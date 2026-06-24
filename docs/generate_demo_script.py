"""Generates the Rethicsec 5-minute demo video script as a Microsoft Word .docx."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

BROWN = RGBColor(0x2D, 0x1B, 0x14)
AMBER = RGBColor(0x7A, 0x5C, 0x00)
GREY = RGBColor(0x49, 0x45, 0x4F)
HDR_HEX = "EFE7DE"
QUOTE_HEX = "F3E9CC"

doc = Document()
n = doc.styles["Normal"]
n.font.name = "Calibri"; n.font.size = Pt(11)


def shade(cell, hexc):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), hexc)
    tcPr.append(sh)


def h(t, lvl=1, color=BROWN):
    p = doc.add_heading(level=lvl); r = p.add_run(t); r.font.color.rgb = color; r.font.name = "Calibri"


def para(t, bold=False, italic=False, color=None, size=11, align=None):
    p = doc.add_paragraph(); r = p.add_run(t)
    r.bold = bold; r.italic = italic; r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    if align: p.alignment = align
    return p


def check(t):
    doc.add_paragraph(t, style="List Bullet")


def quote(text):
    tbl = doc.add_table(rows=1, cols=1); tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = tbl.rows[0].cells[0]; shade(c, QUOTE_HEX)
    r = c.paragraphs[0].add_run(text); r.italic = True; r.font.size = Pt(11)
    doc.add_paragraph()


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, ht in enumerate(headers):
        shade(t.rows[0].cells[i], HDR_HEX)
        rr = t.rows[0].cells[i].paragraphs[0].add_run(ht); rr.bold = True; rr.font.color.rgb = BROWN; rr.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            rr = cells[i].paragraphs[0].add_run(str(v)); rr.font.size = Pt(9)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()


# Title
tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = tp.add_run("Rethicsec — 5-Minute Demo Video Script"); tr.bold = True; tr.font.size = Pt(20); tr.font.color.rgb = BROWN
sp = doc.add_paragraph(); sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sp.add_run("Shot list for the Implementation & Testing submission"); sr.italic = True; sr.font.color.rgb = AMBER
doc.add_paragraph()
para("Goal: demonstrate the core functionality under the rubric — different testing strategies, "
     "different data values, and performance — in about five minutes. Skip sign-up / sign-in "
     "(the rubric says so).", italic=True, color=GREY)

h("Before you hit record (2-minute prep)", 1)
check("Warm up the model first: open the Scanner and run any scan, then wait ~30–60s. The ML Space sleeps when idle; warming it up means your demo scans show the AI model verdict, not the heuristic fallback.")
check("Have the 4 example messages ready to paste (copy them from the table below into a notes app on the phone).")
check("Sign in before recording so you can open straight onto the dashboard.")
check("Screen-record in portrait, clean status bar; speak slowly and clearly.")
check("Have a second device or emulator ready for the 20-second performance shot (or record it separately).")

h("The script (≈5:00 total)", 1)
table(
    ["Time", "On screen", "Say this (narration)"],
    [
        ["0:00–0:20\n(Hook)", "App icon / splash, then dashboard.",
         "This is Rethicsec — a scam-defence app for Africa. Across the continent, mobile-money, advance-fee and phishing scams cost everyday users real money. Rethicsec answers one question instantly: is this message a scam? — and helps you act on it."],
        ["0:20–0:35\n(Orientation)", "Dashboard: stats, feature grid.",
         "When I open the app I land on my dashboard — reports filed, cases resolved, threats blocked — and quick access to every tool."],
        ["0:35–2:05\n(CORE: Scanner)", "Tap Scanner. Paste each of the 4 messages, scan, let the verdict appear.",
         "The heart of the app is the scanner. Let me test it with different kinds of messages. [advance-fee] Here's an advance-fee scam — the AI flags it HIGH RISK and tells me why. [repeat for mobile-money, phishing] [safe] And a normal message comes back SAFE — so it doesn't cry wolf."],
        ["2:05–2:50\n(CORE: Report)", "On a scam result, scroll to Report to authorities. Open the country dropdown (show full list). Tap Call / Email / Report online.",
         "A verdict isn't enough — Rethicsec helps me act. It shows the real cyber-crime and police units for my country, and I can switch country here; all 14 are covered. One tap calls them, emails a report, or opens the official portal, pre-filled."],
        ["2:50–3:20\n(Education)", "Open Education, show a lesson + progress.",
         "To stop the next scam, the education hub has short lessons, with progress and certificates — prevention, not just detection."],
        ["3:20–3:50\n(Assistant)", "Open Wilson, ask: 'How do I know if a MoMo message is fake?'",
         "And Wilson, the built-in assistant, answers cyber-safety questions in plain language, any time."],
        ["3:50–4:20\n(Admin)", "Admin → Emergency Contacts. Tap +, choose Add new country, fill a few fields, save.",
         "For administrators, the whole authority directory is editable in-app — I can add a brand-new country, update it, or delete it, with no app update."],
        ["4:20–4:40\n(Testing + HW)", "Terminal showing `flutter test` → All 36 passed. Then a quick clip on a second device/emulator.",
         "On the engineering side: the app has an automated test suite — unit and widget tests — that passes green, and it runs smoothly across different Android devices."],
        ["4:40–5:00\n(Close)", "Back to dashboard / logo.",
         "All three objectives are met: a labelled scam corpus, a working detection-and-reporting platform, and a custom classifier at 0.955 macro-F1. The next step is collecting more local scam data. Rethicsec — know the scam, stop the scam."],
    ],
    widths=[1.0, 2.3, 3.7],
)

h("The 4 scanner inputs (copy these to the phone)", 1)
table(
    ["#", "Type", "Message"],
    [
        ["1", "Advance-fee", "Congratulations! You have won 2,000,000 in the MTN promo. Send your BVN and a 5,000 activation fee to claim your prize now."],
        ["2", "Mobile-money", "Your MoMo account will be blocked today. Dial *123*PIN# now to verify your wallet and avoid suspension."],
        ["3", "Phishing", "Dear customer, your bank account has been suspended. Click http://bit.ly/secure-restore to reactivate immediately."],
        ["4", "Safe", "Hi, are we still meeting at 3pm tomorrow at the office?"],
    ],
    widths=[0.4, 1.3, 5.3],
)

h("Rubric mapping (so nothing is missed)", 1)
table(
    ["Rubric item", "Covered by"],
    [
        ["Different testing strategies", "4:20 shot — flutter test (unit + widget) + manual functional walkthrough"],
        ["Different data values", "0:35–2:05 — four different scanner inputs → four verdicts"],
        ["Performance on different hardware/software", "4:20 shot — app on a second device/emulator (also fill the README matrix)"],
        ["Analysis vs objectives", "4:40 close — Obj 1/2/3 recap"],
        ["Recommendation / future work", "4:40 close — collect more local scam data"],
    ],
    widths=[2.5, 4.5],
)

h("Analysis & Discussion — read verbatim (closing shot)", 1)
para("Option A — tight close (~20 seconds, for the 4:40–5:00 slot)", bold=True, color=BROWN)
quote("“To wrap up: all three objectives were met. I built a labelled, four-class scam corpus; I "
      "delivered a working detection-and-reporting platform that's deployed and installable; and I trained "
      "a custom classifier that reaches a macro-F1 of about 0.955. The honest limit is that this score "
      "holds best on familiar messages — newer or rarer scams are harder, mainly because the training data "
      "is still light on mobile-money and advance-fee examples. That's also the clearest way forward: more "
      "real, local scam data will lift accuracy the most. What makes Rethicsec matter is that it doesn't "
      "stop at a verdict — it turns detection into action, in the user's own language. Rethicsec: know the "
      "scam, stop the scam.”")

para("Option B — fuller analysis + discussion (~45 seconds, if you add a dedicated segment)", bold=True, color=BROWN)
quote("“Let me analyse the results against the proposal. Objective one, the labelled scam corpus, was "
      "achieved: roughly four thousand four hundred messages across four classes — advance-fee, "
      "mobile-money, phishing, and not-a-scam. Objective two, a working platform, was achieved and is what "
      "you've just seen: scanning, country-aware reporting, education, an assistant, and an admin console, "
      "deployed as an installable app. Objective three, the classifier, was also achieved — a TF-IDF and "
      "multilingual embedding ensemble at about 0.955 macro-F1.”")
quote("“Where it fell short of the ideal is on unfamiliar messages: the model inherits a bias toward "
      "phishing because that class dominates the data, so some genuine mobile-money scams get mislabelled. "
      "Importantly, re-balancing experiments did not fix those cases — which tells me the real constraint "
      "is the amount of authentic local data, not the algorithm.”")
quote("“Why does this milestone matter? Because detection alone doesn't protect anyone. The impact here "
      "is the full loop — a verdict a non-technical user can read, in their language, plus a one-tap path "
      "to report it to a real authority. My recommendation is to keep collecting verified local scam "
      "reports, which both improves the model and builds a data advantage no competitor can copy by adding "
      "a language pack.”")

para("Tip: record Option B once cleanly; if the total runs over 5:00, fall back to Option A. Keep total "
     "length at or under 5:00 — if long, trim the Education and Wilson shots to ~15s each.",
     italic=True, color=GREY)

out = os.path.join(os.path.dirname(__file__), "DEMO_VIDEO_SCRIPT.docx")
doc.save(out)
print("Saved:", out)
