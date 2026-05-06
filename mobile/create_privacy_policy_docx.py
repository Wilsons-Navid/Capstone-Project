from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a new Document
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

# Title
title = doc.add_heading('Privacy Policy for Rethicssec', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Effective Date
effective = doc.add_paragraph()
effective.add_run('Effective Date: ').bold = True
effective.add_run('January 19, 2025')
effective.alignment = WD_ALIGN_PARAGRAPH.CENTER

last_updated = doc.add_paragraph()
last_updated.add_run('Last Updated: ').bold = True
last_updated.add_run('January 19, 2025')
last_updated.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('_' * 80)

# Section 1
doc.add_heading('1. Introduction', 1)
doc.add_paragraph(
    'Welcome to Rethicssec (also known as "Rethicsai," "the App," "we," "us," or "our"). '
    'Rethicssec is a cybercrime reporting and cybersecurity education platform designed specifically for Africa. '
    'We are committed to protecting your privacy and ensuring transparency about how we collect, use, and safeguard your personal information.'
)
doc.add_paragraph('This Privacy Policy explains:')
bullets = ['What information we collect', 'How we use your information',
           'Who we share your information with', 'Your rights and choices regarding your data',
           'How we protect your information']
for bullet in bullets:
    doc.add_paragraph(bullet, style='List Bullet')

doc.add_paragraph(
    'By using Rethicssec, you agree to the collection and use of information in accordance with this Privacy Policy. '
    'If you do not agree with this policy, please do not use our services.'
)

# Section 2
doc.add_heading('2. Information We Collect', 1)
doc.add_heading('2.1 Information You Provide Directly', 2)

doc.add_paragraph().add_run('Account Registration:').bold = True
doc.add_paragraph('When you create an account, we collect:')
account_items = ['Full name (first and last name)', 'Email address', 'Password (encrypted)',
                 'Phone number (optional)', 'Country/location (optional)', 'Language preference',
                 'Profile photo/avatar (optional)']
for item in account_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph().add_run('Incident Reporting:').bold = True
doc.add_paragraph('When you report a cybercrime incident, you may provide:')
incident_items = ['Incident type and description', 'Date and time of incident',
                  'Location where incident occurred', 'Financial loss amount (if applicable)',
                  'Suspect information', 'Evidence files (photos, documents, videos)',
                  'Contact preferences for follow-up']
for item in incident_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph().add_run('AI Assistant Interactions:').bold = True
doc.add_paragraph('When you use Wilson AI Assistant, we collect:')
ai_items = ['Chat messages and conversation history', 'File attachments shared during conversations',
            'Session information for context continuity']
for item in ai_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph().add_run('Education & Learning:').bold = True
doc.add_paragraph('We track:')
education_items = ['Course completion status', 'Learning module progress', 'Achievement certificates',
                   'Time spent on educational content']
for item in education_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph().add_run('Preferences & Settings:').bold = True
doc.add_paragraph('We store your preferences for:')
pref_items = ['Push notifications', 'Email notifications', 'SMS notifications', 'Marketing communications',
              'Security alerts', 'Language and localization settings']
for item in pref_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('2.2 Information Collected Automatically', 2)

doc.add_paragraph().add_run('Device Information:').bold = True
device_items = ['Device brand and model', 'Operating system and version (Android/iOS)', 'App version',
                'Device category (determined by specifications)', 'Unique device identifiers (Firebase Cloud Messaging tokens)']
for item in device_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph().add_run('Location Data:').bold = True
location_items = ['GPS coordinates (only when you grant location permission)',
                  'Approximate location based on IP address', 'Location data for incident reporting']
for item in location_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph().add_run('Usage Data:').bold = True
usage_items = ['App screens and features you access', 'Time spent in the app',
               'Interaction with educational content', 'AI assistant usage patterns',
               'Incident report creation and updates', 'Authentication events']
for item in usage_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph().add_run('Technical Data:').bold = True
tech_items = ['Network connectivity status (WiFi, mobile data, offline)', 'App performance metrics',
              'Error logs and crash reports', 'API response times']
for item in tech_items:
    doc.add_paragraph(item, style='List Bullet')

# Section 3
doc.add_heading('3. How We Use Your Information', 1)
doc.add_paragraph('We use the collected information for the following purposes:')

doc.add_heading('3.1 Service Provision', 2)
service_items = ['Create and manage your account', 'Process and respond to incident reports',
                 'Provide AI-powered cybersecurity guidance through Wilson AI',
                 'Deliver educational content and track your learning progress',
                 'Enable location-based features for incident reporting',
                 'Send notifications about case updates and security alerts']
for item in service_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('3.2 Communication', 2)
comm_items = ['Respond to your inquiries and support requests',
              'Send verification emails and password reset instructions',
              'Notify you about incident status changes',
              'Share security alerts and threat intelligence',
              'Send educational content and safety tips (with your consent)',
              'Deliver marketing communications (with your opt-in consent)']
for item in comm_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('3.3 Improvement & Analytics', 2)
analytics_items = ['Analyze app usage to improve features and user experience',
                   'Monitor app performance and fix technical issues',
                   'Develop new features and services',
                   'Conduct research on cybercrime trends in Africa',
                   'Optimize AI assistant responses and accuracy']
for item in analytics_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('3.4 Security & Legal Compliance', 2)
security_items = ['Detect and prevent fraud, abuse, and security incidents',
                  'Enforce our Terms of Service',
                  'Comply with legal obligations and law enforcement requests',
                  'Protect the rights, property, and safety of our users and the public']
for item in security_items:
    doc.add_paragraph(item, style='List Bullet')

# Section 4
doc.add_heading('4. Legal Basis for Processing (GDPR)', 1)
doc.add_paragraph(
    'If you are located in the European Economic Area (EEA), we process your personal data under the following legal bases:'
)
legal_basis = [
    'Consent: When you opt-in to marketing communications, analytics, or optional features',
    'Contract Performance: To provide services you requested (account creation, incident reporting)',
    'Legitimate Interests: To improve our services, ensure security, and analyze usage patterns',
    'Legal Obligation: To comply with applicable laws and regulations'
]
for item in legal_basis:
    doc.add_paragraph(item, style='List Bullet')

# Section 5
doc.add_heading('5. How We Share Your Information', 1)
doc.add_paragraph(
    'We do not sell, rent, or trade your personal information to third parties. '
    'We share your information only in the following circumstances:'
)

doc.add_heading('5.1 Third-Party Service Providers', 2)

doc.add_paragraph().add_run('Firebase/Google Cloud Platform:').bold = True
doc.add_paragraph('• Purpose: Backend infrastructure, database, file storage, authentication, analytics')
doc.add_paragraph('• Data Shared: User ID, email, name, profile data, incident reports, chat history, device information')
doc.add_paragraph('• Privacy Policy: https://policies.google.com/privacy')

doc.add_paragraph().add_run('Google Vertex AI (Gemini):').bold = True
doc.add_paragraph('• Purpose: AI-powered cybersecurity assistance through Wilson AI')
doc.add_paragraph('• Data Shared: Chat messages, context for intelligent responses')
doc.add_paragraph('• Privacy Policy: https://policies.google.com/privacy')

doc.add_paragraph().add_run('Google Sign-In:').bold = True
doc.add_paragraph('• Purpose: Third-party authentication')
doc.add_paragraph('• Data Shared: Email, name, profile photo')
doc.add_paragraph('• Privacy Policy: https://policies.google.com/privacy')

doc.add_paragraph().add_run('Apple Sign-In:').bold = True
doc.add_paragraph('• Purpose: Third-party authentication for iOS/macOS users')
doc.add_paragraph('• Data Shared: Email (optional), name (optional)')
doc.add_paragraph('• Privacy Policy: https://www.apple.com/privacy/')

doc.add_heading('5.2 Law Enforcement & Legal Requirements', 2)
doc.add_paragraph('We may disclose your information if required by law or in response to:')
law_items = ['Valid legal processes (subpoenas, court orders, warrants)',
             'Law enforcement requests for cybercrime investigations',
             'Protection of our rights, property, or safety',
             'Prevention of fraud, security threats, or illegal activities']
for item in law_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('5.3 Business Transfers', 2)
doc.add_paragraph(
    'In the event of a merger, acquisition, or sale of assets, your information may be transferred to the acquiring entity. '
    'We will notify you via email and/or prominent notice in the app before your data is transferred.'
)

doc.add_heading('5.4 With Your Consent', 2)
doc.add_paragraph('We may share your information with other parties when you explicitly consent or request such sharing.')

# Section 6
doc.add_heading('6. Data Storage & International Transfers', 1)

doc.add_heading('6.1 Storage Location', 2)
storage_items = [
    'Cloud Storage: Your data is stored on Google Cloud Platform servers (Firebase Cloud Firestore and Firebase Storage)',
    'Data Region: Primary storage is in Google Cloud\'s default regions (typically United States)',
    'Local Storage: Some data is cached locally on your device for offline access using encrypted storage'
]
for item in storage_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('6.2 International Data Transfers', 2)
doc.add_paragraph('If you are located outside the United States, please note that:')
transfer_items = [
    'Your information may be transferred to and processed in the United States',
    'The U.S. may have different data protection laws than your country',
    'We rely on standard contractual clauses and Google Cloud\'s compliance frameworks',
    'By using our services, you consent to the transfer of your information to the U.S.'
]
for item in transfer_items:
    doc.add_paragraph(item, style='List Bullet')

# Section 7
doc.add_heading('7. Data Retention', 1)
doc.add_paragraph('We retain your personal information for as long as necessary to fulfill the purposes outlined in this Privacy Policy:')
retention_items = [
    'Account Data: Retained while your account is active',
    'Incident Reports: Retained for investigation, legal, and archival purposes (typically 7+ years)',
    'Chat History: Retained indefinitely unless you delete conversations',
    'Analytics Data: Retained according to Firebase defaults (60+ days for detailed data, aggregated data retained longer)',
    'Deleted Accounts: Personal data is deleted from our systems within 30 days of account deletion, except where retention is required by law'
]
for item in retention_items:
    doc.add_paragraph(item, style='List Bullet')

# Section 8
doc.add_heading('8. Your Privacy Rights', 1)
doc.add_paragraph('Depending on your location, you may have the following rights:')

rights_sections = {
    '8.1 Access & Portability': [
        'Request a copy of your personal data',
        'Export your incident reports and chat history'
    ],
    '8.2 Correction': [
        'Update your profile information through account settings',
        'Correct inaccuracies in your data'
    ],
    '8.3 Deletion': [
        'Delete your account and associated personal data through the app settings',
        'Request deletion of specific incident reports or chat conversations'
    ],
    '8.4 Opt-Out & Restriction': [
        'Disable push notifications, email notifications, or SMS notifications',
        'Opt-out of marketing communications',
        'Revoke permissions (camera, location, microphone) through device settings',
        'Restrict certain data processing activities'
    ],
    '8.5 Withdrawal of Consent': [
        'Withdraw consent for analytics or optional features at any time',
        'Note: Withdrawal does not affect the lawfulness of processing before withdrawal'
    ],
    '8.6 Data Portability (GDPR/CCPA)': [
        'Receive your data in a structured, commonly used, machine-readable format'
    ],
    '8.7 Lodge a Complaint': [
        'Contact us directly with privacy concerns',
        'File a complaint with your local data protection authority (EEA residents)'
    ]
}

for section, items in rights_sections.items():
    doc.add_heading(section, 2)
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph().add_run('To exercise your rights, please contact us at: [Your Contact Email]').bold = True

# Section 9
doc.add_heading('9. Device Permissions', 1)
doc.add_paragraph('Rethicssec requests the following permissions to provide full functionality:')

# Create a simple table representation
permissions_data = [
    ('Camera', 'Capture photos as evidence for incident reports', 'Optional'),
    ('Photo Library', 'Attach existing photos/videos to reports', 'Optional'),
    ('Location', 'Tag incident location with GPS coordinates', 'Optional'),
    ('Microphone', 'Voice input for AI assistant', 'Optional'),
    ('Storage', 'Save files and cache data for offline access', 'Required'),
    ('Notifications', 'Send push notifications for case updates and alerts', 'Optional'),
    ('Internet', 'Communicate with servers and cloud services', 'Required')
]

for perm, purpose, required in permissions_data:
    p = doc.add_paragraph()
    p.add_run(f'{perm}: ').bold = True
    p.add_run(f'{purpose} ({required})')

doc.add_paragraph(
    'You can manage these permissions at any time through your device settings. '
    'Denying optional permissions may limit certain features but will not prevent basic app usage.'
)

# Section 10
doc.add_heading('10. Security Measures', 1)
doc.add_paragraph('We implement industry-standard security measures to protect your data:')

doc.add_heading('10.1 Technical Safeguards', 2)
tech_safeguards = [
    'Encryption in Transit: All data transmitted between your device and our servers uses SSL/TLS encryption',
    'Encryption at Rest: Sensitive data stored locally on your device is encrypted',
    'Secure Authentication: Passwords are hashed and encrypted; we support multi-factor authentication',
    'Access Controls: Role-based access controls limit who can access your data',
    'Firestore Security Rules: Database rules prevent unauthorized access to user data',
    'No Cleartext Traffic: App explicitly disables unencrypted network communication'
]
for item in tech_safeguards:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('10.2 Organizational Safeguards', 2)
org_safeguards = [
    'Regular security audits and vulnerability assessments',
    'Employee training on data protection and privacy',
    'Incident response procedures for data breaches',
    'Limited access to personal data on a need-to-know basis'
]
for item in org_safeguards:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('10.3 Limitations', 2)
doc.add_paragraph(
    'While we strive to protect your data, no security system is 100% secure. '
    'We cannot guarantee absolute security and are not responsible for unauthorized access '
    'resulting from circumstances beyond our control.'
)

# Section 11
doc.add_heading('11. Children\'s Privacy', 1)
doc.add_paragraph('Rethicssec is not intended for children under the age of 13. We do not knowingly collect personal information from children under 13.')
children_items = [
    'If you are under 13, do not use this app or provide any personal information',
    'If we discover that we have collected data from a child under 13, we will delete it immediately',
    'Parents or guardians who believe their child has provided information should contact us immediately'
]
for item in children_items:
    doc.add_paragraph(item, style='List Bullet')

# Section 12
doc.add_heading('12. Third-Party Links', 1)
doc.add_paragraph(
    'Rethicssec may contain links to third-party websites, services, or resources for educational purposes or external support. '
    'We are not responsible for the privacy practices of these third parties. '
    'We encourage you to review their privacy policies before providing any personal information.'
)

# Section 13
doc.add_heading('13. Analytics & Tracking Technologies', 1)
doc.add_paragraph('We use Firebase Analytics to understand how users interact with our app:')
analytics_points = [
    'What We Track: Screen views, feature usage, incident creation events, AI interactions, educational content engagement',
    'Purpose: Improve user experience, identify bugs, optimize features',
    'User Control: Analytics is enabled by default but can be disabled in app settings',
    'No Personal Identifiers: Analytics data is aggregated and does not include directly identifiable information'
]
for item in analytics_points:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('We do not use third-party advertising networks or tracking cookies for marketing purposes.')

# Section 14
doc.add_heading('14. Your Choices & Control', 1)
doc.add_paragraph('You have control over your data and privacy:')

control_data = [
    ('Notification Preferences', 'App Settings → Notifications'),
    ('Location Sharing', 'Device Settings → App Permissions → Location'),
    ('Camera/Microphone Access', 'Device Settings → App Permissions'),
    ('Marketing Emails', 'Unsubscribe link in emails or Account Settings'),
    ('Account Deletion', 'App Settings → Account → Delete Account'),
    ('Data Export', 'Contact us to request a copy of your data'),
    ('Analytics Opt-Out', 'App Settings → Privacy → Disable Analytics')
]

for control, method in control_data:
    p = doc.add_paragraph()
    p.add_run(f'{control}: ').bold = True
    p.add_run(method)

# Section 15
doc.add_heading('15. California Privacy Rights (CCPA)', 1)
doc.add_paragraph('If you are a California resident, you have the following rights under the California Consumer Privacy Act (CCPA):')
ccpa_rights = [
    'Right to Know: Request disclosure of personal information collected, sources, purposes, and third parties shared with',
    'Right to Delete: Request deletion of your personal information (subject to certain exceptions)',
    'Right to Opt-Out: We do not sell personal information, so opt-out is not applicable',
    'Right to Non-Discrimination: We will not discriminate against you for exercising your CCPA rights'
]
for item in ccpa_rights:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('To exercise these rights, contact us at [Your Contact Email].')

# Section 16
doc.add_heading('16. European Privacy Rights (GDPR)', 1)
doc.add_paragraph('If you are located in the European Economic Area (EEA), you have the following rights under the General Data Protection Regulation (GDPR):')
gdpr_rights = [
    'Right to access your personal data',
    'Right to rectification of inaccurate data',
    'Right to erasure ("right to be forgotten")',
    'Right to restrict processing',
    'Right to data portability',
    'Right to object to processing',
    'Right to withdraw consent',
    'Right to lodge a complaint with a supervisory authority'
]
for item in gdpr_rights:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('To exercise these rights, contact us at [Your Contact Email].')

# Section 17
doc.add_heading('17. South African Privacy Rights (POPIA)', 1)
doc.add_paragraph('If you are located in South Africa, we comply with the Protection of Personal Information Act (POPIA):')
popia_rights = [
    'You have the right to access, correct, and delete your personal information',
    'You can object to the processing of your information',
    'You can request that we stop sending marketing communications',
    'You can lodge a complaint with the Information Regulator'
]
for item in popia_rights:
    doc.add_paragraph(item, style='List Bullet')

# Section 18
doc.add_heading('18. Data Breach Notification', 1)
doc.add_paragraph('In the unlikely event of a data breach that affects your personal information:')
breach_items = [
    'We will notify affected users within 72 hours (as required by GDPR)',
    'Notification will be sent via email and/or in-app notification',
    'We will provide details about the breach, data affected, and steps you should take',
    'We will report the breach to relevant authorities as required by law'
]
for item in breach_items:
    doc.add_paragraph(item, style='List Bullet')

# Section 19
doc.add_heading('19. Changes to This Privacy Policy', 1)
doc.add_paragraph('We may update this Privacy Policy from time to time to reflect:')
changes_reasons = [
    'Changes in our practices or services',
    'Legal, regulatory, or security requirements',
    'User feedback and industry best practices'
]
for item in changes_reasons:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph().add_run('When we make changes:').bold = True
changes_items = [
    'We will update the "Last Updated" date at the top of this policy',
    'Material changes will be communicated via email and/or prominent in-app notice',
    'Continued use of the app after changes constitutes acceptance of the updated policy',
    'You can review the current policy at any time in the app\'s Help section'
]
for item in changes_items:
    doc.add_paragraph(item, style='List Bullet')

# Section 20
doc.add_heading('20. Contact Us', 1)
doc.add_paragraph('If you have questions, concerns, or requests regarding this Privacy Policy or our privacy practices, please contact us:')
doc.add_paragraph().add_run('Email: [Your Contact Email Address]').bold = True
doc.add_paragraph().add_run('Address: [Your Business Address]').bold = True
doc.add_paragraph().add_run('Response Time: We will respond to privacy inquiries within 30 days').bold = True

# Section 21
doc.add_heading('21. Data Protection Officer', 1)
doc.add_paragraph('For users in the European Economic Area, you may contact our Data Protection Officer (if applicable):')
doc.add_paragraph().add_run('DPO Email: [DPO Email Address]').bold = True
doc.add_paragraph().add_run('DPO Address: [DPO Business Address]').bold = True

# Section 22
doc.add_heading('22. Consent', 1)
doc.add_paragraph(
    'By using Rethicssec, you acknowledge that you have read, understood, and agree to this Privacy Policy. '
    'You consent to the collection, use, and sharing of your information as described herein.'
)
doc.add_paragraph(
    'If you do not agree with this Privacy Policy, please discontinue use of the app and contact us to delete your account.'
)

# Footer
doc.add_paragraph('_' * 80)
footer = doc.add_paragraph()
footer.add_run('Thank you for trusting Rethicssec with your cybersecurity needs. Your privacy and security are our top priorities.').italic = True
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
copyright = doc.add_paragraph('© 2025 Rethicssec. All rights reserved.')
copyright.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Save the document
doc.save(r'C:\Users\LENOVO\Desktop\Rethicsai\PRIVACY_POLICY.docx')
print("Privacy Policy document created successfully: PRIVACY_POLICY.docx")
