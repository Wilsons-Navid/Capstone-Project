from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

doc = Document()

# --- Styles ---
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# Title page
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('RETHICS AI')
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x2D, 0x1B, 0x14)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Google Stitch Design Prompts')
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0xCC, 0x88, 0x00)

desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = desc.add_run('Futuristic UI Designs with Rich Backgrounds, Animations & User-Friendly Experience')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('Version 1.0 | March 2026')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# --- Table of Contents ---
toc_title = doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Dashboard — Home Base',
    '2. AI Assistant Wilson — Friendly Guardian',
    '3. Threat Scanner — Safety Radar',
    '4. Incident Reporting — Safe Space',
    '5. Case Tracking — Your Journey',
    '6. Education Hub — Learning Garden',
    '7. Emergency Contacts — Lifeline',
    '8. Onboarding — Story of Protection',
    '9. Admin Dashboard — Mission Control',
    '10. User Management — People Hub',
    '11. Case Triage & Assignment — Dispatch Center',
    '12. Threat Database Manager — Vault',
    '13. Analytics & Reports — Intelligence Brief',
    '14. Education Content Manager — Knowledge Forge',
    '15. Role Management & System Config — Control Center',
    '16. Admin Broadcast Center — Signal Tower',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(2)

doc.add_page_break()

# --- Helper ---
def add_section_header(doc, number, title, subtitle):
    doc.add_paragraph()
    h = doc.add_heading(f'{number}. {title}', level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x2D, 0x1B, 0x14)
    sub = doc.add_paragraph()
    run = sub.add_run(subtitle)
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0xCC, 0x88, 0x00)

def add_prompt(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(6)
    # Add "PROMPT:" label
    label = p.add_run('PROMPT:\n')
    label.bold = True
    label.font.size = Pt(11)
    label.font.color.rgb = RGBColor(0xCC, 0x88, 0x00)
    # Add prompt text
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# ============================================
# SECTION 1: Dashboard
# ============================================
add_section_header(doc, '1', 'Dashboard — Home Base', 'Main user dashboard with security score, quick stats, and feature grid')

add_prompt(doc, """Design a warm, futuristic mobile dashboard for "Rethics AI", an African cybersecurity companion app.

BACKGROUND: A full-bleed gradient from deep midnight blue #0A1628 at top to warm charcoal #1A1210 at bottom. Behind the content, a softly glowing illustrated African savanna silhouette skyline (baobab trees, distant mountains) in amber #CC8800 at 6% opacity. Floating soft-focus bokeh particles in gold and teal drift slowly upward like fireflies. A subtle animated aurora ribbon in amber-to-teal weaves across the top 20% of the screen behind the header.

HEADER: Rounded avatar photo with a warm amber ring that breathes (scales 1.0 to 1.03 in 3s loop). "Good evening, Amara" in friendly rounded font. Notification bell with a soft bounce animation when new alerts arrive — count badge scales up with spring physics.

SECURITY SCORE: Large circular ring in the center area, 180px wide. The ring draws itself clockwise with a smooth 1.5s ease-out animation on page load, filling with a gradient from teal #00D4AA to amber #FFAA00 based on score. Inside: large friendly "87" number that counts up from 0. Below the ring: "Your digital safety is strong" in warm subtitle text with a gentle fade-in after the ring completes. Small sparkle particles emit from the ring endpoint.

QUICK STATS: 3 rounded glass-morphism cards in a row with 16px blur backdrop, warm inner glow, and 1px border at 15% white opacity. Each card fades in sequentially (0.2s stagger):
  • Shield icon + "3 Threats Blocked" in teal
  • Folder icon + "2 Active Cases" in amber
  • Book icon + "68% Learned" in soft green
Numbers use counting animation on load. Icons have a gentle float (translateY 2px oscillation, 4s loop).

FEATURE GRID: 6 large rounded cards (2x3 grid, 16px gap) with category background illustrations:
  • "Report Incident" — faint police badge watermark, warm red-amber gradient overlay
  • "Track Cases" — faint timeline watermark, teal gradient overlay
  • "Ask Wilson AI" — faint brain circuit watermark, blue-amber gradient overlay
  • "Scan Threats" — faint radar watermark, green-teal gradient overlay
  • "Learn Safety" — faint book watermark, warm gold gradient overlay
  • "Emergency SOS" — faint cross/phone watermark, crimson gradient overlay
Each card has a friendly rounded icon (32px), bold label, and 1-line subtitle. On tap: card scales to 0.96 with spring-back and ripple effect. Cards stagger fade-in-up on page load (0.1s intervals).

RECENT ACTIVITY: Horizontal scrollable cards showing latest actions with friendly language: "Your phishing report was received", "Wilson answered your question", "New safety tip available". Each card has a left color accent stripe matching action type. Cards slide in from the right on load.

BOTTOM NAV: 5 items with rounded pill indicator that slides with spring animation between tabs. Active icon scales up 1.2x and shifts color to amber. Subtle haptic dot animation on active icon.

Typography: Rounded, friendly sans-serif (like Nunito or Poppins). Warm whites #F5F0EB for primary text, muted #A09888 for secondary.
Mobile portrait, dark mode, feeling safe and protected, not intimidating.""")

doc.add_page_break()

# ============================================
# SECTION 2: AI Assistant Wilson
# ============================================
add_section_header(doc, '2', 'AI Assistant Wilson — Friendly Guardian', 'AI chat interface with sentient orb avatar and contextual threat detection')

add_prompt(doc, """Design a warm, futuristic AI chat screen for "Wilson", the friendly AI assistant in "Rethics AI" cybersecurity app.

BACKGROUND: Deep warm dark #100E0C base. A large subtle illustrated scene behind the chat area: soft watercolor-style African landscape with acacia trees and a setting sun in amber/gold tones at 5% opacity, giving warmth without distraction. Gentle animated particles — tiny floating dots of amber and teal light drifting slowly upward like embers from a campfire. A very subtle radial gradient glow emanates from Wilson's avatar area in teal at 4% opacity.

WILSON AVATAR AREA: Top section with Wilson's avatar — a friendly rounded shield character with eyes (think Baymax meets a shield). The avatar sits inside a soft glowing teal circle with two orbiting ring animations (thin, slow rotation, 20s and 30s loops in opposite directions). Below: "Wilson" in bold friendly font, "Your AI Security Buddy" subtitle. Status: "Online and ready to help" with a green breathing dot (opacity 0.6 to 1.0, 2s loop). When Wilson is typing, the avatar's eyes become animated dots bouncing in sequence, and the rings speed up slightly.

CHAT AREA: Messages appear with smooth slide-up + fade-in animations (0.3s ease-out):
  • Wilson's messages: left-aligned, warm dark glass cards (#1A1714) with rounded corners (16px), soft teal left border (3px) that glows subtly. Text in warm white. When Wilson shares a safety tip, a small animated lightbulb icon pulses beside it. Threat warnings get an amber left border with a gentle pulse animation.
  • User messages: right-aligned, amber #CC8800 background with rounded corners, dark text. Subtle send animation — message scales from 0.8 to 1.0 with spring physics.
  • Timestamps appear as small friendly relative text "just now", "2 min ago" in muted color between message groups.

SMART SUGGESTIONS: Below Wilson's responses, 2-3 suggestion chips float in with staggered pop animation: "Is this link safe?", "How do I spot a scam?", "What's SIM swapping?" — each as a rounded pill with thin amber border and subtle glow on tap. Chips gently bounce (translateY 1px, 3s loop, staggered).

THREAT DETECTION MOMENT: When Wilson detects something dangerous in the conversation, a special animation triggers — the background briefly flashes with a soft red radial pulse from the center, Wilson's avatar border shifts from teal to amber, and the response card has a red-amber gradient left border with shield icon. An animated "Warning: Threat Detected" banner slides down from the top with a smooth 0.4s animation.

TYPING INDICATOR: Three dots in a teal bubble with wave animation (each dot bounces sequentially). Background of the bubble has a subtle shimmer sweep.

INPUT BAR: Bottom-pinned dark glass bar with rounded input field. Microphone icon on the left with a voice waveform animation that appears when held (bars dance to audio amplitude). Send button on right — arrow icon that rotates 360 degrees and shoots upward with a trail when tapped. Placeholder text cycles through helpful prompts: "Ask me anything about online safety..." then "Paste a suspicious link to check..." then "Tell me what happened..." with fade transitions.

CONTEXT ACTIONS: When Wilson's response includes actionable advice, animated action buttons appear below the message: "Scan This Link" (radar icon), "Report This" (flag icon), "Learn More" (book icon) — each as glass pill buttons with category-colored left icon, slide-in animation from left with 0.1s stagger.

Typography: Rounded, friendly (Nunito/Poppins). Wilson's text slightly larger than user text for readability. Warm color palette throughout.
Mobile portrait, dark mode, feeling like chatting with a knowledgeable friend, not a cold AI.""")

doc.add_page_break()

# ============================================
# SECTION 3: Threat Scanner
# ============================================
add_section_header(doc, '3', 'Threat Scanner — Safety Radar', 'Multi-type threat scanning with radar visualization and dual-source verification')

add_prompt(doc, """Design a warm, futuristic threat scanner screen for "Rethics AI" cybersecurity app. It should feel empowering, not scary.

BACKGROUND: Deep warm black #0D0A08 base. A beautiful illustrated African night sky scene at 7% opacity — stars, a crescent moon, and distant mountain silhouettes in warm tones. Animated: stars twinkle with randomized opacity pulses (0.3-0.8, varied timing 2-6s). A soft animated nebula glow in teal and amber shifts slowly in the upper area (like a slow-moving gradient, 30s loop). Faint concentric circles radiate outward from the center of the screen at 3% opacity — like calm water ripples, animating slowly (one new ring every 4s, fading as they expand).

TAB SELECTOR: 4 tabs as rounded pill buttons in a horizontal row with a sliding background indicator in amber that moves with spring physics: URL (link icon), Email (envelope icon), Phone (phone icon), File (document icon). Inactive tabs in muted warm gray, active tab text in dark on amber background. Icons have a subtle scale animation on tab switch.

INPUT AREA: Large friendly rounded input card with glass-morphism effect. Warm placeholder text: "Paste a link to check if it's safe..." with a blinking amber cursor. Paste button (clipboard icon) on the left with a pop animation on tap. Clear button (x) appears with fade-in when text is entered. Below: large rounded "SCAN NOW" button in amber #CC8800, full-width, with a friendly shield-check icon. On press: button compresses slightly, text changes to "Scanning..." and a progress sweep animation fills the button from left to right in a lighter amber.

SCANNING STATE: When scanning, the background ripple animation speeds up and intensifies to 8% opacity. The central area shows a friendly animated shield that rotates slowly with a scanning line sweeping across it (like a lighthouse beam, teal glow). Below: "Checking against 14,000+ known threats..." with an animated ellipsis. Three animated checkpoints appear sequentially with check-mark pop animations:
  Checkmark "Rethics Database checked" (0.5s)
  Checkmark "VirusTotal verified" (1.5s)
  Checkmark "AI analysis complete" (2.5s)

RESULT — SAFE: Shield icon morphs from scanning to a large green checkmark with a celebration burst animation (small confetti particles in teal and gold). Glass result card slides up with spring physics:
  • "All Clear! This looks safe" in friendly large text, green tint
  • Safety score: "98/100" with animated circular fill in green
  • Detail rows fade in with stagger: "No threats found in Rethics database", "VirusTotal: 0/89 engines flagged", "AI Assessment: Legitimate"
  • Friendly tip: "You're doing great staying vigilant!" with lightbulb icon

RESULT — THREAT FOUND: Shield icon morphs to a warning triangle with a controlled amber/red pulse (not aggressive, just attention-getting). Soft red radial gradient appears behind the result area. Glass result card slides up:
  • "Warning: This may be dangerous" in clear large text, amber/red tint
  • Threat level gauge: animated fill bar from left (green to red gradient), pointer lands on detected level with a soft bounce
  • Threat type: "Phishing Attempt" with description in friendly plain language
  • What to do: numbered steps in clear friendly text with action buttons — "Block This" (shield icon), "Report It" (flag icon), "Tell a Friend" (share icon) — each as rounded pill buttons with relevant colors

SCAN HISTORY: Below results, "Recent Scans" section with horizontal scrollable cards showing previous scans — each with URL/email preview, date, and small colored dot (green/amber/red). Cards slide in from right on load.

Typography: Warm, friendly rounded font. Clear hierarchy. Encouraging language throughout.
Mobile portrait, dark mode, empowering and protective feeling.""")

doc.add_page_break()

# ============================================
# SECTION 4: Incident Reporting
# ============================================
add_section_header(doc, '4', 'Incident Reporting — Safe Space', 'Multi-step incident report form with evidence upload and encryption')

add_prompt(doc, """Design a warm, futuristic incident reporting screen for "Rethics AI" cybersecurity app. This should feel like a safe, guided process — not a cold bureaucratic form.

BACKGROUND: Deep warm dark #100E0C. A beautiful soft-focus illustrated background of African shield patterns and woven textile geometry at 4% opacity in warm amber/brown tones — representing protection and community. Gentle floating ember particles in gold drift upward at the edges. A soft warm gradient glow sits behind the form area — teal at the top fading to amber at bottom, very subtle.

PROGRESS STEPPER: Top section with a friendly 4-step horizontal progress bar. Steps: "What happened" then "Tell us more" then "Evidence" then "Review". Connected by a line that fills with animated amber light as you progress. Current step: larger circle with pulse glow in amber, icon inside (magnifier, pencil, camera, checkmark). Completed steps: solid teal circles with animated checkmark pop. Future steps: outlined circles in muted tone. Step labels in friendly text below each circle. The fill animation between steps takes 0.6s with ease-in-out.

STEP 1 — WHAT HAPPENED:
Warm headline: "What happened to you?" with supporting text: "Take your time. Your report is confidential and encrypted."
Incident type grid: 3-column grid of rounded cards with illustrated icons (not generic — hand-drawn style):
  • Phishing (fishing hook icon on teal)
  • Online Fraud (mask icon on amber)
  • Mobile Money Scam (phone-money icon on gold)
  • SIM Swapping (SIM card icon on red)
  • Identity Theft (fingerprint icon on purple)
  • Cyberbullying (chat-sad icon on blue)
  • Romance Scam (heart-broken icon on pink)
  • Investment Fraud (chart-down icon on orange)
  • Other (dots icon on gray)
Each card has a soft background gradient in its category color at 20%. On selection: card border glows in its color with spring scale animation (1.0 to 1.05 to 1.02), checkmark appears in corner with pop animation. Unselected cards dim slightly. Cards stagger fade-in on load.

STEP 2 — DETAILS:
Form fields with warm glass-morphism styling. Each field is a rounded dark glass card:
  • Labels float above in amber when focused, with smooth 0.3s translate-up animation
  • Input text in warm white, cursor amber
  • "Give it a title" — single line with example placeholder that types itself: "Someone stole my..."
  • "Tell us what happened" — multiline with expanding animation as user types, min 4 lines
  • "When did this happen?" — date picker that opens as a friendly bottom sheet with calendar, smooth slide-up
  • "Where did this happen?" — location field with map pin icon, tapping opens mini map with animated pin drop and ripple rings
  • "Any financial loss?" — optional toggle that slides open a currency input with country-specific currency symbol
Focus states: active field gets a soft amber glow border (0.3s fade-in), other fields dim slightly. Validated fields show a small teal checkmark that pops in at the right edge.

STEP 3 — EVIDENCE:
Warm headline: "Add any proof you have" with subtitle: "Screenshots, emails, documents — anything helps"
Upload area: large rounded dashed-border zone with animated border — dashes flow clockwise slowly like a conveyor belt in amber at 30%. Center: camera/upload icon with a gentle bounce animation (translateY 2px, 3s loop). Text: "Tap to upload or take a photo".
On file added: upload area shrinks, file appears as a thumbnail card below with:
  • Image preview or file type icon with rounded corners
  • File name and size
  • Upload progress ring that fills circularly in teal with smooth animation
  • Completion: ring bursts into small particles and transforms into a teal checkmark
  • Remove button (x) in corner with hover glow
Multiple files arrange in a scrollable grid. Add more button as a smaller dashed card with "+" icon.

STEP 4 — REVIEW:
All entered information displayed in a clean summary with section dividers. Each section slides in from the bottom with 0.15s stagger. An animated encryption badge at the top: lock icon with orbiting dots, text "Your report is end-to-end encrypted". Priority selector: 3 large pill buttons with friendly labels:
  • "I can wait" (Low — teal, clock icon)
  • "Please look soon" (Medium — amber, flag icon)
  • "This is urgent" (High — red, alert icon, subtle pulse animation when selected)
Submit button: large rounded amber button "Submit My Report". On press: button shrinks, morphs into a circular progress ring that fills, then bursts into a checkmark with celebration particles. Screen transitions to a confirmation view: "Your case number is RTH-2024-0847" in large friendly text with copy button, "We'll notify you of any updates" with bell icon animation.

NAVIGATION: Friendly "Back" and "Next" buttons at the bottom with arrow icons. Next button disabled (dimmed) until required fields complete, then activates with a glow fade-in. Next button has a forward arrow that slides right on tap.

Typography: Warm rounded font, generous line height. Reassuring, plain language. No legal jargon.
Mobile portrait, dark mode, safe and supportive feeling — like talking to a trusted friend.""")

doc.add_page_break()

# ============================================
# SECTION 5: Case Tracking
# ============================================
add_section_header(doc, '5', 'Case Tracking — Your Journey', 'Visual journey timeline showing case progress from filing to resolution')

add_prompt(doc, """Design a warm, futuristic case tracking screen for "Rethics AI" cybersecurity app. It should feel like following a personal journey toward resolution.

BACKGROUND: Deep warm dark #0F0C0A. A beautiful illustrated African landscape at 5% opacity — rolling hills with a path/road leading toward a sunrise on the horizon, symbolizing resolution and hope. The sunrise has a very subtle animated glow pulse (warm amber, 8s loop). Gentle particles of warm light float along the path direction. Faint topographic map lines at 3% opacity give depth.

CASE HEADER: Glass-morphism card at top with warm inner glow:
  • Case number "RTH-2024-0847" in monospace with copy icon (tap: icon morphs to checkmark with pop animation)
  • Incident type badge: "Phishing Attack" with category icon and color accent
  • Filed date with friendly relative time: "Filed 3 days ago, March 19"
  • Priority badge: colored pill (teal/amber/red) with gentle pulse if high priority
  Card fades in and slides down from top on page load (0.4s).

JOURNEY TIMELINE: The centerpiece — a vertical flowing path visualization (not a rigid line):
  • A curved animated path line flows from top to bottom, like a river or road on a map
  • 6 milestone nodes along the path: REPORTED then RECEIVED then ASSIGNED then INVESTIGATING then RESOLVED then CLOSED
  • Completed nodes: solid teal circles with animated checkmark, connected by solid glowing teal path
  • Current node: larger amber circle with pulse glow animation (scale 1.0-1.15, 2s loop) and ripple rings expanding outward
  • Future nodes: soft outlined circles with dotted path, muted colors
  • Each node has a timestamp and friendly status text on alternating sides of the path:
    "We received your report" — "Mar 19, 2:30 PM" with checkmark
    "An analyst is reviewing" — "Mar 20, 9:15 AM" with checkmark
    "Officer Kwame is on it" — "Mar 21, 11:00 AM" (current, pulsing)
  • Small illustrated icons at each node (envelope for received, person for assigned, magnifier for investigating)
  • The path draws itself with animation on page load (1.5s total, top to bottom)
  • Progress percentage: "50% Complete" with a thin horizontal progress bar below the header that fills with teal gradient

ASSIGNED OFFICER CARD: Glass card with:
  • Officer avatar (or shield icon for anonymity) with teal ring
  • "Officer Kwame M." in friendly bold text
  • Specialty badge: "Phishing Expert" with small star rating
  • "Working on your case since Mar 21" in warm subtitle
  • "Message Officer" button with chat bubble icon, rounded pill shape in teal
  Card slides in from left with 0.3s animation.

EVIDENCE SECTION: Expandable card labeled "Your Evidence (3 files)" with chevron that rotates on expand:
  • Thumbnails in horizontal scroll with rounded corners and subtle shadow
  • Each has file type badge and upload status checkmark
  • "Add More Evidence" button with animated plus icon
  • Expand animation: card height grows with spring physics, thumbnails fade in with stagger

INVESTIGATION NOTES: Timeline of updates from the investigating team, each as a warm glass card:
  • Officer avatar, name, timestamp
  • Note text in friendly language: "I've identified the phishing source and reported it to the hosting provider. Waiting for takedown confirmation."
  • Cards slide in with stagger animation from left, newest at top with a "NEW" amber badge that fades after 5s

ACTIONS: Bottom section with rounded pill buttons:
  • "Add Information" (pencil icon, amber outline)
  • "Contact Support" (headset icon, teal outline)
  • "Share Case" (share icon, muted outline)
  Each button has a subtle hover glow and press scale animation.

RESOLUTION CELEBRATION: When case reaches RESOLVED, special animation triggers — the sunrise in the background intensifies to 15% opacity with warm glow, golden confetti particles cascade down, and a large friendly message appears: "Great news! Your case has been resolved." with details card below.

Typography: Warm rounded font, generous spacing. Hopeful, reassuring language throughout.
Mobile portrait, dark mode, journey and progress feeling.""")

doc.add_page_break()

# ============================================
# SECTION 6: Education Hub
# ============================================
add_section_header(doc, '6', 'Education Hub — Learning Garden', 'Cybersecurity learning platform with categories, progress, and daily tips')

add_prompt(doc, """Design a warm, futuristic education hub screen for "Rethics AI" cybersecurity app. It should feel like entering a nurturing knowledge garden.

BACKGROUND: Deep warm dark #0E0B09. An illustrated African village scene at 5% opacity — round huts, a large tree with people gathered beneath it (symbolizing community learning), warm amber sunset glow. Animated: tree leaves sway gently (subtle transform rotation 0-2 degrees, 5s loop), firefly dots of amber light float around. A soft gradient glow from bottom center in warm gold at 6%.

HEADER: "Learn Digital Safety" in large friendly font with a book-and-shield icon that has a page-turning animation on load. Subtitle: "Knowledge is your best protection". Search bar below: rounded glass input with magnifier icon and animated placeholder cycling: "Search topics..." then "Try 'password safety'..." then "Try 'mobile money tips'..." (fade transitions, 4s each).

LEARNING STREAK: Motivational card at top with:
  • Flame icon with animated flicker (color shifts amber to gold)
  • "7-Day Learning Streak!" with streak count animation
  • Mini calendar showing last 7 days as dots (teal for completed, amber for today, gray for future)
  • "Keep going! You're in the top 15% of learners" in encouraging text
  Card has warm amber border glow and slides in from top.

CONTINUE LEARNING: Horizontal scrollable section with in-progress course cards:
  • Each card: 160px wide, rounded, with category background illustration (soft, hand-drawn style)
  • Course title in bold, lesson count "4 of 8 lessons"
  • Circular progress ring with percentage — ring is animated, filled portion in teal with gradient
  • "Continue" text in amber at bottom
  • Cards slide in from right with stagger on load
  • Active card (leftmost) slightly larger with shadow lift

CATEGORY GRID: "Explore Topics" section header. 2-column grid of illustrated category cards, each with:
  • Beautiful custom illustration as card background (not generic icons — warm, African-context art):
    - "Scam Prevention" — illustrated hand stopping a hook, warm red-amber palette
    - "Password Security" — illustrated golden lock with keys, teal palette
    - "Mobile Money Safety" — illustrated M-Pesa style phone, amber-gold palette
    - "Social Media Privacy" — illustrated shield over chat bubbles, blue palette
    - "Safe Browsing" — illustrated compass/map, teal-green palette
    - "Family Safety Online" — illustrated family under umbrella, warm purple palette
  • Category name in bold white over a dark gradient overlay at bottom
  • Lesson count badge in corner: "8 lessons" in rounded pill
  • Completion indicator: thin progress bar at bottom of card, filled in teal
  • Cards fade-in-up with stagger (0.1s intervals) on page load
  • On tap: card scales to 0.97 with spring-back, navigates to category

DAILY SAFETY TIP: Featured card with:
  • Lightbulb icon with animated glow pulse in amber
  • "Today's Tip" header
  • Practical tip text: "Always verify M-Pesa messages by calling Safaricom directly. Scammers can fake sender IDs."
  • "Learn More" in teal
  • Card has an amber left border stripe (3px) with subtle glow
  • Swipe left/right to see previous tips with smooth card transition

ACHIEVEMENT BADGES: Horizontal scroll of earned badges as small circular icons with:
  • Illustrations: shield, brain, lock, eye, star designs
  • Earned badges in full color with subtle shine animation (highlight sweep)
  • Unearned badges in dark silhouette with "?" overlay
  • Tapping earned badge shows details with pop animation

COMMUNITY STATS: Bottom card with warm glass-morphism:
  • "12,450 learners across Africa" with animated counter
  • Mini Africa map with dots representing learners, soft pulse animations at random points
  • "Together we're safer" tagline

Typography: Warm, rounded, generous spacing. Encouraging and educational tone.
Mobile portrait, dark mode, nurturing and empowering.""")

doc.add_page_break()

# ============================================
# SECTION 7: Emergency Contacts
# ============================================
add_section_header(doc, '7', 'Emergency Contacts — Lifeline', 'Country-specific emergency resources with quick-dial SOS functionality')

add_prompt(doc, """Design a warm, futuristic emergency contacts screen for "Rethics AI" cybersecurity app. It should feel like a lifeline — urgent but not panic-inducing.

BACKGROUND: Deep dark #0C0908 with a subtle illustrated background of interconnected hands forming a safety net pattern at 4% opacity in warm tones. A soft radial gradient in crimson-to-amber pulses very slowly from the top SOS area (8s loop, 3% to 6% opacity). Subtle floating particles in warm red and amber drift slowly.

SOS HEADER: Large prominent section at top:
  • Circular SOS button: 100px diameter, crimson red #CD5C5C background with 3 concentric ring animations pulsing outward (ripple effect, 2s intervals, fading as they expand)
  • "EMERGENCY SOS" text below in bold, warm white
  • Subtitle: "Tap for immediate help" in lighter text
  • On tap: button scales down then up (press feedback), rings intensify, and a quick-dial sheet slides up with the most critical numbers pre-loaded
  • A soft heartbeat animation on the button (subtle scale pulse 1.0-1.02, 1s loop) to draw attention without causing anxiety

COUNTRY SELECTOR: Rounded glass-morphism card below SOS:
  • Current country flag icon (animated flag wave — subtle transform skew, 4s loop) + country name "Kenya"
  • Dropdown arrow that rotates with animation on tap
  • On tap: bottom sheet slides up with list of 16 countries, each with flag + name, scrollable with search. Selected country has amber highlight with checkmark pop animation
  • Switching country: contacts below transition with a smooth crossfade (0.3s)

CONTACT CATEGORIES: Vertically stacked expandable sections, each as a warm glass card with category-colored left border:

  "Police and Cyber Crime" (border: amber, shield icon):
  • National Cyber Crime Unit — phone number, tap-to-call icon with ripple animation
  • Local Police — phone number, 24/7 badge in green with pulse dot
  • Each contact row has: organization name, description in small text, phone number in monospace, call button (phone icon in circle), WhatsApp button if available (green circle)
  • Call button has a ring animation on tap before initiating call

  "Fraud and Financial" (border: gold, bank icon):
  • Central Bank Fraud Line
  • Mobile Money Provider (M-Pesa/MTN/Airtel based on country)
  • Banking Ombudsman

  "Mental Health and Support" (border: teal, heart icon):
  • Crisis Helpline — prominent with warm highlight, "Free and Confidential" badge
  • Cyberbullying Support
  • Gender-Based Violence Hotline

  "Child Protection" (border: purple, child-shield icon):
  • Child Helpline
  • Online Safety for Kids

  Each section: chevron rotates 180 degrees with animation on expand/collapse. Contact rows stagger fade-in on expand (0.1s intervals). Collapsed sections show count: "4 contacts". First section auto-expanded on load.

SAVE OFFLINE TOGGLE: Floating card at bottom:
  • "Available Offline" toggle with download cloud icon
  • When enabled: smooth toggle slide in teal, download animation (cloud icon rains dots downward, then checkmark appears)
  • "All contacts saved to your device" confirmation text fades in
  • This ensures contacts work without internet in emergencies

OPERATING HOURS INDICATOR: Each contact shows:
  • Green dot + "Available now" if within hours (dot breathes)
  • Amber dot + "Opens at 8:00 AM" if outside hours
  • Blue dot + "24/7" for always-available lines (dot has continuous soft pulse)

QUICK ACTIONS BAR: Bottom-pinned glass bar with:
  • "Call Police" (red, phone icon)
  • "Report Online" (amber, file icon)
  • "Chat Support" (teal, chat icon)
  Each as rounded pill buttons with press animations.

Typography: Clear, readable, warm. Critical numbers in larger font. Calm but authoritative tone.
Mobile portrait, dark mode, reassuring lifeline feeling.""")

doc.add_page_break()

# ============================================
# SECTION 8: Onboarding
# ============================================
add_section_header(doc, '8', 'Onboarding — Story of Protection', 'Cinematic 4-screen welcome flow with rich illustrations and parallax transitions')

add_prompt(doc, """Design a 4-screen cinematic onboarding flow for "Rethics AI", African cybersecurity app. Each screen tells a story chapter with rich illustration and animation.

SHARED ELEMENTS: All screens have smooth horizontal swipe transitions with parallax depth (background moves at 40% speed, foreground at 100%). Dot pagination at bottom: active dot is an amber dash (16px wide) that morphs from/to circle with spring animation. Skip button top-right in muted text. Background base: warm charcoal #0E0B09.

SCREEN 1 — "The Digital World Is Beautiful":
  • BACKGROUND IMAGE: A rich illustrated African cityscape at night — buildings with glowing windows, cell towers with signal waves, people walking with phones, a starry sky. Art style: warm flat illustration with soft textures. Colors: deep blues, warm ambers, teal accents.
  • ANIMATION: Stars twinkle (randomized opacity pulse), signal waves animate outward from cell towers (concentric arcs expanding and fading), phone screens in people's hands glow with tiny animated content. A shooting star crosses the sky every 8s.
  • TEXT (bottom third, over dark gradient overlay):
    Title: "The digital world connects us" in large warm font, fades in at 0.5s
    Subtitle: "But it also brings new dangers. Scams, fraud, and cyber threats affect millions across Africa every day." fades in at 1.0s
  • SWIPE HINT: Animated hand-swipe gesture icon at bottom, fades out after 3s

SCREEN 2 — "Meet Your Guardian":
  • BACKGROUND IMAGE: Wilson (the AI) illustrated as a friendly glowing shield character standing protectively in front of a family (mother, child, elder) in an African setting — a home compound with warm lighting. Wilson emanates a protective teal glow dome around the family.
  • ANIMATION: Wilson's protective dome pulses gently (teal glow, 3s loop). The family's phones show green shield icons that pop in sequentially. Small threat icons (skull, hook, bug) approach from the edges but dissolve when touching the dome — with small burst particle effects.
  • TEXT:
    Title: "Meet Wilson, your AI guardian" fades in
    Subtitle: "Wilson helps you spot scams, report incidents, and stay safe online. Available 24/7 in 11 languages." fades in with slight delay

SCREEN 3 — "Protection Tools":
  • BACKGROUND IMAGE: An illustrated command center / toolkit scene — a large glowing table-map of Africa surrounded by floating holographic tool icons: magnifying glass (scanner), document (reports), shield (protection), book (education), phone (emergency). Warm ambient lighting, African patterns on the walls.
  • ANIMATION: Tool icons float with gentle bobbing motion (translateY, staggered 3-5s loops). When each tool is "highlighted" in sequence (every 2s), it scales up 1.3x with glow burst and a label appears: "Threat Scanner", "Incident Reports", etc. The Africa map has gentle pulsing dots at country locations. Connecting lines draw between tools showing the ecosystem.
  • TEXT:
    Title: "Everything you need in one place" fades in
    Subtitle: "Report incidents. Scan suspicious links. Learn safety skills. Access emergency help. All from your phone." fades in

SCREEN 4 — "Join the Movement":
  • BACKGROUND IMAGE: A sunrise over the African continent — illustrated from a slight aerial view. The continent is formed by thousands of tiny connected dots (representing users) linked by thin glowing lines forming a safety network. Sunrise colors: amber, gold, warm orange. Country locations have flag-colored node dots.
  • ANIMATION: The network of dots draws itself progressively from east Africa outward (2s animation), connections form with traveling light pulses along the lines. The sunrise glow intensifies slowly behind the continent. User count number at the bottom counts up from 0 to "12,450+ users" with animated digits. Confetti-like particles in amber and teal rise from the continent.
  • TEXT:
    Title: "Built for Africa, by Africa" fades in
    Subtitle: "Join thousands protecting themselves and their communities. Your safety journey starts now." fades in
  • CTA BUTTON: Large rounded amber button "Get Started — It's Free" with animated shine sweep (highlight moves left to right, 3s loop). On press: button expands to fill screen width, text morphs to "Welcome!" with celebration burst.

Typography: Warm rounded font, large titles (28px), comfortable subtitles (16px). Emotional, human language.
Mobile portrait, dark mode, cinematic storytelling.""")

doc.add_page_break()

# ============================================
# SECTION 9: Admin Dashboard
# ============================================
add_section_header(doc, '9', 'Admin Dashboard — Mission Control', 'Admin nerve center with live metrics, incident feed, and attention queue')

add_prompt(doc, """Design a warm but powerful futuristic admin dashboard for "Rethics AI" cybersecurity platform. It should feel like a high-tech mission control that's still human and approachable.

BACKGROUND: Deep dark #0A0806. An illustrated subtle background showing a large digital shield protecting an African continent outline, rendered in warm amber wireframe at 4% opacity. Animated: data stream particles flow along the continent outline slowly (tiny dots traveling the path, amber color). A soft radial gradient in the center area in dark teal at 5% providing depth. Faint hex grid pattern at 2% opacity covering the screen.

ADMIN HEADER: Glass-morphism top bar:
  • Left: admin avatar with gold ring (pulsing gently, indicating super-admin status), "Commander Amara" in bold, role badge "SUPER ADMIN" in gold with subtle shine animation
  • Right: notification bell with animated badge counter (number scales up with spring on new notification), system health indicator — 3 small circles (green/green/amber) representing Services/Database/AI, each with breathing animation matching their status color
  • Below: greeting "Good evening. 3 items need your attention." with attention items count in amber, animated underline

LIVE METRICS ROW: 4 glass cards in horizontal scroll (card width 40% of screen for peek effect):
  • "Active Users" — person icon, large "2,847" with counting animation, mini sparkline graph (last 7 days) drawn with teal line animation, "up 12% this week" in green with up-arrow bounce animation
  • "Open Cases" — folder icon, large "23", color shifts to amber if >20 or red if >50 with pulse. Mini bar chart shows cases by priority (stacked: red/amber/teal)
  • "Threats Today" — radar icon, large "156" with counting animation, mini world map dots showing threat origins, "47 blocked automatically" in teal
  • "Resolution Rate" — checkmark icon, "78%" with animated circular progress ring filling with teal, "Avg: 4.2 days" subtitle
  Cards stagger fade-in-up on load (0.15s intervals). Each card has a subtle background illustration matching its metric at 8% opacity.

ATTENTION QUEUE: "Needs Your Action" section with amber header glow:
  • 3 priority cards stacked vertically, each with animated left border:
    - Unassigned HIGH priority case — red left border pulsing, case info, "Assign Now" button with amber glow
    - Officer requesting backup — amber left border, case details, "Review" button
    - Escalated case from user — red left border, complaint summary, "Respond" button
  • Cards have a subtle shake micro-animation on the first one (1px horizontal, 0.5s, once on load) to draw attention
  • Swiping right on a card reveals quick-action buttons: "Assign", "Escalate", "Dismiss" with slide-reveal animation

INCIDENT FEED: "Live Incidents" with a red breathing dot and "LIVE" badge:
  • Vertical list of recent incidents as glass rows, newest slides in from the top with a smooth push-down animation that moves existing items
  • Each row: case number monospace, incident type icon + label, country flag, priority dot (colored, pulsing if high), reporter avatar (silhouette for anonymity), "2 min ago" timestamp
  • Tapping a row: expands inline with slide-down animation showing summary, evidence count, and action buttons
  • New incidents appear with a soft flash animation (row background briefly highlights in amber then fades)

CHARTS SECTION: Swipeable analytics cards:
  • Card 1: "Incidents This Week" — animated line chart drawing left to right (1s), teal line with gradient fill, data points glow on completion
  • Card 2: "By Type" — animated donut chart assembling segment by segment (0.3s per segment), each segment in its category color. Legend beside with labels and counts
  • Card 3: "By Country" — mini Africa map heat-map, countries color-fill animated based on volume, warm color scale (amber low to red high)
  Page dots below with slide indicator.

QUICK ACTIONS: Floating action area at bottom:
  • 4 circular glass buttons in a row with labels below:
    "Assign Cases" (person+ icon, amber)
    "Send Alert" (megaphone icon, red)
    "View Reports" (chart icon, teal)
    "Manage Users" (people icon, gold)
  • Icons have a gentle float animation (translateY 2px, staggered 3-4s loops)
  • On tap: icon scales with spring physics, navigates to respective section

SYSTEM STATUS FOOTER: Thin glass bar at very bottom:
  • "All systems operational" with green dot if healthy
  • "Firebase: Connected" | "AI: Active" | "Scanner: Online" — small status tags
  • Last sync time with animated refresh icon

Typography: Clean but warm rounded font. Confident but friendly tone. Data is large and scannable.
Mobile portrait, dark mode, powerful yet welcoming command center.""")

doc.add_page_break()

# ============================================
# SECTION 10: User Management
# ============================================
add_section_header(doc, '10', 'User Management — People Hub', 'Community member management with role control and activity tracking')

add_prompt(doc, """Design a warm, futuristic user management screen for "Rethics AI" admin panel. It should feel like managing a community of protected people, not cold database entries.

BACKGROUND: Deep dark #0B0907. Illustrated background of connected people silhouettes forming a network/constellation pattern across Africa at 4% opacity in warm amber. Animated: connection lines between people pulse with traveling light dots (slow, 8s loops). Soft radial gradient from center in teal at 3%.

HEADER: "Community Members" in warm bold font with people-shield icon. User count "2,847 members" with animated counter. Growth indicator: "+48 this week" in teal with animated up arrow.

SEARCH AND FILTER BAR: Rounded glass input with:
  • Magnifier icon, placeholder "Search by name, email, or country..."
  • On focus: bar expands slightly with smooth width animation, amber border glow appears
  • Filter chips below: scrollable row of pill buttons — "All", "Admins", "Officers", "Verified", "New", "Flagged". Active chip filled with amber, inactive outlined. Tapping animates: chip fills with color sweep from left.

USER CARDS: Vertical list of glass-morphism cards, each representing a user:
  • LEFT: Avatar with role ring — gold ring for admin (shine animation), teal for verified, gray for standard, red pulse for flagged accounts
  • CENTER:
    - Name in bold warm white: "Amara Okafor"
    - Email partially masked: "am***@gmail.com" with eye icon to reveal (tap: text unmasks with typewriter animation)
    - Country flag icon + "Nigeria"
    - Member since: "Joined 6 months ago"
  • RIGHT:
    - Activity dot: green breathing "Active", amber "3d idle", gray "30d+ inactive"
    - Mini stats row: "5 cases" | "12 scans" | "82% learned" in tiny text
  • BOTTOM of card (subtle divider):
    - Role badge pill: "User" (gray), "Admin" (gold), "Officer" (teal) with icon
    - Last active: "2 hours ago" in muted text
  • Cards stagger fade-in-up on load (0.08s intervals)
  • On tap: card expands with smooth spring animation revealing detail panel:
    - Activity sparkline graph (last 30 days engagement)
    - Full stats: cases filed, threats reported, education completion %, account age
    - Role history timeline: "User then Officer (promoted by Admin Kwame, Jan 15)"
    - Action buttons as warm rounded pills:
      "Change Role" (crown icon, gold outline)
      "View Cases" (folder icon, teal outline)
      "Send Message" (chat icon, amber outline)
      "Suspend" (pause icon, red outline, requires long-press with confirmation)

ROLE CHANGE FLOW: Tapping "Change Role" opens a bottom sheet:
  • Glass modal slides up with spring physics
  • Role options as large illustrated cards: User (person icon, gray), Officer (badge icon, teal), Admin (crown icon, gold), Super Admin (star-crown icon, platinum) — each with permission summary text
  • Current role highlighted with glow border
  • Selecting new role: card pulses, confirmation dialog slides up: "Promote Amara to Officer?" with "Confirm" (teal) and "Cancel" buttons
  • On confirm: role badge morphs with color transition animation, success checkmark pop

BULK ACTIONS: Long-press a card to enter selection mode:
  • Checkbox circles appear on all cards with pop animation
  • Selected cards get amber glow border
  • Bottom action bar rises with slide-up: "3 selected" counter, action buttons: "Change Role", "Send Message", "Export", "Suspend"
  • "Select All" toggle in the bar

FLAGGED USERS ALERT: If flagged users exist, a warm warning card at the top:
  • Amber border, "3 flagged accounts need review" with animated attention dot
  • "Review Now" button with arrow slide animation
  • Tapping filters list to show only flagged users with crossfade

Typography: Warm, readable. Users are "members" not "entries". Respectful, community language.
Mobile portrait, dark mode, community management feeling.""")

doc.add_page_break()

# ============================================
# SECTION 11: Case Triage
# ============================================
add_section_header(doc, '11', 'Case Triage & Assignment — Dispatch Center', 'Kanban-style case management with drag-to-assign officer functionality')

add_prompt(doc, """Design a warm, futuristic case triage screen for "Rethics AI" admin panel. Admins assign and manage cases here — it should feel like a responsive dispatch center.

BACKGROUND: Deep dark #0A0806. Illustrated background of a stylized African road network connecting cities/villages at 4% opacity — roads glow softly in amber representing case routing. Animated: small light dots travel along the roads periodically (like dispatched units, 6s loops). Soft gradient in the case columns area.

STATS RIBBON: Top glass bar with live metrics:
  • "Unassigned: 7" (amber, pulses if >5) | "In Progress: 16" (teal) | "Overdue: 2" (red, pulse animation) | "Avg Resolution: 4.2 days"
  • Numbers use counting animation on load
  • Overdue count has a subtle shake animation (1px, once every 5s) to draw attention

KANBAN COLUMNS: Horizontal swipeable layout with 3 columns, each as a section:

  "INCOMING" (amber header glow, count badge):
  • Stack of case cards needing assignment
  • Each card: dark glass with amber left border (3px)
    - Case number monospace, incident type icon + label
    - Country flag + reporter (anonymized)
    - Time since filed: "4 hours ago" — text color shifts from white to amber to red as time increases
    - Priority badge: pill with color (teal Low, amber Medium, red High with pulse)
    - Evidence badge: "3 files" with paperclip icon
    - "ASSIGN" button at bottom — amber pill, scales on press

  "IN PROGRESS" (teal header glow):
  • Cards have teal left border
  • Additional info: assigned officer avatar + name, days active
  • Progress indicator: thin bar at bottom showing investigation stage

  "RESOLVED" (green header glow):
  • Cards have green left border, more compact
  • Resolution summary, time to resolve, officer who handled
  • "Close Case" button in muted outline

  Column transitions: dragging a card between columns shows a glowing guide area (drop zone highlights with amber border), card has lift shadow animation while dragging. On drop: card slides into new column with spring physics, column count badges update with pop animation.

OFFICER PANEL: Collapsible bottom drawer (swipe up from handle):
  • Handle bar with "Available Officers (5)" text
  • Peek state shows top 2 officers. Full expansion shows all.
  • Each officer row:
    - Avatar with status ring (green=available, amber=busy, gray=offline)
    - Name and specialty tags: "Fraud Expert", "Phishing Specialist" in small pills
    - Current caseload: horizontal bar that fills — teal when manageable, amber when heavy, red when overloaded. Animated fill on panel open.
    - Active case count: "4 cases"
  • DRAG-TO-ASSIGN: long-press officer avatar, drag to a case card in INCOMING column
    - Connection line draws from officer to case with animated dashes
    - Drop on case: burst animation, case card updates with officer avatar, slides to IN PROGRESS column
    - Officer caseload bar updates with smooth fill animation

CASE DETAIL MODAL: Tapping any case card opens detail:
  • Glass modal slides up from bottom covering 80% screen
  • Full case info: type, description, reporter details, evidence thumbnails (horizontal scroll with rounded corners), timeline of events
  • "Assign To" dropdown: officer list with availability and caseload
  • Priority change: 3 selector buttons with color transition animation
  • "Add Note" expandable text field with glass styling
  • Action buttons: "Assign" (amber), "Escalate" (red), "Close" (green)
  • Close button (x) top-right with rotation animation on press

ESCALATION: Red "ESCALATE" button in the detail modal:
  • On press: button expands, changes to "Escalating..." with alert wave animation
  • Sends priority notification to super admins
  • Case card border shifts to red with pulse
  • "Escalated" badge appears on the card with slide-in animation

OVERDUE INDICATOR: Cases open >48 hours without assignment:
  • Card border shifts to red with pulse animation
  • "OVERDUE" badge appears in corner with urgent dot animation
  • These cards float to the top of INCOMING column

Typography: Clear command language but warm. Cases are treated as people's problems, not tickets.
Mobile portrait, dark mode, responsive dispatch center feeling.""")

doc.add_page_break()

# ============================================
# SECTION 12: Threat Database
# ============================================
add_section_header(doc, '12', 'Threat Database Manager — Vault', 'Verified threat intelligence curation with bulk import and severity management')

add_prompt(doc, """Design a warm, futuristic threat database management screen for "Rethics AI" admin panel. Admins curate the verified threat intelligence here.

BACKGROUND: Deep dark #090706. Illustrated background of a digital vault door with African shield pattern engravings at 4% opacity. Animated: vault door has a very subtle lock mechanism rotation (gear icon, 20s loop). Digital data streams flow vertically at the edges in amber at 3% opacity — columns of characters falling slowly (gentle, not aggressive). Soft gradient glow behind the content area in teal.

VAULT HEADER:
  • "Threat Vault" in bold with animated vault-shield icon (shield assembles from 4 segments on page load, 0.6s)
  • "14,892 verified threats" with counting animation
  • Last updated: "Updated 2 hours ago" with refresh icon (rotates on tap to sync)

SEARCH AND FILTERS:
  • Glass search bar with animated scan-line sweep when active (thin light bar moves left to right, 2s)
  • Threat type filter chips: "All", "URLs", "Emails", "Phones", "Domains", "IPs" — each with icon, active chip filled with category color, inactive outlined
  • Severity filter: "ALL | CRITICAL | HIGH | MEDIUM | LOW" tabs with bottom indicator line that slides with spring physics, count badges on each

ADD THREAT: Floating action button bottom-right — amber circle with "+" icon:
  • On tap: "+" rotates 45 degrees to become "x" (close), add panel slides up as glass bottom sheet
  • Add form:
    - Threat type selector: horizontal icon buttons with labels, selected one scales up with glow
    - Threat value input: glass text field, placeholder "Enter URL, email, phone, or domain..."
    - Severity slider: custom track with gradient from teal (left/low) through amber (medium) to red (right/critical), thumb shows current level text label, track color updates in real-time as dragged
    - Source selector: dropdown with options "Manual Entry", "Community Report", "VirusTotal", "Partner Feed"
    - Notes: multiline glass textarea
    - "Add to Vault" button: amber, on press shows encryption animation (lock icon + orbiting dots) then success checkmark pop

THREAT ENTRIES LIST: Each entry as a glass card row:
  • LEFT: threat type icon in colored circle (link=teal, email=amber, phone=gold, domain=blue)
  • CENTER:
    - Threat value in monospace (defanged for safety display)
    - Source badge: small pill "VirusTotal" (green), "Community" (amber), "Manual" (gray) with subtle glow
    - First detected and Last seen dates
    - Detection count: "Flagged 847 times" with mini horizontal bar chart, bar fills with red proportional to count
  • RIGHT:
    - Severity badge: colored dot with label — "CRITICAL" (red, pulsing), "HIGH" (amber), "MEDIUM" (yellow), "LOW" (teal)
    - Action icons: Edit (pencil), Toggle Active/Inactive (power icon), Delete (trash)

  Cards stagger fade-in on load. On swipe-left: action buttons reveal with slide animation.

  Tapping a card expands it with spring animation showing:
    - Full threat details, all detection sources
    - Timeline of detections
    - Related threats: linked entries with connection lines
    - Geographic impact: mini map showing where this threat was detected most
    - Edit and Delete buttons with confirmation modals

BULK IMPORT: "Import CSV" button in header:
  • On tap: upload modal with file picker
  • File selected: animated progress bar fills with data-stream effect (moving dashes)
  • Preview: table showing first 5 entries to verify, "Import 250 threats" confirmation button
  • Import animation: entries count up rapidly, progress percentage, success burst on completion

THREAT STATS CARD: Expandable summary at top:
  • "This Month: +342 new threats, 89 deactivated, 12 false positives removed"
  • Mini trend chart: new threats per day sparkline
  • Top threat type breakdown: small horizontal bars with category colors

Typography: Technical but accessible. Monospace for threat values, rounded for UI text. Professional, protective tone.
Mobile portrait, dark mode, secure vault aesthetic.""")

doc.add_page_break()

# ============================================
# SECTION 13: Analytics
# ============================================
add_section_header(doc, '13', 'Analytics & Reports — Intelligence Brief', 'Comprehensive analytics with charts, geographic distribution, and exportable reports')

add_prompt(doc, """Design a warm, futuristic analytics and reporting screen for "Rethics AI" admin panel. It should feel like reading a beautifully crafted intelligence briefing.

BACKGROUND: Deep dark #0A0806. Illustrated background of an African landscape viewed from above (aerial/satellite style) with grid overlay at 3% opacity, warm amber/brown tones — representing surveillance and oversight. Animated: slow pan movement across the landscape (translateX, 60s loop, very subtle). Faint data visualization elements (chart outlines, numbers) float and fade at edges at 2% opacity.

BRIEFING HEADER:
  • "Intelligence Brief" in elegant spaced font with animated classified stamp effect on load (stamp rotates slightly and lands with a small bounce, opacity from 0 to 1)
  • Date range selector: glass pill with calendar icon, showing "Mar 1 — Mar 22, 2026"
  • Preset chips: "24H | 7D | 30D | 90D | CUSTOM" — active chip has amber underline that slides with spring physics
  • "Generated 2 min ago" with auto-refresh indicator dot

SECTION 1 — KEY METRICS:
  4 hero numbers in glass cards with background illustrations:
  • "847 Incidents" — shield watermark, animated counter, trend arrow "up 12%" in amber with arrow bounce
  • "78% Resolved" — checkmark watermark, animated circular ring filling in teal, "+5% vs last month" in green
  • "4.2 Day Avg" — clock watermark, animated number, target indicator "Target: 5 days — Beating target!" in teal with star icon pop
  • "2,847 Users" — people watermark, counter animation, growth sparkline mini-chart drawn in teal
  Cards stagger fade-in-up (0.12s intervals), counters animate after card is visible.

SECTION 2 — INCIDENT TRENDS:
  • Section header: "Incident Volume" with animated chart icon
  • Beautiful animated line chart on dark background:
    - Primary line (current period) draws left to right (1.2s ease) in amber with glow
    - Comparison line (previous period) draws in muted teal, dashed
    - Data points appear as dots that pop in after line passes them
    - Gradient fill below primary line (amber at 20% to transparent)
    - X-axis: dates. Y-axis: count. Labels fade in after animation
    - Touch interaction: vertical guideline follows finger, tooltip bubble shows exact values
    - Legend: "This Month" (amber dot) vs "Last Month" (teal dot) with toggle visibility

SECTION 3 — THREAT LANDSCAPE:
  • "Top Threat Types" header
  • Animated horizontal bar chart, bars fill from left with liquid animation (0.4s each, staggered):
    1. Online Fraud 34% — amber bar
    2. Phishing 28% — teal bar
    3. Mobile Money 19% — gold bar
    4. SIM Swapping 11% — red bar
    5. Other 8% — gray bar
  • Each bar has category icon on the left, percentage on the right that counts up
  • Tapping a bar: expands below with breakdown details — subtypes, affected countries, trend direction

SECTION 4 — GEOGRAPHIC DISTRIBUTION:
  • "Incidents by Country" header
  • Illustrated Africa map with countries as interactive regions:
    - Countries fill with heat-map color on load (animated, sweeping from east to west): green (few incidents) to amber to red (many incidents), 1.5s total animation
    - Tapping a country: it lifts with shadow, tooltip appears showing exact stats
    - Top 5 countries listed below map with flag icons and bar charts
  • Small pulsing dots at major cities with incident clusters

SECTION 5 — RESOLUTION PERFORMANCE:
  • 3 animated gauges side by side:
    - "Resolution Time": semi-circle gauge with needle that swings to "4.2 days", zones marked Green (<3d) Amber (3-5d) Red (>5d), needle swings with spring physics
    - "Closure Rate": circular ring that fills to 78% in teal with percentage counter
    - "Satisfaction": star arc that fills to 4.1/5 with stars lighting up sequentially, gold glow
  • Gauges animate when scrolled into view

SECTION 6 — TEAM PERFORMANCE:
  • "Officer Leaderboard" header with trophy icon
  • Ranked list with podium-style animation for top 3:
    - #1: Gold highlight, crown icon, larger card
    - #2: Silver highlight, medal icon
    - #3: Bronze highlight, medal icon
    - #4-10: Standard rows with rank, avatar, name, key metrics
  • Performance bars (cases resolved vs assigned) animate as progress fills

EXPORT SECTION: Glass card at bottom:
  • "Generate Report" header with document icon
  • Format options as 3 illustrated cards:
    - PDF (red accent, "Formatted brief")
    - CSV (green accent, "Raw data")
    - JSON (amber accent, "API format")
  • Tapping a format: card scales with select animation, "Generate" button appears
  • Generation animation: progress ring fills, then download arrow drops with bounce

Typography: Elegant but readable. Section headers in spaced caps. Data-focused, clear hierarchy.
Mobile portrait (long scroll with section transitions), dark mode, premium intelligence briefing aesthetic.""")

doc.add_page_break()

# ============================================
# SECTION 14: Education Content Manager
# ============================================
add_section_header(doc, '14', 'Education Content Manager — Knowledge Forge', 'Admin content creation and curation for cybersecurity learning modules')

add_prompt(doc, """Design a warm, futuristic education content management screen for "Rethics AI" admin panel. Admins create and curate cybersecurity learning content.

BACKGROUND: Deep dark #0B0908. Illustrated background of a grand African library interior — tall bookshelves, warm candlelight glow, a large tree growing through the center (tree of knowledge) at 5% opacity. Animated: candle flames flicker gently (opacity pulse 0.7-1.0, randomized 2-4s), tree leaves sway (subtle rotation 0-1.5 degrees, 6s loop), tiny sparkle particles float near the tree (knowledge sparks).

HEADER:
  • "Knowledge Forge" in warm bold font with book-hammer icon that has a subtle smithing spark animation on load
  • Stats row: "48 Modules" | "12 Categories" | "2,340 Completions" — each with counting animation and relevant micro-icon
  • "Create New" button: amber rounded pill with "+" icon that rotates to pencil on hover/press

CONTENT OVERVIEW: Horizontal scrollable stat cards:
  • "Published" (42, teal dot, steady glow)
  • "Drafts" (4, amber dot, breathing pulse — reminding admin to complete them)
  • "Archived" (2, gray dot)
  • "Avg Completion" (73%, mini ring chart)
  Cards slide in from right on load.

CATEGORY MANAGEMENT: Horizontal scrollable category pills:
  • Each pill: category color dot + name + module count
  • "Scam Prevention" (amber, 8), "Password Security" (teal, 6), "Mobile Money" (gold, 7)
  • Active/filtered pill has filled background with smooth transition
  • "+" pill at end to add category, opens quick-add bottom sheet
  • Long-press to reorder: pill lifts with shadow, drag to new position with smooth reflow animation

MODULE GRID: 2-column masonry layout of content cards:
  • Each card: glass-morphism with category color stripe at top (4px)
  • Card background: subtle illustrated thumbnail relevant to the topic at 15% opacity
  • Module title in bold, lesson count in subtitle
  • Difficulty badge: "BEGINNER" (green pill), "INTERMEDIATE" (amber), "ADVANCED" (red)
  • Engagement row: views, completions, rating in small muted text
  • Status dot: teal (published, steady), amber (draft, breathing pulse), gray (archived)
  • Cards stagger fade-in-up on load (0.08s intervals)
  • On tap: card scales 0.97 to 1.0 with spring, navigates to editor
  • On long-press: quick action menu pops up — "Edit", "Preview", "Duplicate", "Archive", "Delete"

MODULE EDITOR (New Screen/Modal):
  • Full-screen glass-morphism editor
  • Title input with amber focus border animation
  • Rich text area with formatting toolbar — toolbar icons have press animations
  • Media upload drop zone with illustrated upload cloud icon
  • Language tabs: 11 language flags in horizontal scroll with translation status rings:
    - Green ring: complete
    - Amber ring: partial
    - Red ring: no translation
    - "Auto-translate" button with magic wand icon and sparkle animation
  • Lesson builder: sortable list with drag handles
  • Preview toggle: phone frame mockup with slide-in animation
  • Publish flow: "Publish" with confirmation then 3-step animation: Validated then Secured then Published

Typography: Warm, creative. Educational and inspiring tone. Content titles prominent.
Mobile portrait, dark mode, creative workshop feeling.""")

doc.add_page_break()

# ============================================
# SECTION 15: Role Management
# ============================================
add_section_header(doc, '15', 'Role Management & System Config — Control Center', 'System settings, role hierarchy, API integrations, and audit logging')

add_prompt(doc, """Design a warm, futuristic system settings and role management screen for "Rethics AI" admin panel.

BACKGROUND: Deep dark #090706. Illustrated background of an intricate African beadwork pattern forming a circuit board design at 4% opacity — traditional patterns merged with technology lines in warm amber/teal. Animated: tiny light pulses travel along the circuit paths occasionally (traveling dots, 8s intervals, random paths). Soft radial gradient from center in very subtle teal.

HEADER: "Control Center" in bold with gear-shield icon. System version "v1.0.2" badge with subtle pulse. "Last modified by Admin Kwame, 2 hours ago" in muted subtitle.

SECTION 1 — ROLE ARCHITECTURE:
  • Visual hierarchy tree with animated connecting lines:
    SUPER ADMIN (platinum/gold node, 60px, crown icon, golden glow ring)
    then ADMIN (gold node, 50px, shield icon, amber glow)
    then OFFICER (teal node, 45px, badge icon, teal glow)
    then USER (gray node, 40px, person icon, subtle outline)
  • Nodes connected by lines that draw themselves on page load (1s animation)
  • Light pulses travel down the hierarchy lines periodically
  • Each node has user count
  • Tapping a node: expands showing permissions with toggle switches

SECTION 2 — CURRENT ADMINS AND OFFICERS:
  • List of privileged users as glass cards with role-colored rings
  • "Change Role" button opens role selector with confirmation flow
  • Role badge morphs with color transition animation on change
  • "Add Admin/Officer" button for new assignments

SECTION 3 — ROLE AUDIT LOG:
  • Expandable timeline of role changes
  • Color-coded: promotions in teal, demotions in amber, suspensions in red
  • Entries slide in with stagger

SECTION 4 — SYSTEM TOGGLES:
  Glass cards for each feature with illustrated backgrounds at 8% opacity:
  • "Wilson AI Assistant" — brain circuit illustration, ON/OFF toggle with teal glow
  • "Threat Scanner" — radar illustration, toggle with radar sweep icon
  • "Push Notifications" — bell illustration, toggle with bell ring animation
  • "Maintenance Mode" — wrench illustration, special warning on enable with amber confirmation
  Each toggle has custom animated switch with spring physics.

SECTION 5 — API AND INTEGRATIONS:
  • Glass cards for each service:
    - OpenAI API: masked key with copy/reveal, status green "Connected", latency display
    - VirusTotal API: similar layout
    - Firebase Services: 4 sub-status dots (Auth, Firestore, Storage, Functions)
  • Unhealthy service: red dot with urgent pulse and "Retry" button

SECTION 6 — DANGER ZONE:
  • Collapsed by default behind "Show Advanced" toggle
  • Red-tinted glass card with warning stripe
  • Destructive options requiring typed confirmation
  • Intentionally slow and deliberate UX for safety

Typography: Technical but warm. Settings organized by importance.
Mobile portrait, dark mode, secure control room feeling.""")

doc.add_page_break()

# ============================================
# SECTION 16: Broadcast Center
# ============================================
add_section_header(doc, '16', 'Admin Broadcast Center — Signal Tower', 'Multi-channel notification and broadcast management with automated alerts')

add_prompt(doc, """Design a warm, futuristic notification and broadcast management screen for "Rethics AI" admin panel. Admins communicate with users from here.

BACKGROUND: Deep dark #0A0807. Illustrated background of a traditional African drum tower/signal fire on a hilltop, with signal waves emanating outward across a landscape at 5% opacity — representing communication across distances. Animated: signal waves expand outward from the tower in concentric arcs (3s intervals, fading as they travel). Fire at the tower top flickers (warm amber glow pulse). Stars in the sky twinkle. Small birds fly across occasionally (subtle silhouettes, 15s intervals).

HEADER: "Broadcast Center" in bold with satellite-drum hybrid icon (a drum with signal waves). "Reach 2,847 members across 16 countries" subtitle with globe icon.

COMPOSE SECTION: Primary glass card for creating broadcasts:

  AUDIENCE SELECTOR: Segmented control with smooth slide animation:
  • "Everyone" (globe icon) — "2,847 people"
  • "By Country" (flag icon) — country checkboxes grid with flags and user counts
  • "By Role" (badge icon) — role checkboxes with counts
  • "Custom" (filter icon) — advanced filter combinations
  Switching between options: content crossfades with 0.3s transition.

  PRIORITY LEVEL: 3 large pill buttons with animated signal strength:
  • "Info" — teal, 1 signal bar, gentle pulse
  • "Warning" — amber, 2 signal bars, moderate pulse
  • "Critical" — red, 3 signal bars, urgent pulse animation
  Selecting a priority: pill expands, signal bars animate, other options dim.

  MESSAGE COMPOSER: Glass textarea with:
  • Title field with amber focus border
  • Body field with rich formatting and character count
  • Preview toggle showing phone notification mockup
  • Emoji support and quick-insert icons

  CHANNEL SELECTION: 3 toggles as icon cards:
  • "Push Notification" — phone icon with pop-up animation
  • "Email" — envelope icon with open/close animation
  • "In-App" — app icon with badge dot animation
  Delivery estimate updates live based on selections.

  SEND BUTTON: Large amber button "Broadcast to 2,847 members"
  • On press: signal wave ripples emanate outward from button
  • Success: button morphs to checkmark with celebration particles
  • Delivery counter ticks up in real-time

BROADCAST HISTORY: "Sent Broadcasts" chronological list:
  • Each as a glass card with priority badge, title, message preview
  • Delivery stats: animated fill bars for Sent/Delivered/Read
  • Channel icons as small badges
  • Expandable for full details and engagement charts

AUTOMATED ALERTS SECTION: "System Alerts" with automation icon:
  • Toggleable rules:
    - "High Priority Case Filed" — Alert all admins
    - "Case Unassigned >24h" — Alert super admins
    - "New Threat Added" — Alert officers
    - "User Registration Spike" — Alert admins
  • Each with trigger frequency sparkline and edit capabilities
  • "Add Rule" button for custom automation

Typography: Clear, commanding but warm. Action-oriented language. Delivery stats prominent.
Mobile portrait, dark mode, communication command center feeling.""")

# --- Tips Page ---
doc.add_page_break()
tips_header = doc.add_heading('Tips for Using These Prompts', level=1)
for run in tips_header.runs:
    run.font.color.rgb = RGBColor(0x2D, 0x1B, 0x14)

tips = [
    ('Iterate on results', 'After generating, ask Stitch to "make it more minimal" or "add more African-inspired elements" or "show me the dark mode variant".'),
    ('Compare directions', 'Run the same screen prompt with different color palettes to explore design directions.'),
    ('Component focus', 'Ask for just one component like "design only the threat level gauge widget" for detailed work.'),
    ('Request states', 'Ask for "show this screen in 4 states: empty, loading, populated, error" for development handoff.'),
    ('Responsive variants', 'Add "also show tablet landscape version" to any prompt for adaptive layout exploration.'),
    ('A/B test moods', 'Try adding "make it warmer and more African" vs "make it colder and more cyberpunk" to the same prompt.'),
    ('Export component sheets', 'Ask Stitch to "show me all the button styles, card styles, and input styles from this design as a component library".'),
]

for title, desc in tips:
    p = doc.add_paragraph()
    run = p.add_run(f'{title}: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0xCC, 0x88, 0x00)
    run = p.add_run(desc)

# --- Color Reference ---
doc.add_paragraph()
color_header = doc.add_heading('Color Reference', level=2)
for run in color_header.runs:
    run.font.color.rgb = RGBColor(0x2D, 0x1B, 0x14)

colors = [
    ('Primary (Rich Charcoal Brown)', '#2D1B14'),
    ('Secondary (Sunset Amber)', '#CC8800'),
    ('Accent (Acacia Green)', '#9CAF88'),
    ('Electric Amber', '#FFAA00'),
    ('Cyber Teal', '#00D4AA'),
    ('Warning Crimson', '#FF3B5C'),
    ('Sahara Gold', '#D4A574'),
    ('Baobab Brown', '#8B4513'),
    ('Savanna Tan', '#DEB887'),
    ('Kilimanjaro', '#5D4037'),
    ('Victoria Blue', '#4A6FA5'),
    ('Clay Red / Error', '#CD5C5C'),
    ('Copper Accent', '#B87333'),
    ('Dark Base', '#0D0A08'),
    ('Warm White Text', '#F5F0EB'),
    ('Muted Secondary Text', '#A09888'),
]

for name, hex_val in colors:
    p = doc.add_paragraph()
    run = p.add_run(f'{name}: ')
    run.bold = True
    run = p.add_run(hex_val)
    run.font.name = 'Consolas'

# Save
output_path = r'C:\Users\LENOVO\Desktop\Rethicsai\rethicsai\Rethics_AI_Google_Stitch_Design_Prompts.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
