"""Generate a DOCX version of README.md for sharing as a static document.

The README itself remains the editable canonical source. This script produces a
formatted Word document that mirrors the content of README.md so it can be
attached, printed, or read offline.

Output: Capstone_Project_README.docx (workspace root)
"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

if sys.stdout.encoding and sys.stdout.encoding.lower() != "utf-8":
    try:
        sys.stdout.reconfigure(encoding="utf-8")
    except Exception:
        pass

OUTPUT = Path(__file__).resolve().parent / "Capstone_Project_README.docx"

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(12)
normal.paragraph_format.line_spacing = 1.15
normal.paragraph_format.space_after = Pt(6)


# ============================================================
# Helpers
# ============================================================
def add_page_number(section):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._element.append(fld_begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run._element.append(instr)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._element.append(fld_end)
    for r in p.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(10)


def add_header(section, text):
    h = section.header
    p = h.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    run.italic = True


def heading(text, level=1):
    sizes = {1: 14, 2: 12, 3: 11}
    spaces_before = {1: 16, 2: 10, 3: 8}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(spaces_before.get(level, 8))
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(sizes.get(level, 12))
    run.bold = True
    return p


def para(text, bold=False, italic=False, size=12, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def mixed_para(*runs, align=None):
    """Paragraph with multiple inline runs (e.g. for bold/italic mid-sentence)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if align is not None:
        p.alignment = align
    for text, bold, italic in runs:
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)
        r.bold = bold
        r.italic = italic
    return p


def bullet(text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    p.paragraph_format.space_after = Pt(3)
    for r in p.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)
    return p


def numbered(text):
    p = doc.add_paragraph(text, style="List Number")
    p.paragraph_format.space_after = Pt(3)
    for r in p.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)
    return p


def code_block(text):
    """Render a fenced code block as a monospaced indented paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.space_before = Pt(4)
    for line in text.splitlines():
        run = p.add_run(line + "\n")
        run.font.name = "Consolas"
        run.font.size = Pt(10)
    return p


def table(headers, rows, header_style=True):
    t = doc.add_table(rows=(1 if header_style else 0) + len(rows), cols=len(headers) if header_style else len(rows[0]))
    try:
        t.style = "Light Grid Accent 1"
    except KeyError:
        t.style = "Table Grid"
    if header_style:
        for i, h in enumerate(headers):
            cell = t.rows[0].cells[i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(h)
            run.font.name = "Arial"
            run.font.size = Pt(11)
            run.bold = True
        offset = 1
    else:
        offset = 0
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + offset].cells[ci]
            cell.text = ""
            # First column slightly bolder if it's a label
            run = cell.paragraphs[0].add_run(str(val))
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
            if not header_style and ci == 0:
                run.bold = True
    doc.add_paragraph()
    return t


def hr():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("___________________________________________________________________________")
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    run.font.size = Pt(10)


# ============================================================
# Footer + header
# ============================================================
for sec in doc.sections:
    add_page_number(sec)
    add_header(sec, "Capstone README")


# ============================================================
# Title
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Capstone Project")
r.font.name = "Arial"
r.font.size = Pt(20)
r.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("AI-Powered Cybercrime Reporting & Scam Detection in West Africa")
r.font.name = "Arial"
r.font.size = Pt(13)
r.italic = True
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()

# Author/term/global-challenge/today block
mixed_para(
    ("Author: ", True, False),
    ("Wilsons Navid Wado Tiwa - BSc Software Engineering, African Leadership University", False, False),
)
mixed_para(
    ("Term: ", True, False),
    ("Final-year capstone, ", False, False),
    ("2-month execution window", True, False),
    (" (target completion ~2026-07-05)", False, False),
)
mixed_para(
    ("Global Challenge: ", True, False),
    ("Millennium Project Challenge 12 - transnational organised crime", False, False),
)
mixed_para(
    ("Today: ", True, False),
    ("2026-05-05 - Sprint 1 begins", False, False),
)

hr()


# ============================================================
# Single source of truth - current status
# ============================================================
heading("Single source of truth - current status", level=1)

table(
    ["Item", "Status"],
    [
        ["App",
         "Rethicssec v1.0.2+3 (Flutter + Firebase), already published - code at mobile/rethicsai/ (consolidated into the workspace 2026-05-05; original backup at C:\\Users\\LENOVO\\Desktop\\rethicsai\\ is read-only)."],
        ["ML model",
         "Not built. Production app uses Vertex AI Gemini 1.5 Flash + OpenAI GPT-4o-mini via Cloud Functions; no custom-trained classifier exists."],
        ["Pilot",
         "Not started. Lagos + Douala per the proposal; revised to one city under the 2-month plan."],
        ["Dissertation",
         "Not started. All 4 unit assignments + Pre-Capstone Major Assessment proposal complete (in docs/)."],
        ["Sprint",
         "S1 - Foundation (reconciliation + codebase recovery + first ML baseline)."],
    ],
)

hr()


# ============================================================
# Workspace map
# ============================================================
heading("Workspace map", level=1)

code_block(
    "Capstone-Project/\n"
    "+-- README.md                this file (status board)\n"
    "+-- .gitignore\n"
    "+-- docs/                    all planning deliverables (Units 1-4, Pre-Capstone proposal)\n"
    "+-- mobile/                  Rethicssec Flutter app (consolidated 2026-05-05)\n"
    "|   +-- PRIVACY_POLICY.md/.docx\n"
    "|   +-- PLAYSTORE_DESCRIPTIONS.md\n"
    "|   +-- rethicsai/           the Flutter project itself (lib/, android/, ios/, functions/, ...)\n"
    "+-- ml/                      machine learning research workspace (the missing piece)\n"
    "|   +-- README.md            research framing + sprint deliverables\n"
    "|   +-- requirements.txt\n"
    "|   +-- data/                raw + processed corpora (gitignored)\n"
    "|   +-- notebooks/           exploration, baselines, ad-hoc eval\n"
    "|   +-- src/                 importable training / preprocessing / eval code\n"
    "|   |   +-- dataset.py\n"
    "|   |   +-- preprocessing.py\n"
    "|   |   +-- baselines.py\n"
    "|   |   +-- llm_baselines.py\n"
    "|   |   +-- eval.py\n"
    "|   +-- models/              saved checkpoints (gitignored)\n"
    "|   +-- reports/             model cards, error analysis\n"
    "+-- meetings/                weekly supervisor meeting notes\n"
    "|   +-- _template.md\n"
    "+-- pilot/                   consent forms, interview guides, analysis\n"
    "+-- dissertation/            final report (chapters, generators)\n"
)

hr()


# ============================================================
# Where the app lives
# ============================================================
heading("Where the app lives (and why it is not in mobile/)", level=1)

para(
    "The Flutter app is now in-tree at mobile/rethicsai/. The Cloud Functions "
    "backend lives inside it at mobile/rethicsai/functions/ - no separate "
    "backend/ folder anymore."
)

para(
    "The original location at C:\\Users\\LENOVO\\Desktop\\rethicsai\\ is kept "
    "untouched as a read-only backup. From 2026-05-05 forward, all editing, "
    "building, and committing happens against the workspace copy. Do not edit "
    "the backup; it exists only as a safety net."
)

para(
    "Build artifacts (build/, .dart_tool/, node_modules/) were excluded during "
    "the copy and will regenerate on first flutter pub get / npm install. The "
    ".git/, .firebase/, and IDE configs travelled with the project so deploys "
    "and Android Studio reopen cleanly."
)

hr()


# ============================================================
# The reconciliation
# ============================================================
heading("The reconciliation that has to happen first", level=1)

para(
    "The Pre-Capstone proposal (docs/coursework/Major-assesment/) and the built app diverge "
    "in three load-bearing ways:"
)

table(
    ["", "Proposal", "Built app"],
    [
        ["Pilot scope", "Lagos + Douala only", "Pan-African positioning"],
        ["Languages", "en, fr, pcm (3)", "11 African languages"],
        ["ML approach",
         "Custom Scikit-learn / TensorFlow classifier, >=85% accuracy on 5 categories",
         "Vertex Gemini Flash + GPT-4o-mini, no measured accuracy, no fixed taxonomy"],
    ],
)

mixed_para(
    ("Sprint 1 of the 2-month window decides which of these to amend (the "
     "proposal) and which to add (the missing custom-ML evaluation). The "
     "recommended framing - ", False, False),
    ("\"LLM-based vs. classical ML scam detection for low-resource West African "
     "languages\"", True, False),
    (" - keeps the existing app and delivers the ML rigor the supervisor "
     "asked for, by treating the deployed LLMs as one arm of a comparative "
     "study.", False, False),
)

para("See ml/README.md for the research question and Sprint deliverables.")

hr()


# ============================================================
# Sprint plan
# ============================================================
heading("Sprint plan (2-month overview)", level=1)

para(
    "The original proposal assumed a 6-month plan. The app is already built, so "
    "most of that compresses cleanly. Pilot scope is the main cut."
)

table(
    ["Sprint", "Days", "Focus", "Exit gate"],
    [
        ["S1 - Foundation", "1-7",
         "Reconciliation memo + supervisor sign-off; codebase recovery (6-month dormancy); ML venv; pull public datasets; PII masking; ~500-example labelled set.",
         "TF-IDF + LR baseline running end-to-end."],
        ["S2 - ML core", "8-14",
         "Add RF baseline; LLM zero-shot eval (Gemini + GPT-4o-mini); per-language confusion matrices; dataset + model cards.",
         "3-approach comparison on held-out test set."],
        ["S3 - App + pilot prep", "15-21",
         "Lock 6-category taxonomy in analyzeSuspiciousContent; classification logging to Firestore; deepfake-aware UI warning; consent forms, interview guide, recruitment, IRB.",
         "App emits structured incidents; pilot kit ready."],
        ["S4 - Pilot launch", "22-35",
         "Pilot in one city, 15-25 users; SUS + Firebase analytics; 5-8 mid-pilot interviews.",
         "Pilot wraps with usable data."],
        ["S5 - Analysis + writing", "36-49",
         "SUS scoring, thematic analysis, on-pilot ML accuracy; dissertation chapters in parallel.",
         "Analysis done; chapters 4-5 drafted."],
        ["S6 - Submit + defend", "50-56",
         "Polish, demo video, defense rehearsal.",
         "Submitted, defended."],
    ],
)

mixed_para(
    ("Cuts from the original year-long plan: ", True, False),
    ("pilot 50-100 users -> 15-25; 2 cities -> 1; ML baselines 4 -> 2 (no "
     "transformer fine-tuning); CMU-Africa data assumed unavailable in time "
     "(still send the email).", False, False),
)

hr()


# ============================================================
# Weekly cadence
# ============================================================
heading("Weekly supervisor meeting cadence", level=1)

para(
    "Use meetings/_template.md. One file per meeting, named YYYY-MM-DD.md. "
    "Every meeting covers four sections:"
)
numbered("Last week - committed vs delivered (with evidence: commit links, plots, screenshots).")
numbered("Blockers.")
numbered("Reading and learning the supervisor assigned (what you took away, follow-up questions).")
numbered("This week - commitments.")

hr()


# ============================================================
# Key documents
# ============================================================
heading("Key documents", level=1)

para(
    "Coursework (Units 1-4 + Pre-Capstone Major Assessment) all live under "
    "docs/coursework/. Active project documents are in docs/ root."
)

table(
    ["File", "What it is"],
    [
        ["docs/coursework/Unit_One_Draft - Copy.docx",
         "Chapter One - intro, problem, objectives (5 SMART), research questions (5), scope, significance."],
        ["docs/coursework/WilsonsNavidWadoTiwa-Unit Two Assignment.docx",
         "Refined Unit 2 draft."],
        ["docs/coursework/unit3/WilsonsNavidWadoTiwa_Unit_Three Assignment.docx",
         "Annotated bibliography (22 sources)."],
        ["docs/coursework/unit4/WilsonsNavidWadoTiwa_Unit Four Assignment.docx",
         "Methodology - mixed-methods, SUS, semi-structured interviews, thematic analysis."],
        ["docs/coursework/Major-assesment/WilsonsNavidWadoTiwa_Pre-Capstone_Research_Proposal_v2.docx",
         "Full Pre-Capstone proposal w/ literature review, system design, all 6 UML diagrams."],
        ["docs/coursework/Major-assesment/diagrams/",
         "Agile model, system architecture, ERD, class, use case, mixed-methods diagrams."],
        ["docs/RECONCILIATION_MEMO.md / docs/WilsonsNavidWadoTiwa_Reconciliation_Memo.docx",
         "Supervisor-facing reconciliation memo with the five D1-D5 decisions."],
        ["docs/onboarding/Capstone_Onboarding_Report.docx",
         "Comprehensive 13-section onboarding report for fresh-eyes context."],
    ],
)

hr()


# ============================================================
# Formatting standard
# ============================================================
heading("Formatting standard (locked across all academic deliverables)", level=1)

para(
    "APA 7. Times New Roman 12pt body. Arial 14pt headings, 12pt subheadings. "
    "Double-spaced. 1-inch margins. File naming: "
    "WilsonsNavidWadoTiwa_<Unit/Section>."
)


# ============================================================
# Save
# ============================================================
doc.save(str(OUTPUT))
print("Wrote:", OUTPUT)
